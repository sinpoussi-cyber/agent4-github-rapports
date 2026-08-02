import argparse
import base64
import copy
import io
import json
import os
import sys
from datetime import date, datetime, timedelta, timezone

from dotenv import load_dotenv

from github_downloader import get_latest_word_reports, get_year_word_reports
from word_comparator import compare_documents
from claude_analyzer import analyze
from report_generator import generate
from generate_note_strategique import generate as generate_note
from generate_fiches_societes import generate as generate_fiches
try:
    from brvm_generateur_rapports import generate_all as generate_rapports_derives
    _GENERATEUR_DISPONIBLE = True
except ImportError:
    _GENERATEUR_DISPONIBLE = False
from email_sender import send_report

# ── Archivage durable + synthèse de période (modules optionnels) ─────────────
try:
    import archive_store
    _ARCHIVE_DISPONIBLE = True
except Exception:
    _ARCHIVE_DISPONIBLE = False
try:
    from period_synthesis import build_synthesis_text, insert_synthesis, build_annual_note
    _SYNTHESE_DISPONIBLE = True
except Exception:
    _SYNTHESE_DISPONIBLE = False

load_dotenv()

GH_REPO = os.getenv("GH_REPO", "sinpoussi/mon-repo")
COMPARISON_FILE = "last_comparison.json"


def log(msg):
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    print(f"[{ts}] {msg}")


# ── Helpers note/fiches ───────────────────────────────────────────────────────

_FREQ_LABELS = {
    "JOUR": "JOURNALIÈRE",
    "HEBDO": "HEBDOMADAIRE",
    "MENSUEL": "MENSUELLE",
    "TRIM": "TRIMESTRIELLE",
    "ANNUEL": "ANNUELLE",
}


def _period_info(freq: str, rapports: list) -> dict:
    """Construit le dict period_info à partir des rapports téléchargés."""
    today = date.today()
    run_dates = []
    for r in rapports:
        d_run = r.get("date_run")
        if d_run:
            run_dates.append(d_run.date() if hasattr(d_run, "date") else d_run)
    date_debut = min(run_dates).strftime("%d/%m/%Y") if run_dates else today.strftime("%d/%m/%Y")
    date_fin = max(run_dates).strftime("%d/%m/%Y") if run_dates else today.strftime("%d/%m/%Y")
    month = today.month
    return {
        "date_debut": date_debut,
        "date_fin": date_fin,
        "nb_seances": len(rapports) if rapports else 1,
        "freq_label": _FREQ_LABELS.get(freq, freq),
        "annee": today.year,
        "semaine": today.isocalendar()[1],
        "trimestre": (month - 1) // 3 + 1,
        "mois": today.strftime("%B %Y"),
    }


def _note_fiches_subject(freq: str, pi: dict) -> str:
    today = date.today()
    if freq == "HEBDO":
        return f"Synthèse Hebdomadaire BRVM — Semaine {pi.get('semaine', today.isocalendar()[1])}"
    if freq == "MENSUEL":
        return f"Bilan Mensuel BRVM — {pi.get('mois', today.strftime('%B %Y'))}"
    if freq == "TRIM":
        return f"Bilan Trimestriel BRVM — T{pi.get('trimestre', (today.month - 1) // 3 + 1)} {pi.get('annee', today.year)}"
    if freq == "ANNUEL":
        return f"Bilan Annuel BRVM — {pi.get('annee', today.year)}"
    return f"Note Stratégique BRVM — {today.strftime('%d/%m/%Y')}"


def _note_fiches_html(freq: str, pi: dict, nb_rapports: int, nb_fiches: int) -> str:
    colors = {"HEBDO": "#0F9D58", "MENSUEL": "#E37400", "TRIM": "#7B1FA2", "ANNUEL": "#1A237E"}
    labels = {"HEBDO": "Synthèse Hebdomadaire", "MENSUEL": "Bilan Mensuel",
              "TRIM": "Bilan Trimestriel", "ANNUEL": "Bilan Annuel"}
    color = colors.get(freq, "#1A73E8")
    label = labels.get(freq, "Note Stratégique")
    return f"""<!DOCTYPE html><html lang="fr">
<head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;background:#f0f2f5;font-family:Arial,sans-serif;">
  <div style="max-width:640px;margin:32px auto;background:#fff;border-radius:10px;overflow:hidden;box-shadow:0 2px 12px rgba(0,0,0,.1);">
    <div style="background:{color};padding:24px 32px;">
      <h1 style="margin:0;color:#fff;font-size:20px;">{label} BRVM</h1>
      <p style="margin:6px 0 0;color:rgba(255,255,255,.85);font-size:13px;">
        Période : {pi.get('date_debut', '—')} → {pi.get('date_fin', '—')}
      </p>
    </div>
    <div style="padding:24px 32px;">
      <table style="border-collapse:collapse;width:100%;max-width:440px;">
        <tr style="background:#f5f5f5;"><td style="padding:8px 12px;border:1px solid #ddd;">Séances analysées</td>
          <td style="padding:8px 12px;border:1px solid #ddd;font-weight:bold;">{nb_rapports}</td></tr>
        <tr><td style="padding:8px 12px;border:1px solid #ddd;">Fréquence</td>
          <td style="padding:8px 12px;border:1px solid #ddd;">{pi.get('freq_label', freq)}</td></tr>
        <tr style="background:#f5f5f5;"><td style="padding:8px 12px;border:1px solid #ddd;">Pièces jointes</td>
          <td style="padding:8px 12px;border:1px solid #ddd;font-weight:bold;">Note Stratégique + {nb_fiches} fiche(s)</td></tr>
      </table>
      <p style="margin-top:20px;color:#666;font-size:13px;line-height:1.6;">
        Ce bilan a été généré automatiquement par l'Agent GitHub Rapports BRVM.<br>
        Consultez les pièces jointes pour l'analyse détaillée.
      </p>
    </div>
    <div style="background:#f5f5f5;padding:12px 32px;text-align:center;font-size:11px;color:#999;">
      Analyse multi-IA : DeepSeek · Gemini · Mistral — Agent GitHub Rapports
    </div>
  </div>
</body></html>"""


# ── Fusion de plusieurs .docx en un seul ─────────────────────────────────────

def _merge_docx(docs_bytes_list: list) -> bytes:
    from docx import Document
    merged = Document()
    for p in list(merged.paragraphs):
        p._element.getparent().remove(p._element)
    for i, doc_bytes in enumerate(docs_bytes_list):
        src = Document(io.BytesIO(doc_bytes))
        if i > 0:
            pb = src.element.body.__class__()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            br_p = OxmlElement("w:p")
            br_r = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            br_r.append(br)
            br_p.append(br_r)
            merged.element.body.append(br_p)
        for element in src.element.body:
            merged.element.body.append(copy.deepcopy(element))
    buf = io.BytesIO()
    merged.save(buf)
    return buf.getvalue()


# ── Sérialisation JSON (datetime → str) ─────────────────────────────────────

def _default(obj):
    if isinstance(obj, datetime):
        return obj.isoformat()
    raise TypeError(f"Type non sérialisable : {type(obj)}")


# ── Modes ────────────────────────────────────────────────────────────────────

def cmd_collect():
    log("=== MODE COLLECT ===")

    log(f"Téléchargement des 2 derniers rapports Word depuis '{GH_REPO}'...")
    rapports = get_latest_word_reports(repo_name=GH_REPO, n=2)

    if len(rapports) < 2:
        log(f"ERREUR : {len(rapports)} rapport(s) trouvé(s), 2 requis. Abandon.")
        sys.exit(1)

    doc1, doc2 = rapports[0], rapports[1]
    log(f"Clés disponibles : {list(doc1.keys())}")
    log(f"Document 1 : {doc1['nom']} ({doc1['date_run']})")
    log(f"Document 2 : {doc2['nom']} ({doc2['date_run']})")

    # ── Archivage durable du rapport journalier source (idempotent par date) ──
    # Les artefacts GitHub Actions expirent (90 j max) ; on conserve une copie
    # dans le dépôt pour bâtir des rapports mensuels/annuels fiables.
    # N'altère JAMAIS le rapport quotidien : échec = simple avertissement.
    if _ARCHIVE_DISPONIBLE:
        try:
            path, action = archive_store.archive_daily_report(
                doc1["nom"], doc1["contenu_bytes"], doc1.get("date_run"))
            log(f"Rapport journalier archivé ({action}) : {path}")
        except Exception as e:
            log(f"AVERTISSEMENT : archivage rapport journalier échoué : {e}")

    log("Comparaison des documents...")
    diff_data = compare_documents(doc1["contenu_bytes"], doc2["contenu_bytes"])
    log(f"Taux de changement : {diff_data['taux_changement']}%")
    log(f"Ajoutés : {len(diff_data['paragraphes_ajoutes'])} | "
        f"Supprimés : {len(diff_data['paragraphes_supprimes'])} | "
        f"Modifiés : {len(diff_data['paragraphes_modifies'])}")

    log("Analyse Claude en cours...")
    analysis = analyze(doc1["nom"], doc2["nom"], diff_data)
    log(f"Score d'importance : {analysis.get('score_importance')}/10")
    log(f"Alerte : {analysis.get('alerte', 'non').upper()}")

    log("--- Résumé exécutif ---")
    print(analysis.get("resume_executif", ""))
    log("--- Changements importants ---")
    for i, c in enumerate(analysis.get("changements_importants", []), 1):
        print(f"  {i}. {c}")
    log("--- Interprétation business ---")
    print(analysis.get("interpretation", ""))

    payload = {
        "doc1_nom":      doc1["nom"],
        "doc2_nom":      doc2["nom"],
        "diff_data":     diff_data,
        "analysis":      analysis,
        "collected_at":  datetime.now(timezone.utc).isoformat(),
        "doc1_bytes_b64": base64.b64encode(doc1["contenu_bytes"]).decode("ascii"),
    }
    with open(COMPARISON_FILE, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2, default=_default)
    log(f"Résultat sauvegardé dans '{COMPARISON_FILE}'.")

    log("Collect terminé. Le rapport sera généré et envoyé immédiatement.")


def cmd_rapport_annuel():
    year = datetime.now(timezone.utc).year
    log(f"=== MODE RAPPORT-ANNUEL — {year} ===")

    log(f"Téléchargement de tous les artefacts de l'année {year}...")
    rapports = get_year_word_reports(year=year, repo_name=GH_REPO)

    if len(rapports) < 2:
        log(f"ERREUR : {len(rapports)} rapport(s) trouvé(s) pour {year}, 2 requis. Abandon.")
        sys.exit(1)

    doc1, doc2 = rapports[0], rapports[-1]
    log(f"Premier rapport : {doc1['nom']} ({doc1['date_run']})")
    log(f"Dernier rapport : {doc2['nom']} ({doc2['date_run']})")
    log(f"Nombre total de rapports dans l'année : {len(rapports)}")

    log("Comparaison premier ↔ dernier document...")
    diff_data = compare_documents(doc1["contenu_bytes"], doc2["contenu_bytes"])
    log(f"Taux de changement : {diff_data['taux_changement']}%")

    log("Analyse Claude en cours...")
    analysis = analyze(doc1["nom"], doc2["nom"], diff_data)
    analysis["nb_runs_annee"] = len(rapports)
    analysis["annee"] = year
    analysis["date_premier_run"] = doc1["date_run"].isoformat() if hasattr(doc1["date_run"], "isoformat") else str(doc1["date_run"])
    analysis["date_dernier_run"] = doc2["date_run"].isoformat() if hasattr(doc2["date_run"], "isoformat") else str(doc2["date_run"])

    log(f"Score d'importance : {analysis.get('score_importance')}/10")

    rapport = generate(doc1["nom"], doc2["nom"], diff_data, analysis, "annuel")
    log(f"Objet : {rapport['subject']}")

    log("Envoi de l'email...")
    ok = send_report(rapport["subject"], rapport["body_html"], rapport["body_text"])
    if ok:
        log("Email annuel envoyé avec succès.")
    else:
        log("ERREUR : échec de l'envoi de l'email.")
        sys.exit(1)


def cmd_rapport(type_rapport):
    log(f"=== MODE RAPPORT — {type_rapport.upper()} ===")

    if not os.path.exists(COMPARISON_FILE):
        log(f"ERREUR : '{COMPARISON_FILE}' introuvable. Lancez d'abord 'collect'.")
        sys.exit(1)

    with open(COMPARISON_FILE, "r", encoding="utf-8") as f:
        payload = json.load(f)

    doc1_nom  = payload["doc1_nom"]
    doc2_nom  = payload["doc2_nom"]
    diff_data = payload["diff_data"]
    analysis  = payload["analysis"]
    collected = payload.get("collected_at", "inconnue")

    log(f"Données chargées depuis '{COMPARISON_FILE}' (collecte : {collected})")
    log(f"Génération du rapport {type_rapport}...")

    rapport = generate(doc1_nom, doc2_nom, diff_data, analysis, type_rapport)
    log(f"Objet : {rapport['subject']}")

    attachments = []
    if type_rapport == "quotidien":
        doc1_b64 = payload.get("doc1_bytes_b64")
        if doc1_b64:
            doc1_bytes = base64.b64decode(doc1_b64)
            today_str = date.today().strftime("%Y%m%d")
            pi_jour = _period_info("JOUR", [])

            log("[Pipeline] Étape 1/2 : Note Stratégique BRVM (JOUR)...")
            try:
                _, note_bytes = generate_note([doc1_bytes], "JOUR", pi_jour)
                note_filename = f"Note_Strategique_BRVM_JOUR_{today_str}.docx"
                attachments.append({"filename": note_filename, "data": note_bytes})
                log(f"Note Stratégique générée : {note_filename}")
            except Exception as e:
                log(f"AVERTISSEMENT : échec note stratégique : {e}")

            log("[Pipeline] Étape 2/2 : Fiches Sociétés — Extraction LLM → Enrichissement → Word...")
            try:
                fiches = generate_fiches([doc1_bytes], "JOUR", pi_jour)
                log(f"[Pipeline] {len(fiches)} société(s) extraite(s) → {len(fiches)} fiche(s) générée(s).")
                if fiches:
                    for fname, fbytes in fiches:
                        attachments.append({"filename": fname, "data": fbytes})
                    log(f"{len(fiches)} fiche(s) jointes individuellement (pas de fusion).")
                else:
                    log("AVERTISSEMENT : aucune fiche générée.")
            except Exception as e:
                log(f"AVERTISSEMENT : échec fiches sociétés : {e}")
        else:
            log("AVERTISSEMENT : doc1_bytes_b64 absent — pièces jointes ignorées.")

        # ── Rapports dérivés désactivés ───────────────────────────────────────
        # Sur demande : on ne génère plus les 7 rapports dérivés
        # (InvestisseursInstitutionnels, BRVM_Interne, Regulateur_AMF_UMOA,
        #  SGI, SGO, CIB_Conseillers, SGP_Patrimoine).
        # Seules la Note Stratégique et les Fiches Sociétés sont produites/jointes.

    log("Envoi de l'email...")
    ok = send_report(
        rapport["subject"],
        rapport["body_html"],
        rapport["body_text"],
        attachments=attachments or None,
    )

    if ok:
        log("Email envoyé avec succès.")
    else:
        log("ERREUR : échec de l'envoi de l'email.")
        sys.exit(1)


# ── Mode note + fiches multi-fréquences ──────────────────────────────────────

# ── Sélection de période (mois/année cibles) ────────────────────────────────
# Par défaut on traite la période COMPLÈTE précédente (le job tourne le 1er) :
#   - mensuel  → le mois qui vient de se terminer
#   - annuel   → l'année qui vient de se terminer
# Surcharge possible pour rejeu/backfill via ARCHIVE_TARGET_YEAR / _MONTH.

def _prev_month(today=None):
    today = today or date.today()
    prev_last = today.replace(day=1) - timedelta(days=1)
    return prev_last.year, prev_last.month


def _target_month():
    y = os.getenv("ARCHIVE_TARGET_YEAR")
    m = os.getenv("ARCHIVE_TARGET_MONTH")
    if y and m:
        return int(y), int(m)
    return _prev_month()


def _target_year():
    y = os.getenv("ARCHIVE_TARGET_YEAR")
    return int(y) if y else date.today().year - 1


def _order_docs(rapports):
    """Retourne (docs_asc, docs_recent_first) en normalisant l'ordre par date_run
    quand elle est disponible (l'archive est croissante, get_latest décroissant)."""
    if all(r.get("date_run") is not None for r in rapports):
        asc = sorted(rapports, key=lambda r: r["date_run"])
    else:
        asc = list(rapports)
    docs_asc = [r["contenu_bytes"] for r in asc]
    return docs_asc, list(reversed(docs_asc))


def _produce_and_send(freq, docs_recent_first, docs_asc, pi, rapports_count,
                      archive_month=None):
    """Génère note (+ synthèse d'évolution) et fiches, archive la note si demandé,
    puis envoie l'email. `docs_recent_first` : plus récent en premier (contrat des
    générateurs). `docs_asc` : ancien→récent (pour la synthèse d'évolution)."""
    attachments = []
    nb_fiches = 0

    log(f"[Pipeline] Étape 1/2 : Note Stratégique BRVM ({freq})...")
    try:
        note_filename, note_bytes = generate_note(docs_recent_first, freq, pi)

        # Synthèse d'ÉVOLUTION sur toute la période (additive, jamais bloquante).
        if _SYNTHESE_DISPONIBLE and len(docs_asc) > 1:
            try:
                synth = build_synthesis_text(docs_asc, freq, pi, is_notes=False)
                note_bytes = insert_synthesis(note_bytes, synth, freq, pi)
                if synth:
                    log(f"Synthèse d'évolution insérée dans la note ({len(synth)} chars).")
            except Exception as e:
                log(f"AVERTISSEMENT : synthèse note ignorée : {e}")

        attachments.append({"filename": note_filename, "data": note_bytes})
        log(f"Note générée : {note_filename}")

        # Archivage de la note mensuelle (brique du rapport annuel).
        if archive_month and _ARCHIVE_DISPONIBLE:
            try:
                ay, am = archive_month
                path, action = archive_store.archive_monthly_note(
                    note_bytes, ay, am, nom=note_filename)
                log(f"Note mensuelle archivée ({action}) : {path}")
            except Exception as e:
                log(f"AVERTISSEMENT : archivage note mensuelle échoué : {e}")
    except Exception as e:
        log(f"AVERTISSEMENT : échec note stratégique : {e}")

    log(f"[Pipeline] Étape 2/2 : Fiches Sociétés ({freq}) — Extraction LLM → Enrichissement → Word...")
    try:
        fiches = generate_fiches(docs_recent_first, freq, pi)
        nb_fiches = len(fiches)
        log(f"[Pipeline] {nb_fiches} société(s) extraite(s) → {nb_fiches} fiche(s) générée(s).")
        for fname, fbytes in fiches:
            attachments.append({"filename": fname, "data": fbytes})
    except Exception as e:
        log(f"AVERTISSEMENT : échec fiches sociétés : {e}")

    if not attachments:
        log("ERREUR : aucune pièce jointe générée. Abandon.")
        sys.exit(1)

    subject = _note_fiches_subject(freq, pi)
    body_html = _note_fiches_html(freq, pi, rapports_count, nb_fiches)

    log(f"Envoi de l'email : {subject}")
    ok = send_report(subject, body_html, attachments=attachments)
    if ok:
        log("Email envoyé avec succès.")
    else:
        log("ERREUR : échec de l'envoi de l'email.")
        sys.exit(1)


def _note_fiches_generic(freq: str):
    """HEBDO / TRIM : derniers N artefacts (comportement historique) + synthèse."""
    log(f"=== MODE NOTE-FICHES — {freq} ===")
    n = {"HEBDO": 7, "TRIM": 90}.get(freq, 7)
    log(f"Téléchargement des {n} derniers rapports Word...")
    rapports = get_latest_word_reports(repo_name=GH_REPO, n=n)
    if not rapports:
        log("ERREUR : aucun rapport disponible. Abandon.")
        sys.exit(1)
    log(f"{len(rapports)} rapport(s) disponible(s).")
    docs_asc, docs_recent_first = _order_docs(rapports)
    pi = _period_info(freq, rapports)
    _produce_and_send(freq, docs_recent_first, docs_asc, pi, len(rapports))


def _note_fiches_mensuel():
    """MENSUEL : UNIQUEMENT les rapports journaliers du mois civil concerné,
    lus depuis l'archive durable. Archive ensuite la note produite."""
    freq = "MENSUEL"
    year, month = _target_month()
    log(f"=== MODE NOTE-FICHES — MENSUEL ({year}-{month:02d}) ===")

    rapports = []
    if _ARCHIVE_DISPONIBLE:
        try:
            rapports = archive_store.get_daily_reports_for_month(year, month)
            log(f"{len(rapports)} rapport(s) journalier(s) archivé(s) pour {year}-{month:02d}.")
        except Exception as e:
            log(f"AVERTISSEMENT : lecture de l'archive échouée : {e}")

    if not rapports:
        log("Archive vide/indisponible — repli sur les 30 derniers artefacts.")
        rapports = get_latest_word_reports(repo_name=GH_REPO, n=30)
    if not rapports:
        log("ERREUR : aucun rapport disponible. Abandon.")
        sys.exit(1)

    docs_asc, docs_recent_first = _order_docs(rapports)
    pi = _period_info(freq, rapports)
    _produce_and_send(freq, docs_recent_first, docs_asc, pi, len(rapports),
                      archive_month=(year, month))


def _note_fiches_annuel():
    """ANNUEL : basé sur les 12 notes mensuelles archivées de l'année concernée."""
    year = _target_year()
    log(f"=== MODE NOTE-FICHES — ANNUEL ({year}) ===")

    notes = []
    if _ARCHIVE_DISPONIBLE:
        try:
            notes = archive_store.get_monthly_notes_for_year(year)
            log(f"{len(notes)} note(s) mensuelle(s) archivée(s) pour {year}.")
        except Exception as e:
            log(f"AVERTISSEMENT : lecture des notes mensuelles échouée : {e}")

    if not notes:
        log("ERREUR : aucune note mensuelle archivée pour l'année demandée. "
            "Le rapport annuel s'appuie sur les notes mensuelles "
            "(voir mode note-fiches-mensuel). Abandon.")
        sys.exit(1)

    if not _SYNTHESE_DISPONIBLE:
        log("ERREUR : module de synthèse indisponible (period_synthesis). Abandon.")
        sys.exit(1)

    pi = {
        "date_debut": f"01/01/{year}",
        "date_fin": f"31/12/{year}",
        "nb_seances": len(notes),
        "freq_label": "ANNUELLE",
        "annee": year,
    }

    log("Construction de la note annuelle à partir des notes mensuelles...")
    note_filename, note_bytes = build_annual_note(notes, year, pi)
    attachments = [{"filename": note_filename, "data": note_bytes}]

    subject = _note_fiches_subject("ANNUEL", pi)
    body_html = _note_fiches_html("ANNUEL", pi, len(notes), 0)
    log(f"Envoi de l'email : {subject}")
    ok = send_report(subject, body_html, attachments=attachments)
    if ok:
        log("Email annuel envoyé avec succès.")
    else:
        log("ERREUR : échec de l'envoi de l'email.")
        sys.exit(1)


def cmd_note_fiches(freq: str):
    """Aiguille vers le bon pipeline selon la fréquence."""
    if freq == "MENSUEL":
        return _note_fiches_mensuel()
    if freq == "ANNUEL":
        return _note_fiches_annuel()
    return _note_fiches_generic(freq)


# ── Point d'entrée ───────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Agent GitHub Rapports — comparaison et envoi de rapports Word",
    )
    parser.add_argument(
        "mode",
        choices=[
            "collect",
            "rapport-jour", "rapport-hebdo", "rapport-mensuel", "rapport-annuel",
            "note-fiches-hebdo", "note-fiches-mensuel", "note-fiches-trim", "note-fiches-annuel",
        ],
        help=(
            "collect : télécharge, compare et analyse ; "
            "rapport-* : génère et envoie le rapport comparatif ; "
            "note-fiches-* : génère et envoie note stratégique + fiches sociétés"
        ),
    )
    args = parser.parse_args()

    if args.mode == "collect":
        cmd_collect()
    elif args.mode == "rapport-jour":
        cmd_rapport("quotidien")
    elif args.mode == "rapport-hebdo":
        cmd_rapport("hebdo")
    elif args.mode == "rapport-mensuel":
        cmd_rapport("mensuel")
    elif args.mode == "rapport-annuel":
        cmd_rapport_annuel()
    elif args.mode == "note-fiches-hebdo":
        cmd_note_fiches("HEBDO")
    elif args.mode == "note-fiches-mensuel":
        cmd_note_fiches("MENSUEL")
    elif args.mode == "note-fiches-trim":
        cmd_note_fiches("TRIM")
    elif args.mode == "note-fiches-annuel":
        cmd_note_fiches("ANNUEL")


if __name__ == "__main__":
    main()
