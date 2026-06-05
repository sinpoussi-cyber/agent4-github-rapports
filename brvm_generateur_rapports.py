#!/usr/bin/env python3
"""
brvm_generateur_rapports.py
===========================
Générateur de rapports BRVM dérivés — 7 destinataires + 2 e-mails suggérés.

Usage :
    python brvm_generateur_rapports.py \
        --source "Rapport_Ultimate_BRVM_20260528_1711.docx" \
        --out    "./rapports_derives"

Produit (nommage BRVM_<Destinataire>_AAAAMMJJ.docx) :
    1. BRVM_InvestisseursInstitutionnels_AAAAMMJJ.docx  — rapport complet
    2. BRVM_BRVM_Interne_AAAAMMJJ.docx                 — sans reco
    3. BRVM_Regulateur_AMF_UMOA_AAAAMMJJ.docx          — sans reco, score, pred
    4. BRVM_SGI_AAAAMMJJ.docx
    5. BRVM_SGO_AAAAMMJJ.docx
    6. BRVM_CIB_Conseillers_AAAAMMJJ.docx              — digest
    7. BRVM_SGP_Patrimoine_AAAAMMJJ.docx
    Email_BRVM_informatif_detaille_AAAAMMJJ.txt
    Email_BRVM_informatif_court_AAAAMMJJ.txt
"""

import argparse
import os
import re
import copy
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:
    print("[ERROR] python-docx manquant — pip install python-docx")
    sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════════════
# MATRICE DE ROUTAGE
# Clés = identifiants de sections (sous-chaînes du titre Heading1 en majuscules)
# ═══════════════════════════════════════════════════════════════════════════════

SECTIONS = {
    "SYNTHÈSE GÉNÉRALE":               "synthese",
    "ANALYSE PAR SECTEUR":             "secteurs",
    "MATRICE DE CONVERGENCE":          "matrice_signaux",
    "ANALYSE DE LIQUIDITÉ":            "liquidite",
    "TOP 10 DES DIVERGENCES":          "divergences",
    "MATRICE RISQUE":                  "matrice_risque",
    "ANALYSE MACRO":                   "macro",
    "1. ACTUALITÉS MACRO":             "macro_actu",
    "2. ACTUALITÉS POLITIQUES":        "macro_pol",
    "3. ACTUALITÉS FINANCIÈRES":       "macro_fin",
    "SYNTHÈSE & RECOMMANDATION FINALE":"macro_synth",
    "ACTUALITÉS DU MARCHÉ BRVM":       "actualites",
    "CLASSEMENT DES SOCIÉTÉS":         "classement",
    "PORTEFEUILLES MODÈLES":           "portefeuilles",
    "ALERTES DU JOUR":                 "alertes",
    "RÉCAPITULATIF DES RISQUES":       "risques",
    "SYNTHÈSE RÉCAPITULATIVE DES PRÉDICTIONS": "predictions",
    "ANALYSE FINANCIÈRE COMPARATIVE":  "analyse_fin",
    "TABLE DES MATIÈRES":              "toc",
    "ANALYSES DÉTAILLÉES PAR SOCIÉTÉ": "fiches",
    "NOTES IMPORTANTES":               "notes",
}

# Groupes macro (tous inclus ensemble)
MACRO_KEYS = {"macro", "macro_actu", "macro_pol", "macro_fin", "macro_synth"}

ROUTING = {
    # Destinataire          : sections incluses (None = toutes)
    "InvestisseursInstitutionnels": None,   # tout

    "BRVM_Interne": {
        "synthese", "secteurs", "liquidite", "alertes",
        "actualites", "risques",
    } | MACRO_KEYS,

    "Regulateur_AMF_UMOA": {
        "synthese", "liquidite", "alertes",
        "actualites", "risques",
    } | MACRO_KEYS,

    "SGI": {
        "synthese", "secteurs", "liquidite", "alertes",
        "actualites", "classement", "divergences",
    },

    "SGO": {
        "synthese", "secteurs", "liquidite", "predictions",
        "analyse_fin", "fiches", "portefeuilles", "risques",
    } | MACRO_KEYS,

    "CIB_Conseillers": {
        "synthese", "actualites", "alertes", "macro_synth",
    },

    "SGP_Patrimoine": {
        "synthese", "matrice_risque", "secteurs",
        "portefeuilles",
    } | MACRO_KEYS,
}

# Destinataires où toute recommandation est INTERDITE
NO_RECO = {"BRVM_Interne", "Regulateur_AMF_UMOA"}

# Destinataires où score/portefeuille/prédiction sont AUSSI interdits
NO_SCORE_PRED = {"Regulateur_AMF_UMOA"}

# Mots à neutraliser (remplacés par "—") dans les cellules de tableaux
RECO_WORDS_RX = re.compile(
    r'\b(ACHAT\s+FORT|ACHAT|VENTE|CONSERVER|NEUTRE|céder|opportunité'
    r'|ACHAT_FORT|conseil|recommandation|cible de cours)\b',
    re.IGNORECASE
)

# Colonnes à vider entièrement
RECO_COL_NAMES_RX = re.compile(
    r'(Recom\.|Recommandation|Signal|Reco|Score\s*/100|Score\s*\/\s*100)',
    re.IGNORECASE
)


# ═══════════════════════════════════════════════════════════════════════════════
# UTILITAIRES DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def _para_text(el) -> str:
    return ''.join(n.text or '' for n in el.iter() if n.text)


def _dedup(s: str) -> str:
    n = len(s)
    for d in (3, 2):
        if n % d == 0 and s == s[:n//d] * d:
            return s[:n//d]
    return s


def _para_style(el) -> str:
    pPr = el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
    return pPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '') if pPr is not None else ''


def _section_key(title: str) -> str | None:
    """Mappe un titre Heading1 vers une clé de section."""
    t_up = title.upper()
    for pattern, key in SECTIONS.items():
        if pattern in t_up:
            return key
    return None


def _is_heading1(el) -> bool:
    return _para_style(el) in ('Titre1', 'Heading1')


def _cell_text(cell_el) -> str:
    return _dedup(''.join(n.text or '' for n in cell_el.iter() if n.text).strip())


def _set_run_text(run_el, new_text: str):
    """Remplace le texte d'un run XML."""
    t_el = run_el.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    if t_el is not None:
        t_el.text = new_text
    else:
        t_el = run_el.makeelement(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', {}
        )
        t_el.text = new_text
        run_el.append(t_el)


def _neutralize_para(para_el):
    """Remplace les mots de recommandation dans tous les runs d'un paragraphe."""
    for run in para_el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        t = run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        if t is not None and t.text:
            new_text = RECO_WORDS_RX.sub('—', t.text)
            if new_text != t.text:
                t.text = new_text


def _neutralize_table(tbl_el, no_score: bool = False):
    """
    Dans un tableau :
    - Repère les colonnes de recommandation (header) → vide toutes les cellules
    - Dans le reste, remplace les mots de recommandation par '—'
    - Si no_score : vide aussi les colonnes Score et préfixes de Score
    """
    rows = tbl_el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
    if not rows:
        return

    # Identifier les colonnes à vider depuis le header (1ère ligne)
    header_cells = rows[0].findall(
        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'
    )
    reco_cols: set[int] = set()
    for ci, hc in enumerate(header_cells):
        h_txt = _cell_text(hc)
        if RECO_COL_NAMES_RX.search(h_txt):
            reco_cols.add(ci)
        if no_score and re.search(r'score|signal', h_txt, re.IGNORECASE):
            reco_cols.add(ci)

    # Traiter toutes les lignes
    for ri, row in enumerate(rows):
        cells = row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
        for ci, cell in enumerate(cells):
            if ci in reco_cols and ri > 0:
                # Vider la cellule
                for run in cell.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                    t = run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    if t is not None:
                        t.text = '—'
            else:
                # Neutraliser les mots de reco dans le texte libre
                for run in cell.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                    t = run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    if t is not None and t.text:
                        t.text = RECO_WORDS_RX.sub('—', t.text)


# ═══════════════════════════════════════════════════════════════════════════════
# DÉCOUPAGE DU RAPPORT SOURCE EN SECTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def split_source_by_sections(doc: Document) -> dict:
    """
    Découpe le document source en blocs par section Heading1.
    Retourne dict {section_key: [elements]}.
    """
    body_elems = list(doc.element.body.iterchildren())
    sections: dict[str, list] = {"_header": []}  # éléments avant la 1ère section
    current_key = "_header"

    for el in body_elems:
        tag = el.tag.split('}')[-1]
        if tag == 'p' and _is_heading1(el):
            title = _dedup(_para_text(el).strip())
            key = _section_key(title)
            if key:
                current_key = key
                sections.setdefault(key, [])
        sections.setdefault(current_key, []).append(el)

    return sections


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTRUCTION D'UN RAPPORT DÉRIVÉ
# ═══════════════════════════════════════════════════════════════════════════════

def _clone_element(el):
    """Clone profond d'un élément XML."""
    return copy.deepcopy(el)


def build_derived_doc(
    source_doc: Document,
    sections_data: dict,
    allowed_keys: set | None,
    no_reco: bool = False,
    no_score_pred: bool = False,
    date_str: str = "",
    dest_label: str = "",
) -> Document:
    """
    Construit un Document Word dérivé selon les sections autorisées.

    - allowed_keys : None = tout inclure ; sinon = ensemble des clés autorisées
    - no_reco      : supprimer toute recommandation du contenu
    - no_score_pred: supprimer aussi scores et prédictions de cours
    """
    new_doc = Document()
    # Supprimer le paragraphe vide par défaut
    for p in new_doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Ajouter un bandeau d'en-tête destinataire
    def _add_banner():
        p = new_doc.add_paragraph()
        run = p.add_run(
            f"RAPPORT BRVM — {dest_label.upper().replace('_', ' ')} — {date_str}"
        )
        run.bold = True
        run.font.size = Pt(11)
        if no_reco:
            warn = new_doc.add_paragraph()
            warn_r = warn.add_run(
                "⚠  VERSION SANS CONSEIL NI SIGNAL D'INVESTISSEMENT"
                " — diffusion strictement réglementée."
            )
            warn_r.bold = True
            warn_r.font.color.rgb = RGBColor(0xD9, 0x30, 0x25)
            warn_r.font.size = Pt(9)
        new_doc.add_paragraph()

    _add_banner()

    # Ordre d'insertion : respecter l'ordre naturel du rapport source
    all_keys_ordered = list(SECTIONS.values())

    keys_to_include = (
        [k for k in all_keys_ordered if k in allowed_keys] + ["_header"]
        if allowed_keys is not None
        else list(sections_data.keys())
    )

    # Toujours inclure l'en-tête source (couverture / titre)
    if "_header" in sections_data:
        for el in sections_data["_header"]:
            cloned_h = _clone_element(el)
            if no_reco:
                tag_h = cloned_h.tag.split('}')[-1]
                if tag_h == 'p': _neutralize_para(cloned_h)
                elif tag_h == 'tbl': _neutralize_table(cloned_h, no_score=no_score_pred)
            new_doc.element.body.append(cloned_h)

    for key in all_keys_ordered:
        if allowed_keys is not None and key not in allowed_keys:
            continue
        if key not in sections_data:
            continue
        elems = sections_data[key]
        for el in elems:
            cloned = _clone_element(el)
            tag = cloned.tag.split('}')[-1]

            if no_reco:
                if tag == 'p':
                    _neutralize_para(cloned)
                elif tag == 'tbl':
                    _neutralize_table(cloned, no_score=no_score_pred)

            if no_score_pred and tag == 'p':
                # Masquer les lignes qui contiennent score /100 ou prédiction numérique
                text = _dedup(_para_text(cloned).strip())
                if re.search(r'\b\d+\s*/\s*100\b|\bprédi[ct]|\bJ\+\d+\b', text, re.IGNORECASE):
                    continue

            new_doc.element.body.append(cloned)

    # Ajouter clause de non-responsabilité en pied de document
    new_doc.add_paragraph()
    disclaimer = new_doc.add_paragraph()
    dr = disclaimer.add_run(
        "Cette analyse est diffusée à titre strictement informatif. "
        "Elle ne constitue ni un conseil en investissement, "
        "ni une recommandation d'achat ou de vente de valeurs mobilières. "
        f"Rapport généré le {date_str}."
    )
    dr.italic = True
    dr.font.size = Pt(8)

    return new_doc


# ═══════════════════════════════════════════════════════════════════════════════
# EXTRACTION DES DONNÉES MARCHÉ POUR LES E-MAILS
# ═══════════════════════════════════════════════════════════════════════════════

def extract_market_data(sections_data: dict) -> dict:
    """
    Extrait les indicateurs agrégés neutres depuis la section 'synthese' :
    BRVM Composite, capitalisation, nombre de sociétés.
    """
    data = {
        "composite": "—",
        "var_composite": "—",
        "capitalisation": "—",
        "nb_societes": "47",
        "date": "—",
    }

    synth_elems = sections_data.get("synthese", [])
    rx_composite = re.compile(r'BRVM\s*Composite[^0-9]*([0-9\s,\.]+)\s*points?', re.IGNORECASE)
    rx_var = re.compile(r'\(([+-]?\s*[0-9,\.]+\s*%)\)', re.IGNORECASE)
    rx_capi = re.compile(r'capitalisation[^0-9]*([0-9\s,\.]+)\s*(Mds|milliards|Md)', re.IGNORECASE)
    rx_nb = re.compile(r'([0-9]+)\s*soci[eé]t[eé]', re.IGNORECASE)

    for el in synth_elems:
        t = _dedup(_para_text(el).strip())
        m = rx_composite.search(t)
        if m and data["composite"] == "—":
            data["composite"] = m.group(1).strip().replace('\u202f', ' ')
        mv = rx_var.search(t)
        if mv and data["var_composite"] == "—":
            data["var_composite"] = mv.group(1).strip()
        mc = rx_capi.search(t)
        if mc and data["capitalisation"] == "—":
            data["capitalisation"] = f"{mc.group(1).strip()} {mc.group(2)}"
        mn = rx_nb.search(t)
        if mn and data["nb_societes"] == "47":
            data["nb_societes"] = mn.group(1)

    return data


# ═══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION DES E-MAILS
# ═══════════════════════════════════════════════════════════════════════════════

GARDE_FOU_RX = re.compile(
    r'\b(achat|vente|recommandation|opportunité|conseil|cible|investissez|achetez|vendez)\b',
    re.IGNORECASE
)


def _garde_fou_check(text: str) -> list[str]:
    """Retourne les mots interdits trouvés dans le texte variable."""
    return GARDE_FOU_RX.findall(text)


DISCLAIMER = (
    "\n\n---\n"
    "Cette communication est diffusée à titre strictement informatif. "
    "Elle ne constitue ni un conseil en investissement, "
    "ni une recommandation d'achat ou de vente.\n"
)


def build_email_detailed(market: dict, date_str: str) -> str:
    """E-mail informatif détaillé — agrégats marché uniquement."""
    body_variable = (
        f"Bourse Régionale des Valeurs Mobilières (BRVM) — {date_str}\n\n"
        f"Indice BRVM Composite : {market['composite']} points  "
        f"({market['var_composite']} sur la séance)\n"
        f"Capitalisation globale du marché : {market['capitalisation']} FCFA\n"
        f"Nombre de sociétés suivies : {market['nb_societes']}\n\n"
        "Pour toute information complémentaire sur l'évolution des marchés, "
        "nous vous invitons à consulter le rapport de marché complet disponible "
        "auprès de votre intermédiaire agréé."
    )

    # Garde-fou
    hits = _garde_fou_check(body_variable)
    if hits:
        return (
            f"[BLOQUÉ PAR GARDE-FOU] Mots interdits détectés : {hits}\n"
            "L'e-mail n'a PAS été généré. Corrigez la source et relancez."
        )

    subject = (
        f"BRVM — Indicateurs de marché du {date_str} "
        f"(BRVM Composite : {market['composite']} pts)"
    )
    return (
        f"OBJET : {subject}\n\n"
        f"{body_variable}"
        f"{DISCLAIMER}"
        "⚠  SUGGESTION — à valider avant tout envoi."
    )


def build_email_short(market: dict, date_str: str) -> str:
    """E-mail informatif court."""
    body_variable = (
        f"BRVM — {date_str}\n"
        f"Composite : {market['composite']} pts ({market['var_composite']})  |  "
        f"Capitalisation : {market['capitalisation']} FCFA  |  "
        f"{market['nb_societes']} sociétés suivies."
    )

    hits = _garde_fou_check(body_variable)
    if hits:
        return (
            f"[BLOQUÉ PAR GARDE-FOU] Mots interdits : {hits}\n"
            "L'e-mail n'a PAS été généré."
        )

    subject = f"BRVM Marché {date_str}"
    return (
        f"OBJET : {subject}\n\n"
        f"{body_variable}"
        f"{DISCLAIMER}"
        "⚠  SUGGESTION — à valider avant tout envoi."
    )


# ═══════════════════════════════════════════════════════════════════════════════
# CONTRÔLE QUALITÉ
# ═══════════════════════════════════════════════════════════════════════════════

def quality_check(doc: Document, dest: str) -> dict:
    """
    Vérifie l'absence de recommandations pour BRVM_Interne et Regulateur.
    Retourne un rapport {ok, occurrences, anomalies}.
    """
    result = {"dest": dest, "ok": True, "occurrences": [], "anomalies": []}
    check_reco = dest in NO_RECO

    # Textes à ignorer dans le QC (clause de non-responsabilité interne)
    QC_IGNORE = re.compile(
        r'(ni un conseil en investissement|ni une recommandation d\'achat|'
        r'VERSION SANS CONSEIL|strictement informatif|strictement réglementé)',
        re.IGNORECASE
    )
    for el in doc.element.body.iter():
        tag = el.tag.split('}')[-1]
        if tag in ('t', 'p'):
            text = el.text or ''
            if check_reco and not QC_IGNORE.search(text):
                hits = RECO_WORDS_RX.findall(text)
                if hits:
                    result["ok"] = False
                    result["occurrences"].extend(hits)

    # Anomalie calibrage : tous les signaux en ACHAT (info uniquement)
    if dest == "InvestisseursInstitutionnels":
        all_text = ' '.join(
            el.text or '' for el in doc.element.body.iter()
            if el.tag.split('}')[-1] == 't'
        )
        n_achat = len(re.findall(r'\bACHAT\b', all_text, re.IGNORECASE))
        n_vente = len(re.findall(r'\bVENTE\b', all_text, re.IGNORECASE))
        if n_vente == 0 and n_achat > 20:
            result["anomalies"].append(
                f"⚠  Calibrage : {n_achat} signaux ACHAT, 0 VENTE "
                "— signal marché potentiellement anormal (à signaler)."
            )

    return result


# ═══════════════════════════════════════════════════════════════════════════════
# POINT D'ENTRÉE PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════

def generate_all(source_path: str, out_dir: str) -> dict:
    """
    Génère les 7 rapports dérivés + 2 e-mails.
    Retourne un résumé {files, qc_reports, emails}.
    """
    source_path = Path(source_path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Date du rapport
    date_match = re.search(r'(\d{8})', source_path.name)
    if date_match:
        raw_date = date_match.group(1)
        date_str = f"{raw_date[:4]}-{raw_date[4:6]}-{raw_date[6:]}"
    else:
        date_str = datetime.today().strftime('%Y-%m-%d')

    date_file = date_str.replace('-', '')

    print(f"[GenRapports] Source : {source_path.name}")
    print(f"[GenRapports] Date   : {date_str}")
    print(f"[GenRapports] Sortie : {out_dir}")

    # Charger le document source
    source_doc = Document(str(source_path))
    sections_data = split_source_by_sections(source_doc)
    print(f"[GenRapports] Sections détectées : {[k for k in sections_data if k != '_header']}")

    market_data = extract_market_data(sections_data)
    print(f"[GenRapports] Composite : {market_data['composite']} pts "
          f"({market_data['var_composite']})")

    created_files = []
    qc_reports = []

    # ── Générer les 7 rapports ────────────────────────────────────────────────
    for dest, allowed in ROUTING.items():
        no_reco = dest in NO_RECO
        no_score = dest in NO_SCORE_PRED

        print(f"[GenRapports] Génération {dest}  "
              f"[sections={len(allowed) if allowed else 'TOUTES'}  "
              f"no_reco={no_reco}  no_score={no_score}]")

        derived = build_derived_doc(
            source_doc=source_doc,
            sections_data=sections_data,
            allowed_keys=allowed,
            no_reco=no_reco,
            no_score_pred=no_score,
            date_str=date_str,
            dest_label=dest,
        )

        fname = f"BRVM_{dest}_{date_file}.docx"
        fpath = out_dir / fname
        derived.save(str(fpath))
        size_kb = fpath.stat().st_size // 1024
        created_files.append((fname, size_kb))
        print(f"  ✓  {fname}  ({size_kb} KB)")

        # Contrôle qualité
        qc = quality_check(derived, dest)
        qc_reports.append(qc)
        if not qc["ok"]:
            print(f"  ✗  QC ÉCHEC : {len(qc['occurrences'])} occurrence(s) de reco "
                  f"pour {dest}")
        for anomaly in qc.get("anomalies", []):
            print(f"  {anomaly}")

    # ── Générer les 2 e-mails ─────────────────────────────────────────────────
    emails = {}

    email_long = build_email_detailed(market_data, date_str)
    email_short = build_email_short(market_data, date_str)

    fname_long  = f"Email_BRVM_informatif_detaille_{date_file}.txt"
    fname_short = f"Email_BRVM_informatif_court_{date_file}.txt"

    for fname, content in [(fname_long, email_long), (fname_short, email_short)]:
        fpath = out_dir / fname
        fpath.write_text(content, encoding='utf-8')
        size_kb = fpath.stat().st_size // 1024
        created_files.append((fname, max(size_kb, 1)))
        print(f"  ✓  {fname}")
        emails[fname] = content

    # ── Résumé final ──────────────────────────────────────────────────────────
    print("\n[GenRapports] === RÉSUMÉ ===")
    print(f"  {len([f for f in created_files if f[0].endswith('.docx')])} fichiers .docx")
    print(f"  {len([f for f in created_files if f[0].endswith('.txt')])}  fichiers .txt")

    qc_failures = [q for q in qc_reports if not q["ok"]]
    if qc_failures:
        print(f"\n  ⚠  QC : {len(qc_failures)} rapport(s) avec occurrences de reco :")
        for q in qc_failures:
            print(f"     {q['dest']} — {len(q['occurrences'])} occurrence(s)")
    else:
        print("\n  ✓  QC : 0 occurrence de recommandation dans BRVM_Interne et Régulateur.")

    all_anomalies = [a for q in qc_reports for a in q.get("anomalies", [])]
    if all_anomalies:
        for a in all_anomalies:
            print(f"\n  {a}")

    print("\n[GenRapports] Fichiers générés :")
    for fname, size_kb in created_files:
        print(f"  {fname:<60s}  {size_kb:>5d} KB")

    return {
        "files": created_files,
        "qc_reports": qc_reports,
        "emails": emails,
        "market_data": market_data,
    }


def main():
    parser = argparse.ArgumentParser(
        description="Générateur de rapports BRVM dérivés (7 destinataires + 2 e-mails)"
    )
    parser.add_argument('--source', required=True,
                        help='Chemin du rapport source .docx')
    parser.add_argument('--out', required=True,
                        help='Dossier de sortie pour les rapports dérivés')
    args = parser.parse_args()

    if not Path(args.source).exists():
        print(f"[ERROR] Source introuvable : {args.source}")
        sys.exit(1)

    result = generate_all(args.source, args.out)

    # Afficher les e-mails en suggestion
    print("\n" + "="*70)
    print("SUGGESTIONS D'E-MAILS (à valider avant tout envoi)")
    print("="*70)
    for fname, content in result["emails"].items():
        print(f"\n--- {fname} ---")
        print(content)
        print()


if __name__ == '__main__':
    main()
