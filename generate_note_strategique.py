import io
import json
import logging
import os
import re
from datetime import date

logger = logging.getLogger(__name__)

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from dotenv import load_dotenv

load_dotenv()

_DESTINATAIRE = (
    "À l'attention de Madame Corinne Houmou ORMON\n"
    "Directrice de l'Antenne Nationale de Bourse de Côte d'Ivoire"
)


# ── Helpers docx ──────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _rgb(hex6: str) -> RGBColor:
    return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))


def _bold(paragraph, text: str, size: int = 11, color: str = None):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = _rgb(color)
    return run


def _normal(paragraph, text: str, size: int = 10):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    return run


def _heading(doc, text: str, level: int = 1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(style=style_map.get(level, "Heading 1"))
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = _rgb("1A73E8")
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = _rgb("333333")
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = _rgb("555555")
    return p


def _para(doc, text: str, size: int = 10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _add_page_number(run):
    for tag, text in [("begin", None), (None, " PAGE "), ("end", None)]:
        if tag:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
            run._r.append(el)
        else:
            el = OxmlElement("w:instrText")
            el.text = text
            el.set(qn("xml:space"), "preserve")
            run._r.append(el)


def _setup_header_footer(doc, date_str: str):
    section = doc.sections[0]

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(2)
    r1 = hp.add_run(f"NOTE STRATÉGIQUE BRVM — {date_str}    ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = _rgb("1A73E8")
    r2 = hp.add_run("CONFIDENTIEL")
    r2.bold = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = _rgb("D93025")
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A73E8")
    pBdr.append(bot)
    pPr.append(pBdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.space_before = Pt(0)
    r3 = fp.add_run("Page ")
    r3.font.size = Pt(8)
    r3.font.color.rgb = _rgb("999999")
    pn = fp.add_run()
    _add_page_number(pn)
    pn.font.size = Pt(8)
    pn.font.color.rgb = _rgb("999999")


def _reco_color(reco: str) -> str:
    r = str(reco).upper()
    if "ACHAT" in r:
        return "C6EFCE"
    if "VENTE" in r:
        return "FFC7CE"
    return "FFEB9C"


def _tbl_header(tbl, headers, bg="E8F0FE", fg="333333"):
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        _cell_bg(c, bg)
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb(fg)


def _bloc(doc, label: str, texte: str, bg: str, fg: str = "1A1A1A"):
    """Bloc coloré POINT CLÉ / RISQUE / OPPORTUNITÉ inséré dans le document."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "100")
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = _rgb(fg)
    r2 = p.add_run(texte)
    r2.font.size = Pt(9)
    r2.font.color.rgb = _rgb("333333")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(3)


def _sep_synthese(doc):
    """Séparateur visuel '▌ SYNTHÈSE EXÉCUTIVE'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("▌  SYNTHÈSE EXÉCUTIVE")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb("1558A7")


def _sep_detail(doc):
    """Séparateur visuel '▌ ANALYSE DÉTAILLÉE'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("▌  ANALYSE DÉTAILLÉE")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb("666666")


def _graph_ph(doc, texte: str):
    """Placeholder de graphique centré en italique gris."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texte)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb("888888")


def _graph_lecture(doc, texte: str):
    """Lecture du graphique en italique discret sous le placeholder."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Lecture du graphique — " + texte)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb("444444")


def _interp_tendance(tendance: str, perf: str) -> str:
    """Interprétation analytique approfondie d'une tendance de marché."""
    t = str(tendance or "").lower()
    positif = str(perf or "").strip().startswith("+")
    if "haussier" in t and positif:
        return (
            "Cette dynamique haussière traduit un regain d'appétit pour le risque, "
            "porté par une demande institutionnelle soutenue sur les valeurs de référence. "
            "La conjonction d'une tendance de fond positive et d'une performance cumulée favorable "
            "constitue un signal de continuité : les investisseurs peuvent envisager de maintenir, "
            "voire de renforcer leurs expositions sur les titres affichant les meilleurs scores composites. "
            "La sélectivité reste néanmoins de mise — les valeurs à faible liquidité doivent être traitées "
            "avec prudence même dans un contexte porteur."
        )
    if "haussier" in t:
        return (
            "La tendance de fond demeure haussière, même si la performance récente indique "
            "une phase de consolidation pouvant précéder une reprise du momentum. "
            "Ce type de configuration — tendance structurelle intacte mais progression momentanément freinée — "
            "représente souvent une opportunité d'entrée tactique pour les investisseurs n'ayant pas encore "
            "initié de position, sous réserve d'une confirmation technique à court terme. "
            "Il convient de surveiller les niveaux de support clés : un franchissement à la baisse "
            "remettrait en cause le scénario haussier de base."
        )
    if "baissier" in t:
        return (
            "Cette configuration baissière signale une prudence accrue des opérateurs, "
            "susceptible de refléter des prises de bénéfices massives ou des incertitudes macro-économiques "
            "pesant sur la valorisation des actifs cotés. "
            "Dans ce contexte, la priorité est à la préservation du capital : réduire les expositions "
            "sur les titres à faible score fondamental, privilégier le cash et les valeurs défensives "
            "à dividendes stables, et éviter tout renforcement sur momentum baissier. "
            "Un retournement ne pourra être confirmé qu'après une clôture franche au-dessus "
            "des résistances techniques identifiées."
        )
    return (
        "Le marché traverse une phase de neutralité et de consolidation, caractérisée par "
        "une absence de conviction directionnelle des opérateurs et des volumes en repli. "
        "Cette stabilité apparente peut masquer des tensions latentes — elle est souvent "
        "précurseur d'une rupture de tendance dont la direction doit être anticipée. "
        "La stratégie adaptée consiste à maintenir une allocation équilibrée, "
        "en réduisant l'exposition aux valeurs les plus volatiles et en conservant "
        "une réserve de liquidités permettant d'agir rapidement à la première confirmation directionnelle."
    )


def _lecture_graph_indice(idx: dict) -> str:
    """Analyse graphique approfondie (4 phrases) pour l'indice BRVM Composite."""
    niveau = idx["niveau"]
    haut = idx["plus_haut"]
    bas = idx["plus_bas"]
    perf = idx["perf_100j"]
    t = str(idx["tendance"]).lower()
    if "haussier" in t:
        return (
            f"La courbe décrit une progression structurée entre {bas} et {haut} points sur la période, "
            f"avec des phases de correction limitées qui ont systématiquement débouché sur de nouveaux plus hauts — "
            f"témoignant d'une structure technique fondamentalement saine. "
            f"Le gain cumulé de {perf} reflète un réel intérêt acheteur, non une simple effervescence spéculative. "
            f"À {niveau} points, l'indice se positionne en zone haute de son couloir de valorisation : "
            f"la résistance à {haut} constitue le prochain verrou à franchir pour ouvrir un nouveau palier haussier. "
            f"Un repli sous {bas} invaliderait le scénario haussier et inviterait à réviser l'allocation."
        )
    if "baissier" in t:
        return (
            f"L'évolution graphique révèle une pression vendeuse structurelle depuis {haut} vers {bas} points, "
            f"sans rebond technique significatif capable d'interrompre la séquence de plus bas successifs. "
            f"La performance de {perf} sur la période matérialise l'ampleur de la correction et "
            f"souligne l'absence d'acheteurs suffisamment convaincus pour absorber les flux vendeurs. "
            f"Le niveau actuel de {niveau} points représente un support critique : "
            f"son franchissement à la baisse ouvrirait la voie vers {bas} en première instance. "
            f"Toute tentative de rebond doit être confirmée en clôture avant d'envisager un repositionnement."
        )
    return (
        f"L'indice évolue dans un couloir de consolidation compris entre {bas} et {haut} points, "
        f"sans orientation directionnelle franche sur la période analysée. "
        f"La performance de {perf} illustre l'équilibre des forces en présence — "
        f"les acheteurs et les vendeurs s'annulant mutuellement sans qu'un camp ne prenne durablement l'avantage. "
        f"À {niveau} points, l'indice navigue en zone médiane de son range : "
        f"un bris au-dessus de {haut} serait le signal acheteur attendu, "
        f"un bris sous {bas} ouvrirait à la baisse. "
        f"La stratégie recommandée dans ce contexte : attente et préservation du capital."
    )


def _lecture_graph_capi(cap: dict) -> str:
    """Analyse graphique approfondie (4 phrases) pour la capitalisation boursière."""
    evolution = str(cap.get("evolution") or "").strip()
    valeur = cap["valeur"]
    pic = cap.get("pic", "—")
    plancher = cap.get("plancher", "—")
    if evolution.startswith("+"):
        pic_str = f", après avoir atteint un pic à {pic}" if pic != "—" else ""
        return (
            f"La progression de la capitalisation à {valeur}{pic_str} confirme une création de "
            f"valeur nette sur le marché — signe que les flux entrants surpassent les sorties de capitaux. "
            f"Cette dynamique positive reflète un intérêt institutionnel soutenu et une confiance "
            f"des opérateurs dans la solidité des fondamentaux des sociétés cotées. "
            f"Une capitalisation en hausse améliore mécaniquement la profondeur du marché et "
            f"réduit le risque de distorsion de prix sur les titres de référence. "
            f"Pour les investisseurs, ce contexte favorise les stratégies d'accumulation progressive "
            f"sur les valeurs à score composite élevé."
        )
    if evolution.startswith("-"):
        plancher_str = f", avec un plancher à {plancher} atteint sur la période" if plancher != "—" else ""
        return (
            f"Le recul de la capitalisation à {valeur}{plancher_str} signale une sortie nette de capitaux, "
            f"traduisant soit des prises de bénéfices coordonnées, soit une défiance des investisseurs "
            f"face à un contexte macro-économique ou sectoriel dégradé. "
            f"Une capitalisation en repli réduit la profondeur du marché et amplifie la sensibilité "
            f"du prix des titres aux ordres de vente — le risque de slippage est accru. "
            f"Cette configuration invite à une révision des pondérations : alléger les positions "
            f"sur les valeurs à faible liquidité en priorité, et renforcer la part de cash du portefeuille."
        )
    return (
        f"La capitalisation stabilisée à {valeur} traduit un équilibre entre flux acheteurs et vendeurs, "
        f"caractéristique d'une phase d'attente sur les marchés. "
        f"Cette stabilité peut être interprétée positivement — absence de panique vendeuse — "
        f"ou négativement — faiblesse des convictions acheteuses. "
        f"La profondeur de marché reste préservée, ce qui facilite les ajustements tactiques de portefeuille "
        f"sans impact majeur sur les prix. "
        f"La direction du prochain mouvement de capitalisation sera un indicateur avancé "
        f"de la tendance de l'indice dans les semaines à venir."
    )


# ── Refonte par recopie du rapport source (par titres Heading 1) ─────────────
#
# Le rapport source (Rapport_Ultimate_BRVM_*.docx) contient toutes les sections
# requises (Synthèse, Sectoriel, Liquidité, Macro, Actualités, Classement,
# Portefeuilles, Alertes) sous forme de paragraphes Heading 1. La note
# stratégique est désormais construite en recopiant ces sections, plutôt qu'en
# ré-extrayant les données via Claude. Les ancres ci-dessous sont des
# sous-chaînes de chaque titre Heading 1.

_SOURCE_ANCHORS = {
    "synthese":         "SYNTHÈSE GÉNÉRALE",
    "secteurs":         "ANALYSE PAR SECTEUR",
    "matrice_signaux":  "MATRICE DE CONVERGENCE DES SIGNAUX",
    "liquidite":        "ANALYSE DE LIQUIDITÉ",
    "top_divergences":  "TOP 10 DES DIVERGENCES MAJEURES",
    "matrice_risque":   "MATRICE RISQUE vs HORIZON",
    "macro":            "ANALYSE MACRO",
    "macro_actu":       "1. ACTUALITÉS MACRO",
    "macro_pol":        "2. ACTUALITÉS POLITIQUES",
    "macro_fin":        "3. ACTUALITÉS FINANCIÈRES",
    "macro_synth":      "SYNTHÈSE & RECOMMANDATION FINALE",
    "actualites":       "ACTUALITÉS DU MARCHÉ BRVM",
    "classement":       "CLASSEMENT DES SOCIÉTÉS",
    "portefeuilles":    "PORTEFEUILLES MODÈLES",
    "alertes":          "ALERTES DU JOUR",
    "recap_risques":    "RÉCAPITULATIF DES RISQUES",
    "toc_detail":       "TABLE DES MATIÈRES - ANALYSES DÉTAILLÉES",
    "predictions":      "PRÉDICTIONS",
    "analyse_financiere": "ANALYSE FINANCIÈRE",
}


def _para_style(p_elem) -> str:
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return ""
    pStyle = pPr.find(qn("w:pStyle"))
    return pStyle.get(qn("w:val")) if pStyle is not None else ""


def _para_text(p_elem) -> str:
    return "".join(t.text for t in p_elem.iter(qn("w:t")) if t.text)


def _open_source_doc(docs_bytes):
    if isinstance(docs_bytes, bytes):
        return Document(io.BytesIO(docs_bytes))
    return Document(io.BytesIO(docs_bytes[0]))


def _split_source_by_h1(source_doc) -> dict:
    """Découpe le corps du document source en buckets par Heading 1.

    Retourne {anchor_key: [(kind, xml_element), ...]} où kind ∈ {"p","tbl"}.
    Les blocs avant le premier Heading 1 connu sont ignorés.
    """
    body = source_doc.element.body
    bucket = {k: [] for k in _SOURCE_ANCHORS}
    current_key = None
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag not in ("p", "tbl"):
            continue
        if tag == "p" and _para_style(child) == "Heading1":
            txt = _para_text(child)
            current_key = None
            for k, anchor in _SOURCE_ANCHORS.items():
                if anchor in txt:
                    current_key = k
                    break
        if current_key is not None:
            bucket[current_key].append((tag, child))
    return bucket


def _image_blobs_in_blocks(source_doc, blocks):
    """Renvoie [(blob_bytes, ext), ...] des images des blocs (ordre du flux)."""
    out = []
    for kind, el in blocks:
        if kind != "p":
            continue
        for drawing in el.iter(qn("w:drawing")):
            blip = drawing.find(".//" + qn("a:blip"))
            if blip is None:
                continue
            rId = blip.get(qn("r:embed"))
            if not rId:
                continue
            try:
                part = source_doc.part.related_parts[rId]
                blob = part.blob
                ct = getattr(part, "content_type", "image/png")
                ext = "jpg" if ("jpeg" in ct or "jpg" in ct) else "png"
                out.append((blob, ext))
            except (KeyError, AttributeError):
                continue
    return out


def _set_table_borders(tbl, val: str = "nil"):
    """Force toutes les bordures du tableau (val='nil' pour invisible)."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tblPr.append(borders)


def _insert_images_side_by_side(doc, blob_left, blob_right, caption_left, caption_right,
                                 width_cm: float = 8.0):
    """Insère 2 images dans une table 2×2 sans bordure ; légendes en ligne 2."""
    if not blob_left and not blob_right:
        return
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_table_borders(tbl, "nil")
    for i, blob in enumerate((blob_left, blob_right)):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if blob:
            try:
                p.add_run().add_picture(io.BytesIO(blob), width=Cm(width_cm))
            except Exception as e:
                logger.warning(f"Échec insertion image: {e}")
    for i, cap in enumerate((caption_left, caption_right)):
        cell = tbl.rows[1].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(cap or "")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb("555555")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def _copy_source_table(doc, source_tbl_xml, header_bg: str = "1A73E8", header_fg: str = "FFFFFF",
                        max_rows: int = None):
    """Recopie le contenu d'un <w:tbl> source dans 'doc' avec restyle.

    Détecte et déduplique un header répété en ligne 2 (pattern fréquent dans le
    rapport source). max_rows limite le nombre de lignes de données copiées.
    """
    rows = source_tbl_xml.findall(qn("w:tr"))
    if not rows:
        return None
    data = []
    for tr in rows:
        cells = []
        for tc in tr.findall(qn("w:tc")):
            cells.append("".join(t.text or "" for t in tc.iter(qn("w:t"))).strip())
        data.append(cells)
    if not data:
        return None
    if len(data) > 1 and data[0] == data[1]:
        data = [data[0]] + data[2:]
    if max_rows is not None and len(data) > max_rows + 1:
        data = data[: max_rows + 1]
    n_cols = max(len(r) for r in data)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Table Grid"
    headers = (data[0] + [""] * n_cols)[:n_cols]
    _tbl_header(tbl, headers, header_bg, header_fg)
    for row_data in data[1:]:
        row = tbl.add_row()
        padded = (row_data + [""] * n_cols)[:n_cols]
        for i, val in enumerate(padded):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl


def _make_histogram_png(labels, values, title: str, color: str = "#1A73E8") -> bytes:
    """Génère un histogramme PNG via matplotlib."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(6.0, 3.6), dpi=140)
    bars = ax.bar(range(len(labels)), values, color=color,
                  edgecolor="white", linewidth=0.8)
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=30, ha="right", fontsize=8)
    ax.set_title(title, fontsize=10.5, fontweight="bold", color="#222222", pad=10)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    ax.tick_params(axis="y", labelsize=7)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    for bar, v in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(),
                f"{v:,.0f}".replace(",", " "),
                ha="center", va="bottom", fontsize=6.8, color="#444444")
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# ── Extraction texte source ───────────────────────────────────────────────────

def _extract_text(doc_bytes: bytes) -> str:
    doc = Document(io.BytesIO(doc_bytes))
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    text = "\n".join(lines)
    print(f"[DEBUG] Longueur texte extrait : {len(text)} chars")
    print(f"[DEBUG] Texte extrait (500 premiers chars) : {text[:500]}")
    return text


# ── Contexte multi-documents ─────────────────────────────────────────────────

def _build_context(docs_bytes: list, freq: str) -> str:
    if len(docs_bytes) == 1:
        return _extract_text(docs_bytes[0])

    max_older = {"HEBDO": 6, "MENSUEL": 9, "TRIM": 12, "ANNUEL": 14}.get(freq, 5)
    chars_older = {"HEBDO": 3000, "MENSUEL": 2000, "TRIM": 1500, "ANNUEL": 1000}.get(freq, 2000)

    recent = _extract_text(docs_bytes[0])[:25000]
    print(f"[DEBUG] Texte envoyé à Claude (200 premiers chars) : {recent[:200]}")
    older_parts = []
    for i, db in enumerate(docs_bytes[1:max_older + 1], 1):
        excerpt = _extract_text(db)[:chars_older]
        older_parts.append(f"--- Document J-{i} ---\n{excerpt}")

    return (
        "=== DOCUMENT LE PLUS RÉCENT (J) ===\n"
        + recent
        + ("\n\n=== HISTORIQUE (extraits) ===\n" + "\n\n".join(older_parts) if older_parts else "")
    )


# ── Extraction structurée via Claude ─────────────────────────────────────────

def _extract_data_with_claude(full_text: str, freq: str = "JOUR", period_info: dict = None) -> dict:
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    text = full_text[:40000]

    _period_descs = {
        "HEBDO": "7 derniers jours (synthèse hebdomadaire)",
        "MENSUEL": "30 derniers jours (bilan mensuel)",
        "TRIM": "dernier trimestre (bilan trimestriel)",
        "ANNUEL": "dernière année (bilan annuel complet)",
    }
    if freq != "JOUR":
        nb = (period_info or {}).get("nb_seances", "?")
        period_prefix = (
            f"Tu analyses une SYNTHÈSE sur les {_period_descs.get(freq, freq)} "
            f"couvrant {nb} séances. "
            f"Le document le plus récent est présenté en premier, suivi d'extraits historiques. "
            f"Dans tes champs textuels, décris l'ÉVOLUTION sur la période. "
            f"Pour les risques et opportunités, tiens compte des tendances observées.\n\n"
        )
    else:
        period_prefix = ""

    prompt = f"""Tu es un expert financier BRVM. {period_prefix}Extrais les données du rapport ci-dessous.
Retourne UNIQUEMENT un JSON valide, sans texte avant ou après. Utilise null si une donnée est absente.

RAPPORT SOURCE :
{text}

FORMAT JSON ATTENDU :
{{
  "sentiment": "Haussier|Baissier|Neutre",
  "brvm_composite": "245.67",
  "brvm_variation": "+0.45%",
  "brvm_composite_plus_haut_100j": "250.00",
  "brvm_composite_plus_bas_100j": "230.00",
  "brvm_composite_tendance": "haussière|baissière|latérale",
  "capitalisation": "6 750 Mds FCFA",
  "capitalisation_evolution_100j": "+3.2%",
  "capitalisation_pic_100j": "7 000 Mds FCFA",
  "capitalisation_plancher_100j": "6 400 Mds FCFA",
  "nb_achat": 12,
  "nb_neutre": 20,
  "nb_vente": 15,
  "top_opportunite": "SGBCI",
  "principale_divergence": "description",
  "perf_historique_100j": "+5.2%",
  "risques": ["risque 1", "risque 2", "risque 3"],
  "opportunites_majeures": ["opp 1", "opp 2", "opp 3"],
  "brvm_composite_100j": "de 230 à 245",
  "top10_achats": [{{"symbole": "SGBCI", "score": 82, "reco": "ACHAT"}}],
  "flop10": [{{"symbole": "XXXX", "score": 20, "reco": "VENTE"}}],
  "secteurs": [{{
    "secteur": "Banque",
    "nb_societes": 10,
    "perf_100j": "+8.2%",
    "vs_brvm": "+3.0%",
    "position": "Surperformance",
    "sentiment": "Positif",
    "risque": "Moyen",
    "prix_moyen": "12500",
    "societes": ["SGBCI", "BICICI", "BICIA-CI"]
  }}],
  "focus_industriels_telecoms": "analyse comparative",
  "implication_brvm": "concentration bancaire et diversification nécessaire",
  "top5_opportunites": [{{"symbole": "SGBCI", "score": 82, "raison": "raison"}}],
  "macro_international": "contexte mondial",
  "macro_africain": "contexte africain",
  "macro_uemoa": "contexte UEMOA",
  "impact_bourses_mondiales": "impact",
  "impact_indicateurs_brvm": "impact",
  "impact_societes_cotees": "impact",
  "synthese_macro": "synthèse narrative du contexte macro",
  "recommandation_macro": "recommandation finale macro",
  "sources_macro": [{{"source": "FMI", "resume": "résumé", "alerte": "vert"}}],
  "documents_officiels": [{{"type": "AG", "societe": "SGBCI", "date": "2026-04-30", "detail": "détail"}}],
  "alertes_google": [{{"mot_cle": "BRVM", "sentiment": "positif", "resume": "résumé", "score_pertinence": 8}}],
  "liquidite_haute": [{{
    "symbole": "SGBCI",
    "volume_moyen": "50M FCFA/j",
    "valeur_moyenne": "45M FCFA/j",
    "perf_100j": "+5%",
    "recommandation": "ACHAT",
    "detail": "volume 50M FCFA/j"
  }}],
  "liquidite_risque": [{{
    "symbole": "XXXX",
    "volume_moyen": "2M FCFA/j",
    "valeur_moyenne": "1.5M FCFA/j",
    "perf_100j": "-3%",
    "recommandation": "VENTE",
    "detail": "faible volume"
  }}],
  "matrice_risques": [{{"symbole": "SGBCI", "risque": "faible", "horizon": "MT", "detail": "stable"}}],
  "divergences": [{{"symbole": "XXXX", "detail": "tech haussier / fond baissier"}}],
  "classement_47": [{{"symbole": "SGBCI", "secteur": "Banque", "prix": "14500", "score": 82, "signal_tech": "Haussier", "signal_fond": "Positif", "reco": "ACHAT"}}],
  "portefeuille_defensif": [{{"symbole": "SGBCI", "poids": 20}}, {{"symbole": "CASH", "poids": 15}}],
  "portefeuille_equilibre": [{{"symbole": "SGBCI", "poids": 25}}, {{"symbole": "CASH", "poids": 10}}],
  "portefeuille_offensif": [{{"symbole": "SGBCI", "poids": 30}}, {{"symbole": "CASH", "poids": 5}}],
  "rsi_surachat": [{{
    "symbole": "XXXX",
    "rsi": 75,
    "cours_100j_max": "5490",
    "signal_tech": "Vente",
    "signal_fond": "Achat",
    "reco": "ACHAT",
    "detail": "RSI surachat, cours proche borne haute"
  }}],
  "rsi_survente": [{{
    "symbole": "YYYY",
    "rsi": 25,
    "cours_100j_min": "1200",
    "signal_tech": "Achat",
    "signal_fond": "Vente",
    "reco": "VENTE",
    "detail": "RSI survente, cours proche borne basse"
  }}],
  "cours_bornes_100j": [{{"symbole": "XXXX", "detail": "proche borne haute"}}],
  "divergences_tech_fond": [{{"symbole": "XXXX", "reco": "NEUTRE", "detail": "RSI surachat + fondamentaux faibles"}}],
  "valeur_a_eviter": "XXXX",
  "recommandations_brvm": ["Améliorer la liquidité", "Diversifier avec nouvelles IPO", "Transparence des données"],
  "conclusion_sentiment": "Haussier modéré avec prudence sectorielle",
  "top3_opportunites_conclusion": [{{"symbole": "SGBCI", "raison": "raison courte justifiée"}}],
  "top3_risques_conclusion": [{{"risque": "Risque de liquidité", "detail": "plusieurs titres peu liquides"}}],
  "recommandation_finale": "Le marché présente..."
}}"""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = message.content[0].text.strip()
    start = raw.find("{")
    end = raw.rfind("}") + 1
    if start == -1 or end == 0:
        return {}
    try:
        return json.loads(raw[start:end])
    except json.JSONDecodeError:
        return {}


# ── Helpers sûrs ─────────────────────────────────────────────────────────────

def _s(data, key, default="—"):
    v = data.get(key)
    return v if v is not None else default


def _sl(data, key):
    v = data.get(key)
    return v if isinstance(v, list) else []


def _get_score(item) -> str:
    for key in ("score", "score_composite", "score_technique"):
        v = item.get(key)
        if v is not None:
            return str(v)
    return "N/D"


def _get_score_float(item) -> float:
    for key in ("score", "score_composite", "score_technique"):
        v = item.get(key)
        if v is not None:
            try:
                return float(v)
            except (ValueError, TypeError):
                pass
    return 0.0


def _no_data(text) -> str:
    """Remplace 'Données insuffisantes' et variantes par une formulation neutre."""
    s = str(text).strip() if text else ""
    if s in ("", "—", "null", "None"):
        return "Aucune information significative n'a été collectée sur ce point."
    lower = s.lower()
    if any(x in lower for x in ("données insuffisantes", "données non disponibles", "non disponible")):
        return "Aucune information significative n'a été collectée sur ce point."
    return s


# ── Fonctions d'extraction et de synthèse ────────────────────────────────────

def extract_brvm_index_data(data: dict) -> dict:
    """Extrait les champs de l'indice BRVM Composite depuis les données structurées."""
    return {
        "niveau":    _s(data, "brvm_composite"),
        "variation": _s(data, "brvm_variation"),
        "perf_100j": _s(data, "perf_historique_100j"),
        "plus_haut": _s(data, "brvm_composite_plus_haut_100j"),
        "plus_bas":  _s(data, "brvm_composite_plus_bas_100j"),
        "tendance":  _s(data, "brvm_composite_tendance"),
        "range":     _s(data, "brvm_composite_100j"),
    }


def extract_market_cap_data(data: dict) -> dict:
    """Extrait les champs de la capitalisation boursière depuis les données structurées."""
    return {
        "valeur":    _s(data, "capitalisation"),
        "evolution": _s(data, "capitalisation_evolution_100j"),
        "pic":       _s(data, "capitalisation_pic_100j"),
        "plancher":  _s(data, "capitalisation_plancher_100j"),
    }


def summarize_section(items: list, fmt_fn, max_items: int = 5) -> list:
    """Retourne les max_items premiers éléments formatés par fmt_fn."""
    return [fmt_fn(item) for item in items[:max_items] if item]


# ── Section 1 — En-tête institutionnel ───────────────────────────────────────

def _section_entete(doc, d, date_str: str):
    """Section 1 — En-tête institutionnel."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r1 = p.add_run(_DESTINATAIRE + "\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(f"Date : {date_str}")
    r2.font.size = Pt(10)

    pi = d.get("_period_info")
    if pi and pi.get("nb_seances", 1) > 1:
        p_pi = doc.add_paragraph()
        p_pi.paragraph_format.space_after = Pt(6)
        _bold(
            p_pi,
            f"Période : {pi.get('date_debut', '—')} → {pi.get('date_fin', '—')}"
            f"   |   Séances : {pi.get('nb_seances', '—')}"
            f"   |   Fréquence : {pi.get('freq_label', '—')}",
            9, "283593",
        )


def _section_synthese_generale(doc, source_buckets, source_doc):
    """Section 2 — Synthèse générale (page 2 source).

    Recopie l'introduction du rapport source puis place les 2 graphiques de la
    synthèse (BRVM Composite + capitalisation) côte à côte avec leur commentaire
    extrait du source en légende.
    """
    _heading(doc, "SYNTHÈSE GÉNÉRALE")

    blocks = source_buckets.get("synthese", [])
    if not blocks:
        _para(doc, "Section Synthèse non identifiée dans le rapport source.")
        return

    # 1. Récupère le 1er paragraphe non-titre comme intro
    intro = ""
    for kind, el in blocks:
        if kind != "p":
            continue
        if _para_style(el).startswith("Heading"):
            continue
        txt = _para_text(el).strip()
        if txt:
            intro = txt
            break
    if intro:
        _para(doc, intro)

    # 2. Extrait les 2 images dans l'ordre du flux
    images = _image_blobs_in_blocks(source_doc, blocks)
    if len(images) < 2:
        _para(doc, f"Avertissement : seulement {len(images)} graphique(s) trouvé(s) dans la section Synthèse source.")
        return

    # 3. Récupère les commentaires juste AVANT chaque image (paragraphes
    #    contenant "BRVM Composite —" et "Capitalisation globale —")
    cap_indice, cap_capi = "", ""
    for kind, el in blocks:
        if kind != "p":
            continue
        txt = _para_text(el).strip()
        if not cap_indice and "BRVM Composite —" in txt:
            cap_indice = txt
        elif not cap_capi and "Capitalisation globale —" in txt:
            cap_capi = txt

    # 4. Insère les 2 images côte à côte
    _insert_images_side_by_side(
        doc,
        blob_left=images[0][0],
        blob_right=images[1][0],
        caption_left=cap_indice or "Évolution de l'indice BRVM Composite (100 derniers jours)",
        caption_right=cap_capi or "Évolution de la capitalisation boursière (100 derniers jours)",
        width_cm=8.0,
    )


# ── Section 4 — Analyse sectorielle ──────────────────────────────────────────

_SECTEUR_RX = {
    "nb":        re.compile(r"Nombre de sociétés\s*:\s*(\d+)"),
    "perf":      re.compile(r"Performance moyenne[^:]*:\s*([\-+]?[\d.,]+\s*%)"),
    "vs_brvm":   re.compile(r"vs BRVM[^:]*:\s*([\-+]?[\d.,]+\s*%)\s*([🟢🔴🟡⚪]?\s*\S*)"),
    "sentiment": re.compile(r"Sentiment général\s*:\s*([A-ZÉÈÀ ]+?)(?:\s*\||$)"),
    "risque":    re.compile(r"Risque moyen\s*:\s*([A-Za-zéè ]+?)(?:\s*\||$)"),
    "prix":      re.compile(r"Prix moyen\s*:\s*([\d\s.,]+)"),
}


def _parse_secteur_paragraph(text: str) -> dict:
    """Parse une ligne 'Nombre de sociétés: X | Performance moyenne (100j): Y% | …'."""
    out = {}
    for k, rx in _SECTEUR_RX.items():
        m = rx.search(text)
        if m:
            if k == "vs_brvm":
                out[k] = (m.group(1) or "").strip()
                out["vs_brvm_label"] = (m.group(2) or "").strip()
            else:
                out[k] = m.group(1).strip()
    return out


def _section_secteurs(doc, source_buckets, source_doc):
    """Section 3 — Analyse par secteur (pages 3-6 source).

    Construit un tableau 7 colonnes à partir des paragraphes Heading 3 'Secteur: X'
    et de la ligne de données qui suit. Pas de listing des sociétés.
    """
    _heading(doc, "ANALYSE PAR SECTEUR")
    blocks = source_buckets.get("secteurs", [])
    if not blocks:
        _para(doc, "Section sectorielle non identifiée dans le rapport source.")
        return

    # Intro : 1er paragraphe non-titre
    intro = ""
    for kind, el in blocks:
        if kind == "p" and not _para_style(el).startswith("Heading"):
            txt = _para_text(el).strip()
            if txt:
                intro = txt
                break
    if intro:
        _para(doc, intro)

    # Parcours : Heading3 "Secteur: X" suivi du paragraphe de données
    rows = []
    pending = None
    for kind, el in blocks:
        if kind != "p":
            continue
        st = _para_style(el)
        txt = _para_text(el).strip()
        if st == "Heading3" and txt.lower().startswith("secteur:"):
            pending = txt.split(":", 1)[1].strip()
        elif pending and not st.startswith("Heading"):
            data = _parse_secteur_paragraph(txt)
            if data:
                data["secteur"] = pending
                rows.append(data)
                pending = None

    if not rows:
        _para(doc, "Aucun secteur extrait du rapport source.")
        return

    def _perf_f(r):
        try:
            return float(str(r.get("perf", "0")).replace("%", "").replace("+", "").replace(",", ".").strip())
        except (ValueError, TypeError):
            return 0.0

    rows.sort(key=_perf_f, reverse=True)

    # Tableau 7 colonnes
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    _tbl_header(
        tbl,
        ["Secteur", "Nb sociétés", "Perf moy (100j)", "vs BRVM",
         "Sentiment", "Risque", "Prix moy"],
        "1A73E8", "FFFFFF",
    )
    for r in rows:
        tr = tbl.add_row()
        tr.cells[0].text = r.get("secteur", "—")
        tr.cells[1].text = r.get("nb", "—")
        tr.cells[2].text = r.get("perf", "—")
        tr.cells[3].text = r.get("vs_brvm", "—")
        tr.cells[4].text = r.get("sentiment", "—")
        tr.cells[5].text = r.get("risque", "—")
        tr.cells[6].text = r.get("prix", "—")
        pf = _perf_f(r)
        _cell_bg(tr.cells[2], "C6EFCE" if pf > 0 else ("FFC7CE" if pf < 0 else "FFEB9C"))
        for cell in tr.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()

    # Court commentaire global
    leader = rows[0]
    laggard = rows[-1] if len(rows) > 1 else None
    txt_lead = (
        f"Secteur leader : {leader.get('secteur', '—')} "
        f"({leader.get('perf', '—')} sur 100j, {leader.get('vs_brvm', '—')} vs BRVM). "
    )
    if laggard and laggard["secteur"] != leader["secteur"]:
        txt_lead += (
            f"Lanterne rouge : {laggard.get('secteur', '—')} ({laggard.get('perf', '—')}). "
        )
    txt_lead += (
        f"{len(rows)} secteurs couverts. La sélection sectorielle reste le levier "
        "principal de surperformance dans ce contexte."
    )
    _para(doc, txt_lead)


# ── Section 4 — Analyse de liquidité ─────────────────────────────────────────

def _read_source_tbl_data(tbl_xml):
    """Lit un <w:tbl> source en list[list[str]] (déduplique header répété)."""
    rows = tbl_xml.findall(qn("w:tr"))
    data = []
    for tr in rows:
        cells = []
        for tc in tr.findall(qn("w:tc")):
            cells.append("".join(t.text or "" for t in tc.iter(qn("w:t"))).strip())
        data.append(cells)
    if len(data) > 1 and data[0] == data[1]:
        data = [data[0]] + data[2:]
    return data


def _parse_fcfa(s: str) -> float:
    """Parse '29801676 FCFA' ou '29 801 676 FCFA' → 29801676.0."""
    if not s:
        return 0.0
    cleaned = re.sub(r"[^\d.,\-]", "", s).replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def _ticker_short(s: str) -> str:
    """'ETIT (ETIT - Ecobank...)' → 'ETIT'."""
    return s.split("(", 1)[0].strip() or s.strip()[:8]


def _section_liquidite(doc, source_buckets, source_doc):
    """Section 4 — Analyse de liquidité (page 8 source).

    Recopie les 2 tableaux TOP/FLOP du source (limités à 5 lignes) puis insère
    2 histogrammes côte à côte (TOP 5 et FLOP 5 en valeur de transactions).
    """
    _heading(doc, "ANALYSE DE LIQUIDITÉ")
    blocks = source_buckets.get("liquidite", [])
    if not blocks:
        _para(doc, "Section liquidité non identifiée dans le rapport source.")
        return

    # Intro
    intro = ""
    for kind, el in blocks:
        if kind == "p" and not _para_style(el).startswith("Heading"):
            txt = _para_text(el).strip()
            if txt and "ATTENTION" not in txt.upper():
                intro = txt
                break
    if intro:
        _para(doc, intro)

    # Récupère les 2 premières tables (TOP / FLOP)
    tables = [el for kind, el in blocks if kind == "tbl"]
    if len(tables) < 2:
        _para(doc, f"Avertissement : seulement {len(tables)} tableau(x) liquidité dans le source.")
        return

    top_data = _read_source_tbl_data(tables[0])
    flop_data = _read_source_tbl_data(tables[1])

    # Tableau TOP 5
    _heading(doc, "Top 5 — Titres les plus liquides", 2)
    top_5 = [top_data[0]] + top_data[1:6]
    n_cols = max(len(r) for r in top_5)
    tbl_t = doc.add_table(rows=1, cols=n_cols)
    tbl_t.style = "Table Grid"
    _tbl_header(tbl_t, (top_5[0] + [""] * n_cols)[:n_cols], "1A73E8", "FFFFFF")
    for r in top_5[1:]:
        tr = tbl_t.add_row()
        padded = (r + [""] * n_cols)[:n_cols]
        for i, v in enumerate(padded):
            tr.cells[i].text = v
            for p in tr.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()

    # Tableau FLOP 5
    _heading(doc, "Flop 5 — Titres les moins liquides (risque élevé)", 2)
    flop_5 = [flop_data[0]] + flop_data[1:6]
    n_cols = max(len(r) for r in flop_5)
    tbl_f = doc.add_table(rows=1, cols=n_cols)
    tbl_f.style = "Table Grid"
    _tbl_header(tbl_f, (flop_5[0] + [""] * n_cols)[:n_cols], "C0392B", "FFFFFF")
    for r in flop_5[1:]:
        tr = tbl_f.add_row()
        padded = (r + [""] * n_cols)[:n_cols]
        for i, v in enumerate(padded):
            tr.cells[i].text = v
            for p in tr.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()

    # 2 histogrammes côte à côte (TOP 5 + FLOP 5 en valeur de transactions)
    # Colonne attendue : "Valeur Moy." (index 2 dans le source)
    def _build_chart_data(rows):
        labels, values = [], []
        for r in rows[1:6]:
            if len(r) < 3:
                continue
            labels.append(_ticker_short(r[0]))
            values.append(_parse_fcfa(r[2]))
        return labels, values

    top_labels, top_values = _build_chart_data(top_data)
    flop_labels, flop_values = _build_chart_data(flop_data)

    if top_labels and flop_labels:
        try:
            png_top = _make_histogram_png(
                top_labels, top_values,
                "TOP 5 — Valeur moyenne de transactions (FCFA)",
                color="#1A73E8",
            )
            png_flop = _make_histogram_png(
                flop_labels, flop_values,
                "FLOP 5 — Valeur moyenne de transactions (FCFA)",
                color="#C0392B",
            )
            _insert_images_side_by_side(
                doc,
                blob_left=png_top, blob_right=png_flop,
                caption_left="Volumes les plus élevés du marché.",
                caption_right="Volumes les plus faibles — risque de slippage.",
                width_cm=8.2,
            )
        except Exception as e:
            logger.warning(f"Échec génération histogrammes liquidité: {e}")

    _para(doc,
          f"{len(top_data) - 1} titres à haute liquidité et {len(flop_data) - 1} titres à liquidité risquée "
          "ont été identifiés. La liquidité conditionne la flexibilité tactique du portefeuille : "
          "les volumes faibles amplifient le risque de slippage et de blocage en cas de retournement.")


# ── Section 5 — Analyse macro ────────────────────────────────────────────────

_MACRO_DROP_RX = re.compile(
    r"(?i)(données insuffisantes|aucune information|absolument\.|^---$)"
)



def _section_macro(doc, source_buckets, source_doc):
    """Section 5 — Analyse macro (pages 14-22 source).

    Lit directement le document source pour extraire, pour chaque plan
    (Mondial, Africain, Afrique de l'Ouest, UEMOA, BRVM) et chaque
    sous-thème (macro-économique, politique, financier) :
      - les actualités (style Listepuces)
      - le texte complet de l'impact BRVM (paragraphe après Titre3 '⚡')
    Structure source : Titre2 → Listepuces* → Titre3 → para(impact BRVM)
    """
    _heading(doc, "ANALYSE MACRO — CONTEXTE INTERNATIONAL, AFRICAIN & UEMOA")

    # Intro
    intro_blocks = source_buckets.get("macro", [])
    for kind, el in intro_blocks:
        if kind != "p":
            continue
        if _para_style(el).startswith(("Heading", "Titre")):
            continue
        txt = _para_text(el).strip()
        if not txt or _MACRO_DROP_RX.search(txt):
            continue
        _para(doc, txt)
        break

    # ── Couleurs impact ───────────────────────────────────────────────────────
    _IMPACT_BG = {"Positif": "C6EFCE", "Négatif": "FFC7CE", "Neutre": "FFEB9C"}
    _IMPACT_FG = {"Positif": "0F9D58", "Négatif": "D93025", "Neutre": "7D6608"}

    def _impact_level(txt: str) -> str:
        t = txt.lower()
        if any(w in t for w in ("positif", "favorable", "hausse", "soutien",
                                 "opportunit", "bénéfique")):
            return "Positif"
        if any(w in t for w in ("négatif", "negatif", "baisse", "pression",
                                 "risque", "fragilise", "pénalise", "incertitude",
                                 "ralentissement", "sorties de capitaux")):
            return "Négatif"
        return "Neutre"

    def _render_impact_block(doc, impact_txt: str):
        """Affiche le bloc ⚡ Impact BRVM coloré avec le texte complet."""
        if not impact_txt:
            return
        level = _impact_level(impact_txt)
        bg = _IMPACT_BG[level]
        fg = _IMPACT_FG[level]

        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        _cell_bg(cell, bg)
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side in ("top", "bottom", "left", "right"):
            e = OxmlElement(f"w:{side}")
            e.set(qn("w:w"), "100")
            e.set(qn("w:type"), "dxa")
            tcMar.append(e)
        tcPr.append(tcMar)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Label
        r1 = p.add_run(f"⚡ Impact BRVM — {level} : ")
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = _rgb(fg)

        # Texte complet de l'impact (nettoyé du préfixe éventuel)
        clean = re.sub(
            r"^⚡\s*Impact\s*(estimé\s*)?sur\s*la\s*BRVM\s*:?\s*",
            "", impact_txt, flags=re.IGNORECASE
        ).strip()
        r2 = p.add_run(clean)
        r2.font.size = Pt(9)
        r2.font.color.rgb = _rgb("222222")

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(3)

    def _scan_theme_bucket(bucket_key: str, theme_label: str):
        """
        Lit les blocs du bucket et reconstruit, pour chaque Titre2 (plan),
        la liste des actualités et le texte d'impact BRVM.
        Renvoie list[ (plan_name, actualites_list, impact_txt) ].
        """
        sub_blocks = source_buckets.get(bucket_key, [])
        if not sub_blocks:
            return []

        plans = []
        current_plan = None
        actus = []
        impact = ""
        next_is_impact = False   # True quand on vient de passer un Titre3 ⚡

        def _flush():
            if current_plan is not None:
                plans.append((current_plan, list(actus), impact))

        for kind, el in sub_blocks:
            if kind != "p":
                continue
            style = _para_style(el)
            raw = _para_text(el).strip()
            txt = _dedup_para(raw)
            if not txt:
                continue

            if style in ("Titre2", "Heading2"):
                _flush()
                # Nettoyer le nom du plan (enlever emoji + "Plan ")
                plan_clean = re.sub(r"^[^\w]+", "", txt)
                plan_clean = re.sub(r"^Plan\s+", "", plan_clean)
                current_plan = plan_clean
                actus = []
                impact = ""
                next_is_impact = False

            elif style in ("Titre3", "Heading3"):
                # C'est le titre "⚡ Impact estimé sur la BRVM"
                # Le paragraphe SUIVANT contient le texte de l'impact
                next_is_impact = True

            elif next_is_impact:
                # Premier paragraphe après le Titre3 = texte d'impact BRVM
                if not _MACRO_DROP_RX.search(txt):
                    impact = txt
                next_is_impact = False

            elif style in ("Listepuces", "ListParagraph", ""):
                if not _MACRO_DROP_RX.search(txt) and current_plan is not None:
                    # Nettoyer les ** du markdown
                    clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", txt)
                    actus.append(clean)

        _flush()
        return plans

    def _dedup_para(s: str) -> str:
        n = len(s)
        for d in (3, 2):
            if n % d == 0:
                p = s[:n // d]
                if s == p * d:
                    return p
        return s

    # ── 3 sous-thèmes ─────────────────────────────────────────────────────────
    themes = [
        ("macro_actu", "Actualités macro-économiques"),
        ("macro_pol",  "Actualités politiques & géopolitiques"),
        ("macro_fin",  "Actualités financières & marchés"),
    ]

    for bucket_key, theme_label in themes:
        plans = _scan_theme_bucket(bucket_key, theme_label)
        if not plans:
            continue

        _heading(doc, theme_label, 2)

        for plan_name, actus, impact_txt in plans:
            # Titre du plan
            p_plan = doc.add_paragraph()
            p_plan.paragraph_format.space_before = Pt(6)
            p_plan.paragraph_format.space_after = Pt(2)
            r_plan = p_plan.add_run(f"▶  {plan_name}")
            r_plan.bold = True
            r_plan.font.size = Pt(10)
            r_plan.font.color.rgb = _rgb("1A237E")

            # Actualités (liste à puces)
            if actus:
                for actu in actus[:4]:   # max 4 actualités par plan
                    p_a = doc.add_paragraph(style="List Bullet")
                    p_a.paragraph_format.space_before = Pt(1)
                    p_a.paragraph_format.space_after = Pt(1)
                    p_a.paragraph_format.left_indent = Pt(12)
                    r_a = p_a.add_run(actu[:300] + ("…" if len(actu) > 300 else ""))
                    r_a.font.size = Pt(9)
            else:
                p_nd = doc.add_paragraph()
                p_nd.paragraph_format.space_after = Pt(2)
                r_nd = p_nd.add_run("Aucune actualité disponible pour ce plan.")
                r_nd.italic = True
                r_nd.font.size = Pt(9)
                r_nd.font.color.rgb = _rgb("999999")

            # Impact BRVM — texte complet depuis le rapport source
            _render_impact_block(doc, impact_txt)

    # ── Synthèse & Recommandation finale ──────────────────────────────────────
    synth_blocks = source_buckets.get("macro_synth", [])
    if synth_blocks:
        _heading(doc, "Synthèse & Recommandation finale", 2)
        for kind, el in synth_blocks:
            if kind != "p":
                continue
            st = _para_style(el)
            if st in ("Titre1", "Heading1"):
                continue
            txt = _dedup_para(_para_text(el).strip())
            if not txt:
                continue
            if st in ("Titre3", "Heading3", "Titre4", "Heading4"):
                continue
            if txt.startswith(("🌐", "🌍", "🏦", "📈", "📋")) and len(txt) < 30:
                continue
            if _MACRO_DROP_RX.search(txt):
                continue
            _para(doc, txt)



_BRVM_SECTOR_TICKERS = {
    "bancaire": ("Banque", [
        ("BOAB", "BOA Burkina"), ("BOAC", "BOA Côte d'Ivoire"),
        ("BOAM", "BOA Mali"), ("BOAN", "BOA Niger"), ("BOAS", "BOA Sénégal"),
        ("SGBC", "SGBCI"), ("BICB", "BICI Bénin"), ("BICC", "BICI CI"),
        ("NSBC", "NSIA Banque CI"), ("SIBC", "SIB"), ("CBIBF", "Coris Bank Burkina"),
    ]),
    "agricole": ("Agroalimentaire", [
        ("SOGC", "SOGB"), ("SPHC", "SAPH"), ("PALC", "PALMCI"),
        ("SCRC", "SUCRIVOIRE"), ("STBC", "SITAB"), ("SLBC", "SOLIBRA"),
    ]),
    "énergie": ("Énergie & Distribution", [
        ("SHEC", "Vivo Energy CI"), ("SMBC", "SMB"), ("TTLC", "Total CI"),
        ("TTLS", "Total Sénégal"), ("SIVC", "SIVOA"),
    ]),
    "télécoms": ("Télécoms", [
        ("SNTS", "Sonatel"), ("ORAC", "Orange CI"), ("ONTBF", "Onatel BF"),
    ]),
    "industrie": ("Industrie", [
        ("CABC", "CABC"), ("FTSC", "Filtisac"), ("SDSC", "SODE CI"),
        ("SEMC", "SEMC"), ("STAC", "STA CI"),
    ]),
}

_SECTOR_KEYWORDS = {
    "bancaire": (
        "banque", "bancaire", "crédit", "bceao", "monétaire", "taux directeur",
        "refinancement", "liquidité bancaire", "boa ", "sgbci", "bici", "nsia",
        "coris", "ecobank", "sib ",
    ),
    "agricole": (
        "agricole", "agriculture", "cacao", "café", "coton", "hévéa", "caoutchouc",
        "palmier", "huile de palme", "sucre", "récolte", "campagne",
        "sogb", "saph", "palmci", "sucrivoire", "sitab", "solibra",
    ),
    "énergie": (
        "pétrole", "essence", "carburant", "hydrocarbure", "gaz", "raffinerie",
        "électricité", "énergie", "kwh", "barils", "brent", "shell", "total",
        "vivo energy",
    ),
    "télécoms": (
        "télécom", "telecom", "mobile money", "internet", "artci", "artp",
        "fibre", "5g", "4g", "sonatel", "orange ", "onatel", "moov",
    ),
    "industrie": (
        "industrie", "ciment", "usine", "production industrielle", "manufacture",
        "cimenterie", "filtisac", "sode", "fabrication",
    ),
}

_IMPACT_POS_KW = (
    "hausse", "croissance", "record", "succès", "bénéfice", "dividende",
    "investissement", "expansion", "augmentation", "amélioration", "positif",
    "accord", "signature", "rebond", "performance", "progression", "gain",
    "renforcement", "levée de fonds", "obligataire réussi", "souscription",
)
_IMPACT_NEG_KW = (
    "baisse", "perte", "déficit", "crise", "chute", "recul", "défaut",
    "sanction", "fermeture", "grève", "négatif", "alerte", "risque",
    "dégradation", "contraction", "ralentissement", "tension", "litige",
    "suspension", "report",
)


def _detect_brvm_sector(text: str):
    t = (text or "").lower()
    for key, keywords in _SECTOR_KEYWORDS.items():
        if any(kw in t for kw in keywords):
            return key
    return None


def _assess_brvm_impact(text: str) -> str:
    t = (text or "").lower()
    pos = sum(1 for k in _IMPACT_POS_KW if k in t)
    neg = sum(1 for k in _IMPACT_NEG_KW if k in t)
    if pos > neg:
        return "Positif"
    if neg > pos:
        return "Négatif"
    return "Neutre"


def _brvm_impact_explanation(sector_key, impact: str) -> str:
    sense = {
        "Positif": "favoriser",
        "Négatif": "peser sur",
        "Neutre": "influencer",
    }.get(impact, "influencer")
    sector_phrases = {
        "bancaire": (
            f"Une évolution du contexte bancaire et monétaire est susceptible de "
            f"{sense} les revenus d'intérêts, le coût du risque et la valorisation "
            f"des banques cotées à la BRVM."
        ),
        "agricole": (
            f"Les conditions de marché des matières premières agricoles peuvent "
            f"{sense} les revenus des producteurs cotés (cacao, hévéa, palmier, "
            f"sucre) et leur capacité à distribuer un dividende."
        ),
        "énergie": (
            f"L'évolution des prix de l'énergie et des hydrocarbures tend à "
            f"{sense} les marges des distributeurs et la consommation des ménages."
        ),
        "télécoms": (
            f"Les décisions réglementaires ou commerciales du secteur télécoms "
            f"peuvent {sense} l'ARPU, les investissements réseau et la rentabilité "
            f"des opérateurs cotés."
        ),
        "industrie": (
            f"Les dynamiques de production industrielle peuvent {sense} les "
            f"volumes, les marges et les carnets de commandes des industriels cotés."
        ),
    }
    if sector_key in sector_phrases:
        return sector_phrases[sector_key]
    return (
        "Une actualité d'ordre macroéconomique influence le sentiment global des "
        "investisseurs, la liquidité du marché et peut " + sense +
        " l'ensemble des compartiments de la cote BRVM."
    )

def _detect_brvm_sector(text: str):
    t = (text or "").lower()
    for key, keywords in _SECTOR_KEYWORDS.items():
        if any(kw in t for kw in keywords):
            return key
    return None


def _assess_brvm_impact(text: str) -> str:
    t = (text or "").lower()
    pos = sum(1 for k in _IMPACT_POS_KW if k in t)
    neg = sum(1 for k in _IMPACT_NEG_KW if k in t)
    if pos > neg:
        return "Positif"
    if neg > pos:
        return "Négatif"
    return "Neutre"


def _brvm_impact_explanation(sector_key, impact: str) -> str:
    sense = {
        "Positif": "favoriser",
        "Négatif": "peser sur",
        "Neutre": "influencer",
    }.get(impact, "influencer")
    sector_phrases = {
        "bancaire": (
            f"Une évolution du contexte bancaire et monétaire est susceptible de "
            f"{sense} les revenus d'intérêts, le coût du risque et la valorisation "
            f"des banques cotées à la BRVM."
        ),
        "agricole": (
            f"Les conditions de marché des matières premières agricoles peuvent "
            f"{sense} les revenus des producteurs cotés (cacao, hévéa, palmier, "
            f"sucre) et leur capacité à distribuer un dividende."
        ),
        "énergie": (
            f"L'évolution des prix de l'énergie et des hydrocarbures tend à "
            f"{sense} les marges des distributeurs et la consommation des ménages."
        ),
        "télécoms": (
            f"Les décisions réglementaires ou commerciales du secteur télécoms "
            f"peuvent {sense} l'ARPU, les investissements réseau et la rentabilité "
            f"des opérateurs cotés."
        ),
        "industrie": (
            f"Les dynamiques de production industrielle peuvent {sense} les "
            f"volumes, les marges et les carnets de commandes des industriels cotés."
        ),
    }
    if sector_key in sector_phrases:
        return sector_phrases[sector_key]
    return (
        "Une actualité d'ordre macroéconomique influence le sentiment global des "
        "investisseurs, la liquidité du marché et peut " + sense +
        " l'ensemble des compartiments de la cote BRVM."
    )


def _section_actualites(doc, source_buckets, source_doc):
    """Section 6 — Actualités du marché BRVM (pages 23-29 source).

    Résumé court : compte par catégorie de documents officiels + 1 tableau
    consolidé des 8 actualités les plus récentes.
    """
    _heading(doc, "ACTUALITÉS DU MARCHÉ BRVM")
    blocks = source_buckets.get("actualites", [])
    if not blocks:
        _para(doc, "Section actualités non identifiée dans le rapport source.")
        return

    # Intro
    intro = ""
    for kind, el in blocks:
        if kind == "p" and not _para_style(el).startswith("Heading"):
            txt = _para_text(el).strip()
            if txt:
                intro = txt
                break
    if intro:
        _para(doc, intro)

    # Catégorie courante = dernier Heading3 vu avant chaque tableau
    rx_cat = re.compile(r"\(\s*(\d+)\s*document")
    categories = []  # [(label, count, table_xml), ...]
    current_cat = None
    current_count = "?"
    for kind, el in blocks:
        if kind == "p" and _para_style(el) == "Heading3":
            txt = _para_text(el).strip()
            current_cat = txt
            m = rx_cat.search(txt)
            current_count = m.group(1) if m else "?"
        elif kind == "tbl" and current_cat:
            categories.append((current_cat, current_count, el))
            current_cat = None

    if not categories:
        _para(doc, "Aucune actualité corporate recensée.")
        return

    # Récap court par catégorie
    recap = " · ".join(f"{c[0].split('(')[0].strip()} : {c[1]}" for c in categories)
    _para(doc, f"Recensement par catégorie — {recap}.")

    # Consolidation : prend les 2 lignes les plus récentes de chaque catégorie
    consolidated = []
    for label, _, tbl_xml in categories:
        data = _read_source_tbl_data(tbl_xml)
        if len(data) < 2:
            continue
        cat_short = label.split("(")[0].strip()
        for row in data[1:3]:  # 2 lignes par catégorie
            if len(row) >= 5 and row[0]:
                consolidated.append([cat_short] + row[:5])

    if consolidated:
        _heading(doc, "Top actualités récentes (toutes catégories)", 2)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        _tbl_header(
            tbl,
            ["Catégorie", "Date", "Société", "Titre", "Impact", "Résumé"],
            "1A73E8", "FFFFFF",
        )
        for row in consolidated[:8]:
            tr = tbl.add_row()
            for i, v in enumerate(row[:6]):
                tr.cells[i].text = (v or "")[:80]
                for p in tr.cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
        doc.add_paragraph()

        # Analyse d'impact détaillée pour chaque actualité présentée
        _heading(doc, "Analyse d'impact détaillée sur la BRVM", 2)
        for row in consolidated[:8]:
            padded = (list(row) + [""] * 6)[:6]
            cat_short, date, societe, titre, _src_impact, resume = padded
            context_text = " ".join([cat_short, titre, resume])
            sector_key = _detect_brvm_sector(context_text)
            impact = _assess_brvm_impact(context_text)

            entete = doc.add_paragraph()
            entete.paragraph_format.space_after = Pt(2)
            _bold(entete, f"• {date or ''} — {societe or 'BRVM'}", 10, "1558A7")
            if titre:
                _normal(entete, f" : {titre}", 10)

            p_impact = doc.add_paragraph()
            p_impact.paragraph_format.space_after = Pt(2)
            _bold(p_impact, "⚡ Impact sur la BRVM : ", 10, "1A73E8")
            _normal(p_impact, f"{impact}.", 10)

            sector_info = _BRVM_SECTOR_TICKERS.get(sector_key)
            p_soc = doc.add_paragraph()
            p_soc.paragraph_format.space_after = Pt(2)
            _bold(p_soc, "Sociétés concernées : ", 10)
            if sector_info:
                secteur_label, tickers = sector_info
                soc_str = ", ".join(
                    f"{tk} ({secteur_label})" for tk, _nom in tickers[:6]
                )
                _normal(p_soc, soc_str, 10)
            else:
                _normal(
                    p_soc,
                    "Ensemble du marché BRVM (impact transversal sur la cote).",
                    10,
                )

            _para(doc, _brvm_impact_explanation(sector_key, impact))
            doc.add_paragraph()


# ── Section 9 — Alertes du jour (pages 34-37 source) ─────────────────────────

# Détecte une ligne d'en-tête "TICKER — Nom complet"
_ALERTE_HEADER_RX = re.compile(r"^([A-Z]{3,6})\s+—\s+(.+)$")


def _section_alertes(doc, source_buckets, source_doc):
    """Section 9 — Alertes du jour.

    Parse les paragraphes source en groupes (Ticker → liste d'alertes) et
    construit un tableau Symbole | Société | Alertes.
    """
    _heading(doc, "ALERTES DU JOUR")
    blocks = source_buckets.get("alertes", [])
    if not blocks:
        _para(doc, "Section alertes non identifiée dans le rapport source.")
        return

    intro = ""
    groups = []  # [(ticker, nom, [alertes]), ...]
    current = None
    for kind, el in blocks:
        if kind != "p":
            continue
        st = _para_style(el)
        txt = _para_text(el).strip()
        if not txt:
            continue
        if st == "Heading1":
            continue
        m = _ALERTE_HEADER_RX.match(txt)
        if m:
            ticker, nom = m.group(1), m.group(2)
            current = (ticker, nom, [])
            groups.append(current)
        elif current is None and not intro:
            intro = txt
        elif current is not None:
            current[2].append(txt)

    if intro:
        _para(doc, intro)
    if not groups:
        _para(doc, "Aucune alerte recensée.")
        return

    # Compteurs pour récap
    n_surachat = sum(1 for _, _, a in groups for x in a if "surachat" in x.lower())
    n_survente = sum(1 for _, _, a in groups for x in a if "survente" in x.lower())
    n_div = sum(1 for _, _, a in groups for x in a if "divergence" in x.lower())
    n_borne = sum(1 for _, _, a in groups for x in a if "plus haut 100j" in x.lower() or "plus bas 100j" in x.lower())
    _para(doc,
          f"{len(groups)} sociétés signalées — {n_surachat} surachat · {n_survente} survente · "
          f"{n_div} divergence tech/fond · {n_borne} cours en borne 100j.")

    # Tableau récapitulatif
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _tbl_header(tbl, ["Symbole", "Société", "Alerte(s)"], "C0392B", "FFFFFF")
    for ticker, nom, alertes in groups:
        tr = tbl.add_row()
        tr.cells[0].text = ticker
        # Nettoie le nom de la duplication "TICKER - Nom"
        nom_clean = nom.split(" - ", 1)[1] if " - " in nom else nom
        tr.cells[1].text = nom_clean
        tr.cells[2].text = "\n".join(alertes) if alertes else "—"
        # Coloration selon la 1ère alerte
        first = (alertes[0] if alertes else "").lower()
        if "surachat" in first:
            _cell_bg(tr.cells[0], "FFC7CE")
        elif "survente" in first:
            _cell_bg(tr.cells[0], "C6EFCE")
        elif "divergence" in first:
            _cell_bg(tr.cells[0], "FFEB9C")
        else:
            _cell_bg(tr.cells[0], "E8F0FB")
        for cell in tr.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    doc.add_paragraph()


# ── Section 7 — Classement 47 sociétés (page 30 source) ──────────────────────

def _section_classement(doc, source_buckets, source_doc):
    _heading(doc, "CLASSEMENT DES SOCIÉTÉS — SCORE COMPOSITE /100")
    blocks = source_buckets.get("classement", [])

    intro = ""
    for kind, el in blocks:
        if kind == "p" and not _para_style(el).startswith("Heading"):
            txt = _para_text(el).strip()
            if txt:
                intro = txt
                break
    if intro:
        _para(doc, intro)

    tables = [el for kind, el in blocks if kind == "tbl"]
    if not tables:
        _para(doc, "Tableau de classement non identifié dans le rapport source.")
        return
    _copy_source_table(doc, tables[0], header_bg="1A73E8", header_fg="FFFFFF")


# ── Section 8 — Portefeuilles modèles (pages 32-33 source) ───────────────────

def _section_portefeuilles(doc, source_buckets, source_doc):
    _heading(doc, "PORTEFEUILLES MODÈLES")
    blocks = source_buckets.get("portefeuilles", [])
    if not blocks:
        _para(doc, "Section portefeuilles non identifiée dans le rapport source.")
        return

    intro = ""
    for kind, el in blocks:
        if kind == "p" and not _para_style(el).startswith("Heading"):
            txt = _para_text(el).strip()
            if txt:
                intro = txt
                break
    if intro:
        _para(doc, intro)

    # Récupère les paires (Heading3 label + description, table)
    pairs = []
    pending_label = None
    pending_desc = ""
    for kind, el in blocks:
        if kind == "p":
            st = _para_style(el)
            txt = _para_text(el).strip()
            if st == "Heading3":
                pending_label = txt
                pending_desc = ""
            elif pending_label and not st.startswith("Heading") and txt:
                pending_desc = txt
        elif kind == "tbl" and pending_label:
            pairs.append((pending_label, pending_desc, el))
            pending_label = None
            pending_desc = ""

    bg_colors = ["BDE9F7", "C6EFCE", "FFC7CE"]  # Défensif/Équilibré/Offensif
    for i, (label, desc, tbl_xml) in enumerate(pairs):
        _heading(doc, label, 2)
        if desc:
            p = doc.add_paragraph(desc)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = _rgb("666666")
        bg = bg_colors[i] if i < len(bg_colors) else "1A73E8"
        _copy_source_table(doc, tbl_xml, header_bg=bg, header_fg="333333")


# ── Section 10 — Prédictions IA (pages 39-41 source) ─────────────────────────

_HORIZONS_IA = ("J+1", "J+3", "J+5", "J+10")


def _h1_blocks_matching(source_doc, anchors_upper):
    """Scan body for H1 sections whose title contains any of anchors_upper.
    Returns list of (kind, el) blocks of all matched sections."""
    body = source_doc.element.body
    out = []
    inside = False
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag not in ("p", "tbl"):
            continue
        if tag == "p" and _para_style(child) == "Heading1":
            txt = _para_text(child).upper()
            inside = any(a in txt for a in anchors_upper)
        if inside:
            out.append((tag, child))
    return out


def _scan_prediction_tables(source_doc):
    """Find tables that look like J+1..J+10 prediction tables.
    Header includes 'Prix prédit'; rows start with 'J+N'. Returns list[list[list[str]]]."""
    body = source_doc.element.body
    found = []
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag != "tbl":
            continue
        data = _read_source_tbl_data(child)
        if not data or not data[0]:
            continue
        header_blob = " ".join(data[0])
        if "Prix prédit" not in header_blob and "Prix predit" not in header_blob:
            continue
        first_col = " ".join((r[0] if r else "") for r in data[1:])
        if "J+" not in first_col:
            continue
        found.append(data)
    return found


def _aggregate_predictions_by_horizon(tables):
    """Aggregate Var % per horizon across companies.
    Returns {horizon: (tendance, niveau, confiance)}."""
    series = {h: [] for h in _HORIZONS_IA}
    rx_var = re.compile(r"([\-+]?\d+(?:[.,]\d+)?)\s*%")
    for data in tables:
        for row in data[1:]:
            if not row or len(row) < 2:
                continue
            first = row[0].strip()
            for h in _HORIZONS_IA:
                if first == h or first.startswith(h + " ") or first.startswith(h + "\t"):
                    last_cell = row[-1] if len(row) >= 6 else " ".join(row)
                    m = rx_var.search(last_cell)
                    if m:
                        try:
                            series[h].append(float(m.group(1).replace(",", ".")))
                        except ValueError:
                            pass
                    break
    out = {}
    for h in _HORIZONS_IA:
        vals = series[h]
        if not vals:
            out[h] = ("—", "—", "—")
            continue
        avg = sum(vals) / len(vals)
        n_pos = sum(1 for v in vals if v > 0)
        n_neg = sum(1 for v in vals if v < 0)
        n_tot = len(vals)
        if avg > 0.3:
            tendance = "Haussière"
            agree = n_pos
        elif avg < -0.3:
            tendance = "Baissière"
            agree = n_neg
        else:
            tendance = "Neutre"
            agree = max(n_pos, n_neg, n_tot - n_pos - n_neg)
        confiance = f"{int(round(100 * agree / n_tot))}% ({n_tot} titres)"
        niveau = f"{avg:+.2f}% (moyenne)"
        out[h] = (tendance, niveau, confiance)
    return out


def _section_predictions_ia(doc, source_doc, source_buckets):
    """Section 10 — Synthèse récapitulative des prédictions IA (J+1 → J+10)."""
    _heading(doc, "SYNTHÈSE RÉCAPITULATIVE DES PRÉDICTIONS IA (J+1 → J+10)")

    blocks = source_buckets.get("predictions", [])
    print(f"  [Note/predictions] bucket 'predictions' contient {len(blocks)} blocs")

    if not blocks:
        alt = ["PREDICTION", "PRÉDICTION", "J+1", "HORIZON", "PRÉVISION", "PREVISION"]
        blocks = _h1_blocks_matching(source_doc, alt)
        print(f"  [Note/predictions] fallback ancres alternatives → {len(blocks)} blocs")

    pred_tables_data = _scan_prediction_tables(source_doc)
    print(f"  [Note/predictions] {len(pred_tables_data)} tableaux J+1→J+10 détectés dans le corps")

    j10_data = _scan_j10_per_ticker(source_doc)
    print(f"  [Note/predictions] {len(j10_data)} prédictions J+10 extraites par ticker")

    horizon_data = _aggregate_predictions_by_horizon(pred_tables_data)
    if not any(h[0] != "—" for h in horizon_data.values()) and not blocks:
        _para(doc, "Section prédictions IA non identifiée dans le rapport source.")
        return

    # Intro
    intro = ""
    for kind, el in blocks:
        if kind == "p" and not _para_style(el).startswith("Heading"):
            txt = _para_text(el).strip()
            if txt:
                intro = txt
                break
    if not intro and pred_tables_data:
        intro = (
            f"Synthèse agrégée à partir des prédictions individuelles de {len(pred_tables_data)} "
            "valeurs cotées. Les modèles GRU/LSTM publient un intervalle de confiance à 90% ; "
            "les chiffres ci-dessous reflètent la moyenne des variations attendues par horizon."
        )
    if intro:
        _para(doc, intro[:600] + ("…" if len(intro) > 600 else ""))

    # ── Tableau synthèse par horizon
    _heading(doc, "Tableau de synthèse par horizon", 2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _tbl_header(tbl, ["Horizon", "Tendance prévue", "Variation moyenne", "Confiance"], "673AB7", "FFFFFF")
    for h in _HORIZONS_IA:
        tendance, niveau, confiance = horizon_data[h]
        tr = tbl.add_row()
        tr.cells[0].text = h
        tr.cells[1].text = tendance
        tr.cells[2].text = niveau
        tr.cells[3].text = confiance
        t_low = tendance.lower()
        if "hauss" in t_low:
            _cell_bg(tr.cells[1], "C6EFCE")
        elif "baiss" in t_low:
            _cell_bg(tr.cells[1], "FFC7CE")
        elif tendance != "—":
            _cell_bg(tr.cells[1], "FFEB9C")
        for cell in tr.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

    bucket_tables = [el for kind, el in blocks if kind == "tbl"]
    if bucket_tables:
        _heading(doc, "Détails des prédictions par horizon (source)", 2)
        _copy_source_table(doc, bucket_tables[0], header_bg="673AB7", header_fg="FFFFFF", max_rows=10)

    if j10_data:
        hausse_top5 = sorted(j10_data, key=lambda x: x["var_j10"], reverse=True)[:5]
        baisse_top5 = sorted(j10_data, key=lambda x: x["var_j10"])[:5]

        # ── TOP 5 HAUSSE
        _heading(doc, "TOP 5 — Cours attendus en HAUSSE à J+10", 2)
        tbl_h = doc.add_table(rows=1, cols=5)
        tbl_h.style = "Table Grid"
        _tbl_header(tbl_h,
                    ["Société", "Cours actuel (FCFA)", "Prix prédit J+10", "Var. %", "Borne haute"],
                    "0F9D58", "FFFFFF")
        for item in hausse_top5:
            tr = tbl_h.add_row()
            tr.cells[0].text = item["ticker"]
            tr.cells[1].text = _fmt_prix(item["cours_actuel"])
            tr.cells[2].text = _fmt_prix(item["prix_j10"])
            tr.cells[3].text = f"{item['var_j10']:+.1f}%"
            tr.cells[4].text = _fmt_prix(item["borne_haut"])
            _cell_bg(tr.cells[3], "C6EFCE")
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        doc.add_paragraph()

        # ── TOP 5 BAISSE
        _heading(doc, "TOP 5 — Cours attendus en BAISSE à J+10", 2)
        tbl_b = doc.add_table(rows=1, cols=5)
        tbl_b.style = "Table Grid"
        _tbl_header(tbl_b,
                    ["Société", "Cours actuel (FCFA)", "Prix prédit J+10", "Var. %", "Borne basse"],
                    "D93025", "FFFFFF")
        for item in baisse_top5:
            tr = tbl_b.add_row()
            tr.cells[0].text = item["ticker"]
            tr.cells[1].text = _fmt_prix(item["cours_actuel"])
            tr.cells[2].text = _fmt_prix(item["prix_j10"])
            tr.cells[3].text = f"{item['var_j10']:+.1f}%"
            tr.cells[4].text = _fmt_prix(item["borne_bas"])
            _cell_bg(tr.cells[3], "FFC7CE")
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        doc.add_paragraph()

        # ── Tableau complet TOUTES sociétés
        _heading(doc, "Tableau complet — Prédictions J+10 de toutes les sociétés", 2)
        tbl_all = doc.add_table(rows=1, cols=6)
        tbl_all.style = "Table Grid"
        _tbl_header(tbl_all,
                    ["#", "Société", "Cours actuel", "Prix prédit J+10", "Var. %", "Intervalle [bas – haut]"],
                    "283593", "FFFFFF")
        all_sorted = sorted(j10_data, key=lambda x: x["var_j10"], reverse=True)
        for rank, item in enumerate(all_sorted, 1):
            tr = tbl_all.add_row()
            tr.cells[0].text = str(rank)
            tr.cells[1].text = item["ticker"]
            tr.cells[2].text = _fmt_prix(item["cours_actuel"])
            tr.cells[3].text = _fmt_prix(item["prix_j10"])
            var = item["var_j10"]
            tr.cells[4].text = f"{var:+.1f}%"
            tr.cells[5].text = f"{_fmt_prix(item['borne_bas'])} – {_fmt_prix(item['borne_haut'])}"
            _cell_bg(tr.cells[4], "C6EFCE" if var > 0 else ("FFC7CE" if var < 0 else "FFEB9C"))
            row_bg = "F5F8FF" if rank % 2 == 0 else "FFFFFF"
            for ci in (0, 1, 2, 3, 5):
                _cell_bg(tr.cells[ci], row_bg)
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)
        doc.add_paragraph()

    # Conclusion
    n_haut = sum(1 for h in _HORIZONS_IA if "hauss" in horizon_data[h][0].lower())
    n_bas  = sum(1 for h in _HORIZONS_IA if "baiss" in horizon_data[h][0].lower())
    if n_haut > n_bas:
        reco = (
            "Le modèle IA penche majoritairement vers une orientation haussière sur les "
            "prochains horizons. Un repositionnement progressif vers les valeurs à score "
            "composite élevé peut être envisagé, en privilégiant les titres bénéficiant "
            "d'une confluence technique et fondamentale positive."
        )
    elif n_bas > n_haut:
        reco = (
            "Le modèle IA anticipe un environnement majoritairement baissier sur les "
            "prochains horizons. La priorité doit être donnée à la préservation du capital, "
            "au renforcement des positions défensives et à la constitution d'une réserve "
            "de liquidités permettant d'agir aux points de retournement."
        )
    else:
        reco = (
            "Les prédictions IA présentent un équilibre entre hausse et baisse selon "
            "l'horizon retenu. Une posture tactique flexible — combinant allocation neutre "
            "et sélectivité accrue sur les meilleurs scores — est recommandée."
        )
    _para(doc, reco)


# ── Section 11 — Analyse financière comparative par secteur (pages 42-66) ────

_RX_PER         = re.compile(r"(?:'?PER'?|P/E)\s*:?\s*'?\s*([\-+]?\d+(?:[.,]\d+)?)", re.IGNORECASE)
_RX_ROE         = re.compile(r"'?ROE'?\s*:?\s*'?\s*([\-+]?\d+(?:[.,]\d+)?)\s*%?", re.IGNORECASE)
_RX_MARGE       = re.compile(r"marge_?nette\s*:?\s*'?\s*([\-+]?\d+(?:[.,]\d+)?)\s*%?", re.IGNORECASE)
_RX_CROISS      = re.compile(r"(?:variation_ca|croissance_?ca|croissance\s+(?:du\s+)?CA)\s*:?\s*'?\s*([\-+]?\d+(?:[.,]\d+)?)\s*%?", re.IGNORECASE)


def _parse_ratios_from_text(text: str) -> dict:
    """Parse PER/ROE/Marge nette/Croissance CA from concatenated ratio text."""
    out = {}
    for label, rx, suffix in [
        ("PER",           _RX_PER,    ""),
        ("ROE",           _RX_ROE,    "%"),
        ("Marge nette",   _RX_MARGE,  "%"),
        ("Croissance CA", _RX_CROISS, "%"),
    ]:
        m = rx.search(text)
        if m:
            v = m.group(1).strip().rstrip(".,")
            out[label] = v + suffix
    return out


def _build_sector_ticker_map(source_buckets) -> dict:
    """From the 'secteurs' bucket, build {sector_name: [tickers]} by parsing the
    'Sociétés: TICKER (...), TICKER (...)' tail of each sector summary line."""
    blocks = source_buckets.get("secteurs", [])
    mapping = {}
    current = None
    for kind, el in blocks:
        if kind != "p":
            continue
        st = _para_style(el)
        txt = _para_text(el).strip()
        if st == "Heading3" and txt.lower().startswith("secteur:"):
            current = txt.split(":", 1)[1].strip()
            mapping.setdefault(current, [])
        elif current and "Sociétés" in txt and ":" in txt:
            tail = txt.split("Sociétés", 1)[1].split(":", 1)[-1]
            mapping[current] = re.findall(r"\b([A-Z]{2,6})\s*\(", tail)
    return {k: v for k, v in mapping.items() if v}


def _scan_company_ratio_tables(source_doc) -> list:
    """Scan body for per-company ratio tables. Returns [(ticker, ratios_dict), ...].
    A ratio table contains marge_nette/capitaux_propres/ROE and is preceded
    (somewhere above) by a Heading2 of the form 'N. TICKER - ...'."""
    body = source_doc.element.body
    out = []
    current_ticker = None
    rx_ticker = re.compile(r"^\s*\d+\.\s*([A-Z]{2,6})\b")
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            if _para_style(child) == "Heading2":
                m = rx_ticker.match(_para_text(child))
                if m:
                    current_ticker = m.group(1)
            continue
        if tag != "tbl" or not current_ticker:
            continue
        data = _read_source_tbl_data(child)
        if not data:
            continue
        text_blob = " ".join(" ".join(r) for r in data)
        if not any(k in text_blob for k in ("marge_nette", "capitaux_propres", "ROE")):
            continue
        ratios = _parse_ratios_from_text(text_blob)
        if ratios:
            out.append((current_ticker, ratios))
    return out


def _aggregate_ratios_by_sector(sector_map: dict, company_ratios: list) -> dict:
    """Aggregate per-company ratios by sector. Returns {sector: {label: 'avg%', ...}}."""
    by_ticker = {}
    for ticker, ratios in company_ratios:
        by_ticker.setdefault(ticker, {}).update(ratios)

    def _to_float(s):
        try:
            return float(str(s).replace("%", "").replace(",", ".").strip())
        except (ValueError, AttributeError):
            return None

    out = {}
    for sector, tickers in sector_map.items():
        agg = {}
        for label in ("PER", "ROE", "Marge nette", "Croissance CA"):
            vals = []
            for t in tickers:
                v = by_ticker.get(t, {}).get(label)
                if v is not None:
                    f = _to_float(v)
                    if f is not None:
                        vals.append(f)
            if vals:
                avg = sum(vals) / len(vals)
                suffix = "" if label == "PER" else "%"
                agg[label] = f"{avg:.2f}{suffix} ({len(vals)})"
            else:
                agg[label] = "—"
        out[sector] = agg
    return out




def _extract_all_company_data(source_doc) -> list:
    """
    Parcourt toutes les sections sociétés dans le rapport source et extrait :
    - Ticker, Nom court, Secteur
    - PER, ROA, Liquidité générale, Ratio endettement
    - Bêta numérique réel
    - Capitalisation boursière (numérique)
    Retourne list[dict], 1 dict par société.
    """
    import re as _re2
    elements = list(source_doc.element.body.iterchildren())

    def _ps(el):
        pPr = el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        return pPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '') if pPr is not None else ''

    def _pt2(el): return ''.join(n.text or '' for n in el.iter() if n.text)

    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s

    def _rt(tbl_el):
        rows = []
        for tr in tbl_el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
            cells = [_dd(''.join(n.text or '' for n in tc.iter() if n.text).strip())
                     for tc in tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')]
            if cells: rows.append(cells)
        return rows

    def _fnum(s, pct=False):
        """Parse un nombre (avec ou sans %)."""
        if not s or s in ('—', '', 'N/A', 'null'): return None
        s2 = str(s).replace('%', '').replace(' ', '').replace('\u00a0', '')
        s2 = s2.replace('\u202f', '').strip()
        # Format français : virgule décimale
        if ',' in s2 and '.' not in s2: s2 = s2.replace(',', '.')
        elif ',' in s2 and '.' in s2: s2 = s2.replace(',', '')
        try: return float(s2)
        except (ValueError, TypeError): return None

    rx_h2 = _re2.compile(r'^\s*(\d+)\.\s*([A-Z]{2,6})\s*[-—]')
    companies = []
    current = None
    current_end = None

    # Trouver toutes les sections sociétés
    sections = []
    for i, child in enumerate(elements):
        if child.tag.split('}')[-1] != 'p': continue
        s = _ps(child); t = _dd(_pt2(child).strip())
        if s in ('Titre2', 'Heading2'):
            m = rx_h2.match(t)
            if m:
                sections.append((i, m.group(2), t))

    for sec_i, (idx, ticker, title) in enumerate(sections):
        next_idx = sections[sec_i + 1][0] if sec_i + 1 < len(sections) else min(idx + 180, len(elements))
        entry = {
            'ticker': ticker,
            'nom': _dd(title.split('-')[-1].strip())[:50] if '-' in title else ticker,
            'per': None, 'roa': None, 'liquidite': None,
            'endettement': None, 'beta': None, 'capi': None,
        }

        for i in range(idx, next_idx):
            child = elements[i]
            tag = child.tag.split('}')[-1]

            if tag == 'p':
                txt = _dd(_pt2(child).strip())
                # Capitalisation depuis la table identité (texte ou table)
                if entry['capi'] is None and 'capitalisation' in txt.lower():
                    m_capi = _re2.search(r'capitalisation[^:]*:\s*([\d\s,\.]+)\s*(Mds|milliards|M|millions)?', txt, _re2.IGNORECASE)
                    if m_capi:
                        v = _fnum(m_capi.group(1).replace(' ', ''))
                        unit = (m_capi.group(2) or '').lower()
                        if v:
                            if 'mds' in unit or 'milliards' in unit: v *= 1
                            elif 'm' in unit or 'millions' in unit: v /= 1000
                            entry['capi'] = v
                # PER depuis les paragraphes narratifs (PARTIE 0, table identité)
                if entry['per'] is None and 'PER' in txt:
                    m_per = _re2.search(r'\bPER\s+(?:de\s+)?([0-9]+[,\.][0-9]+)x?', txt, _re2.IGNORECASE)
                    if not m_per:
                        m_per = _re2.search(r'\bPER\b[^0-9]*([0-9]+[,\.][0-9]+)', txt)
                    if m_per:
                        v = _fnum(m_per.group(1))
                        if v and 0 < v < 200: entry['per'] = v
                # ROA depuis les paragraphes
                if entry['roa'] is None and 'ROA' in txt:
                    m_roa = _re2.search(r'\bROA\b[^0-9%]*([0-9]+[,\.][0-9]+)\s*%?', txt)
                    if m_roa:
                        v = _fnum(m_roa.group(1))
                        if v and 0 < v < 100: entry['roa'] = v

            elif tag == 'tbl':
                data = _rt(child)
                if not data: continue

                # Table score risque → bêta numérique
                full = str(data)
                if ('êta' in full or 'Bêt' in full) and entry['beta'] is None:
                    for row in data:
                        cell = ' '.join(row)
                        m_b = _re2.search(r'β\s*=\s*([-]?[\d.,]+)', cell)
                        if m_b:
                            entry['beta'] = _fnum(m_b.group(1))
                            break

                # Table identité → capitalisation
                if len(data[0]) >= 3 and 'IDENTITÉ' in str(data[0]):
                    for row in data:
                        cell_txt = ' '.join(row)
                        if 'capitalisation' in cell_txt.lower() and entry['capi'] is None:
                            m_c = _re2.search(r'(\d+[\d\s,\.]+)\s*(Mds|Mrd|milliards)', cell_txt, _re2.IGNORECASE)
                            if m_c:
                                v = _fnum(m_c.group(1).replace(' ', ''))
                                if v: entry['capi'] = v

                # Tables financières → PER, ROA, liquidité, endettement
                if len(data[0]) >= 2 and data[0][0] == 'Indicateur':
                    for row in data[1:]:
                        if len(row) < 2: continue
                        ind = row[0].lower()
                        val = row[1] if len(row) > 1 else ''
                        if 'per' in ind and entry['per'] is None:
                            entry['per'] = _fnum(val)
                        elif 'roa' == ind.strip() and entry['roa'] is None:
                            entry['roa'] = _fnum(val)
                        elif 'liquidité générale' in ind and entry['liquidite'] is None:
                            entry['liquidite'] = _fnum(val)
                        elif ('ratio endett' in ind or 'endettement' in ind) and entry['endettement'] is None:
                            entry['endettement'] = _fnum(val)

        companies.append(entry)

    return [c for c in companies if c['ticker']]


def _section_analyse_financiere_sectorielle(doc, source_doc, source_buckets):
    """Section 11 — Analyse financière comparative par secteur."""
    _heading(doc, "ANALYSE FINANCIÈRE COMPARATIVE PAR SECTEUR")

    blocks = source_buckets.get("analyse_financiere", [])
    print(f"  [Note/analyse_fin] bucket 'analyse_financiere' contient {len(blocks)} blocs")

    if not blocks:
        alt = ["ANALYSE FINANCIÈRE", "ANALYSE FINANCIERE", "COMPARATIF", "PER", "ROE", "RATIO"]
        blocks = _h1_blocks_matching(source_doc, alt)
        print(f"  [Note/analyse_fin] fallback ancres alternatives → {len(blocks)} blocs")

    sector_map = _build_sector_ticker_map(source_buckets)
    company_ratios = _scan_company_ratio_tables(source_doc)
    print(f"  [Note/analyse_fin] {len(sector_map)} secteurs avec liste de tickers, "
          f"{len(company_ratios)} tables de ratios société détectées")

    sector_ratios = _aggregate_ratios_by_sector(sector_map, company_ratios)
    has_any_ratio = any(v != "—" for r in sector_ratios.values() for v in r.values())

    if not blocks and not has_any_ratio:
        _para(doc, "Section analyse financière sectorielle non identifiée dans le rapport source.")
        return

    intro = ""
    for kind, el in blocks:
        if kind == "p" and not _para_style(el).startswith("Heading"):
            txt = _para_text(el).strip()
            if txt:
                intro = txt
                break
    if not intro and has_any_ratio:
        intro = (
            f"Synthèse agrégée des ratios financiers publiés pour {len(company_ratios)} sociétés "
            f"réparties sur {sum(1 for v in sector_ratios.values() if any(x != '—' for x in v.values()))} secteurs. "
            "Les moyennes sont calculées sur les valeurs disponibles ; le nombre entre parenthèses "
            "indique l'échantillon par ratio."
        )
    if intro:
        _para(doc, intro[:600] + ("…" if len(intro) > 600 else ""))

    _heading(doc, "Ratios financiers clés par secteur", 2)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    _tbl_header(tbl, ["Secteur", "PER", "ROE", "Marge nette", "Croissance CA"], "1A73E8", "FFFFFF")

    sorted_sectors = sorted(
        sector_ratios.items(),
        key=lambda kv: sum(1 for v in kv[1].values() if v != "—"),
        reverse=True,
    )
    for name, r in sorted_sectors[:12]:
        tr = tbl.add_row()
        tr.cells[0].text = name[:40]
        tr.cells[1].text = r.get("PER", "—")
        tr.cells[2].text = r.get("ROE", "—")
        tr.cells[3].text = r.get("Marge nette", "—")
        tr.cells[4].text = r.get("Croissance CA", "—")
        for cell in tr.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()

    # ── Tableau comparatif complet — tous les secteurs d'activité ────────────
    _heading(doc, "Tableau comparatif des indicateurs clés — tous les secteurs d'activité", 2)
    _para(doc,
          "Ce tableau présente les ratios moyens par secteur calculés sur l'ensemble des "
          "sociétés disposant de données publiées. Les moyennes sont pondérées par le nombre "
          "de valeurs disponibles (indiqué entre parenthèses).", 9)

    # Construire le tableau complet : 1 ligne par secteur, colonnes enrichies
    # En-têtes : Secteur | Nb sociétés | PER moy. | ROE moy. | Marge nette moy. | Croissance CA moy. | Signal
    def _signal_secteur(r: dict) -> tuple[str, str]:
        """Retourne (signal_txt, bg_color) selon les ratios du secteur."""
        roe_s = r.get("ROE", "—")
        marge_s = r.get("Marge nette", "—")
        try:
            roe_v = float(str(roe_s).split("(")[0].replace("%","").replace(",",".").strip())
            marge_v = float(str(marge_s).split("(")[0].replace("%","").replace(",",".").strip())
            if roe_v >= 12 and marge_v >= 8:
                return "✅ Solide", "C6EFCE"
            if roe_v >= 6 or marge_v >= 4:
                return "🟡 Correct", "FFEB9C"
            return "⚠️ Faible", "FFC7CE"
        except (ValueError, AttributeError):
            return "—", "F5F5F5"

    # Compter les tickers par secteur
    sector_ticker_counts = {s: len(t) for s, t in _build_sector_ticker_map(source_buckets).items()}

    tbl_comp = doc.add_table(rows=1, cols=7)
    tbl_comp.style = "Table Grid"
    _tbl_header(tbl_comp,
                ["Secteur", "Nb val.", "PER moy.", "ROE moy.", "Marge nette", "Croissance CA", "Signal"],
                "1A237E", "FFFFFF")

    # Trier par ROE décroissant
    def _sort_key(kv):
        r = kv[1]
        try:
            return float(str(r.get("ROE","0")).split("(")[0].replace("%","").replace(",",".").strip())
        except (ValueError, AttributeError):
            return -999.0

    for name, r in sorted(sector_ratios.items(), key=_sort_key, reverse=True):
        signal_txt, signal_bg = _signal_secteur(r)
        nb = sector_ticker_counts.get(name, "—")
        tr = tbl_comp.add_row()
        tr.cells[0].text = name[:35]
        tr.cells[1].text = str(nb)
        tr.cells[2].text = r.get("PER", "—")
        tr.cells[3].text = r.get("ROE", "—")
        tr.cells[4].text = r.get("Marge nette", "—")
        tr.cells[5].text = r.get("Croissance CA", "—")
        tr.cells[6].text = signal_txt
        _cell_bg(tr.cells[6], signal_bg)
        _cell_bg(tr.cells[0], "EBF0FA")
        for cell in tr.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()

    # ── Tableau complet — toutes les sociétés avec leurs ratios ──────────────
    if company_ratios:
        _heading(doc, "Tableau comparatif des indicateurs clés — toutes les sociétés", 2)
        _para(doc,
              "Ratios individuels publiés dans le rapport source pour chaque société cotée "
              "(valeurs extraites des sections 'Analyse financière' et 'Données financières').", 9)

        tbl2 = doc.add_table(rows=1, cols=6)
        tbl2.style = "Table Grid"
        _tbl_header(tbl2,
                    ["Société", "Secteur", "PER", "ROE", "Marge nette", "Croissance CA"],
                    "283593", "FFFFFF")

        # Construire dict ticker → secteur depuis sector_map
        ticker_to_sector = {}
        for sect, tickers in _build_sector_ticker_map(source_buckets).items():
            for t in tickers:
                ticker_to_sector[t] = sect

        seen = set()
        rank = 0
        for ticker, ratios in company_ratios:
            if ticker in seen:
                continue
            seen.add(ticker)
            filled = sum(1 for v in ratios.values() if v)
            if filled < 1:
                continue
            rank += 1
            tr = tbl2.add_row()
            tr.cells[0].text = ticker
            tr.cells[1].text = ticker_to_sector.get(ticker, "—")[:25]
            tr.cells[2].text = ratios.get("PER", "—")
            tr.cells[3].text = ratios.get("ROE", "—")
            tr.cells[4].text = ratios.get("Marge nette", "—")
            tr.cells[5].text = ratios.get("Croissance CA", "—")
            # Fond alterné
            row_bg = "F5F8FF" if rank % 2 == 0 else "FFFFFF"
            for ci, cell in enumerate(tr.cells):
                _cell_bg(cell, "EBF0FA" if ci == 0 else row_bg)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8.5)
        doc.add_paragraph()

    n_sectors_with_data = sum(1 for v in sector_ratios.values() if any(x != "—" for x in v.values()))
    _para(doc,
          f"{n_sectors_with_data} secteurs documentés sur {len(sector_ratios) or '—'}. "
          "La sélection des titres doit s'appuyer sur la combinaison ROE élevé + marge nette positive "
          "+ croissance du CA, tout en restant attentif aux niveaux de valorisation (PER) qui "
          "conditionnent le potentiel de réappréciation à moyen terme.")

    # ── Données enrichies depuis le rapport source ────────────────────────────
    all_co = _extract_all_company_data(source_doc)
    all_co_valid = [c for c in all_co if c.get('ticker')]

    # ── 1. Comparaison ROE, ROA, Liquidité, Endettement ──────────────────────
    _heading(doc, "Comparaison ROE, ROA, Liquidité & Endettement", 2)
    _para(doc,
          "Indicateurs de rentabilité et de structure financière par société, "
          "extraits des états financiers officiels publiés dans le rapport source.", 9)

    roe_data = sorted(
        [(c['ticker'], c['roa'], c.get('liquidite'), c.get('endettement'))
         for c in all_co_valid if c.get('roa') is not None],
        key=lambda x: x[1] or 0, reverse=True
    )
    # Tableau comparatif ROE/ROA/Liquidité/Endettement (depuis _scan_company_ratio_tables)
    tbl_roe = doc.add_table(rows=1, cols=5)
    tbl_roe.style = "Table Grid"
    _tbl_header(tbl_roe, ["Société", "ROA (%)", "Liquidité gén. (%)", "Endettement (%)", "Signal"], "1A237E", "FFFFFF")

    def _fmt_pct(v):
        if v is None: return "—"
        return f"{v:.1f}%"

    def _risk_color_val(v, seuil_vert, seuil_rouge, inverse=False):
        if v is None: return "F5F5F5"
        if not inverse:
            if v >= seuil_vert: return "C6EFCE"
            if v <= seuil_rouge: return "FFC7CE"
        else:
            if v <= seuil_vert: return "C6EFCE"
            if v >= seuil_rouge: return "FFC7CE"
        return "FFEB9C"

    for ticker, roa, liq, endet in roe_data[:20]:  # top 20 par ROA
        sig = "✅" if (roa or 0) > 5 else ("⚠️" if (roa or 0) > 1 else "🔴")
        tr = tbl_roe.add_row()
        tr.cells[0].text = ticker
        tr.cells[1].text = _fmt_pct(roa)
        tr.cells[2].text = _fmt_pct(liq)
        tr.cells[3].text = _fmt_pct(endet)
        tr.cells[4].text = sig
        _cell_bg(tr.cells[1], _risk_color_val(roa, 5, 1))
        _cell_bg(tr.cells[2], _risk_color_val(liq, 120, 80))
        _cell_bg(tr.cells[3], _risk_color_val(endet, 60, 80, inverse=True))
        for cell in tr.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()

    # ── 2. Section Risque ─────────────────────────────────────────────────────
    _heading(doc, "Section Risque — Vue d'ensemble", 2)
    _para(doc,
          "Classement des sociétés par score de risque calculé (agrège volatilité, bêta, "
          "liquidité, divergence signaux et stabilité des rendements). "
          "Score 0 = risque nul, 100 = risque maximal.", 9)

    risk_data = sorted(
        [(c['ticker'], c.get('beta')) for c in all_co_valid if c.get('beta') is not None],
        key=lambda x: abs(x[1]) if x[1] is not None else 0, reverse=True
    )

    # Copier le tableau risque depuis le bucket source (déjà dans la section récap risques)
    # Mais ici on fait un résumé plus concis : top risqués + top sûrs
    recap_blocks = source_buckets.get("recap_risques", [])
    if recap_blocks:
        # Copier les tables importantes (vue d'ensemble + top 5 chaque côté)
        n_tbl = 0
        for kind, el in recap_blocks:
            if kind == "tbl" and n_tbl < 3:  # vue d'ensemble + top5 risqués + top5 sûrs
                _copy_source_table(doc, el, header_bg="1A237E", header_fg="FFFFFF", max_rows=10)
                doc.add_paragraph()
                n_tbl += 1
    else:
        _para(doc, "Données de risque non disponibles dans le rapport source.", 9)

    # ── 3. Classement PER ─────────────────────────────────────────────────────
    _heading(doc, "Classement PER — Meilleurs et mauvais PER du marché", 2)
    _para(doc,
          "PER = Cours / BPA. Un PER faible peut indiquer une sous-valorisation, "
          "un PER élevé une surévaluation ou des attentes de forte croissance. "
          "Seules les sociétés avec un PER positif et disponible sont classées.", 9)

    per_data = [(c['ticker'], c['per']) for c in all_co_valid
                if c.get('per') is not None and c['per'] > 0]
    per_data.sort(key=lambda x: x[1])

    if per_data:
        # Tableau top 10 meilleurs PER + top 10 mauvais
        top_per = per_data[:10]
        bad_per = per_data[-10:][::-1]

        tbl_per = doc.add_table(rows=1, cols=4)
        tbl_per.style = "Table Grid"
        _tbl_header(tbl_per, ["Rang", "Société", "PER", "Signal"], "1A237E", "FFFFFF")

        _para(doc, "🟢 Meilleurs PER (sous-évaluation potentielle)", 9)
        for rank, (ticker, per_v) in enumerate(top_per, 1):
            tr = tbl_per.add_row()
            tr.cells[0].text = str(rank)
            tr.cells[1].text = ticker
            tr.cells[2].text = f"{per_v:.1f}x"
            sig = "🟢 Attractif" if per_v < 8 else ("🟡 Correct" if per_v < 15 else "🔴 Cher")
            tr.cells[3].text = sig
            _cell_bg(tr.cells[2], "C6EFCE" if per_v < 8 else ("FFEB9C" if per_v < 15 else "FFC7CE"))
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for run in p.runs: run.font.size = Pt(8.5)
        doc.add_paragraph()

        _para(doc, "🔴 PER les plus élevés (surévaluation potentielle)", 9)
        tbl_per2 = doc.add_table(rows=1, cols=4)
        tbl_per2.style = "Table Grid"
        _tbl_header(tbl_per2, ["Rang", "Société", "PER", "Signal"], "D93025", "FFFFFF")
        for rank, (ticker, per_v) in enumerate(bad_per, 1):
            tr = tbl_per2.add_row()
            tr.cells[0].text = str(rank)
            tr.cells[1].text = ticker
            tr.cells[2].text = f"{per_v:.1f}x"
            sig = "🟢 Attractif" if per_v < 8 else ("🟡 Correct" if per_v < 15 else "🔴 Cher")
            tr.cells[3].text = sig
            _cell_bg(tr.cells[2], "C6EFCE" if per_v < 8 else ("FFEB9C" if per_v < 15 else "FFC7CE"))
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for run in p.runs: run.font.size = Pt(8.5)
        doc.add_paragraph()

        # Tableau classement complet par PER
        _heading(doc, "Classement complet par PER", 3)
        tbl_per3 = doc.add_table(rows=1, cols=3)
        tbl_per3.style = "Table Grid"
        _tbl_header(tbl_per3, ["Rang", "Société", "PER"], "283593", "FFFFFF")
        for rank, (ticker, per_v) in enumerate(per_data, 1):
            tr = tbl_per3.add_row()
            tr.cells[0].text = str(rank)
            tr.cells[1].text = ticker
            tr.cells[2].text = f"{per_v:.1f}x"
            bg = "C6EFCE" if per_v < 8 else ("FFEB9C" if per_v < 15 else "FFC7CE")
            _cell_bg(tr.cells[2], bg)
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for run in p.runs: run.font.size = Pt(8.5)
        _para(doc, f"{len(per_data)} sociétés avec PER disponible sur {len(all_co_valid)}.", 8)
        doc.add_paragraph()
    else:
        _para(doc, "Données PER non disponibles dans ce rapport.", 9)

    # ── 4. Classement Capitalisation ─────────────────────────────────────────
    _heading(doc, "Classement par Capitalisation Boursière", 2)
    _para(doc,
          "Classement des sociétés par capitalisation boursière décroissante "
          "(en milliards FCFA). Source : rapport source BRVM.", 9)

    capi_data = sorted(
        [(c['ticker'], c['capi']) for c in all_co_valid if c.get('capi') is not None and c['capi'] > 0],
        key=lambda x: x[1], reverse=True
    )

    if capi_data:
        tbl_capi = doc.add_table(rows=1, cols=3)
        tbl_capi.style = "Table Grid"
        _tbl_header(tbl_capi, ["Rang", "Société", "Capitalisation (Mds FCFA)"], "1A237E", "FFFFFF")
        for rank, (ticker, capi_v) in enumerate(capi_data, 1):
            tr = tbl_capi.add_row()
            tr.cells[0].text = str(rank)
            tr.cells[1].text = ticker
            tr.cells[2].text = f"{capi_v:,.0f}".replace(",", " ")
            bg = "C6EFCE" if capi_v > 500 else ("FFEB9C" if capi_v > 50 else "F5F5F5")
            _cell_bg(tr.cells[2], bg)
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for run in p.runs: run.font.size = Pt(8.5)
        _para(doc, f"{len(capi_data)} sociétés avec capitalisation disponible.", 8)
        doc.add_paragraph()
    else:
        _para(doc, "Données de capitalisation non disponibles.", 9)

    # ── 5. Classement Bêta ────────────────────────────────────────────────────
    _heading(doc, "Classement par Bêta (β) — Sensibilité au marché", 2)
    _para(doc,
          "Bêta = sensibilité du titre aux mouvements du BRVM Composite. "
          "β>1 = plus volatile que le marché (agressif). "
          "β≈1 = neutre. β<1 = défensif. β<0 = contre-cyclique.", 9)

    beta_ranked = sorted(
        [(c['ticker'], c['beta']) for c in all_co_valid if c.get('beta') is not None],
        key=lambda x: x[1]
    )

    if beta_ranked:
        tbl_beta = doc.add_table(rows=1, cols=4)
        tbl_beta.style = "Table Grid"
        _tbl_header(tbl_beta, ["Rang", "Société", "Bêta (β)", "Profil"], "1A237E", "FFFFFF")
        for rank, (ticker, beta_v) in enumerate(beta_ranked, 1):
            tr = tbl_beta.add_row()
            tr.cells[0].text = str(rank)
            tr.cells[1].text = ticker
            tr.cells[2].text = f"{beta_v:.4f}"
            if beta_v < 0:
                profil = "🔵 Contre-cyclique"; bg = "E8F0FB"
            elif beta_v < 0.5:
                profil = "🟢 Très défensif"; bg = "C6EFCE"
            elif beta_v < 0.8:
                profil = "🟢 Défensif"; bg = "EBF7EE"
            elif beta_v < 1.2:
                profil = "🟡 Neutre"; bg = "FFEB9C"
            elif beta_v < 1.5:
                profil = "🟠 Agressif"; bg = "FFE0B2"
            else:
                profil = "🔴 Très agressif"; bg = "FFC7CE"
            tr.cells[3].text = profil
            _cell_bg(tr.cells[2], bg)
            _cell_bg(tr.cells[3], bg)
            for cell in tr.cells:
                for p in cell.paragraphs:
                    for run in p.runs: run.font.size = Pt(8.5)
        _para(doc, f"{len(beta_ranked)} sociétés avec bêta numérique disponible.", 8)
        doc.add_paragraph()
    else:
        _para(doc, "Données bêta non disponibles.", 9)


# ── Build complet ─────────────────────────────────────────────────────────────

def _dedup_cell(s: str) -> str:
    """Déduplique une cellule triplée 'XYZXYZXYZ' → 'XYZ'."""
    n = len(s)
    for div in (3, 2):
        if n % div == 0:
            part = s[: n // div]
            if s == part * div:
                return part
    return s



def _parse_brvm_price(raw: str):
    """Parse un prix BRVM depuis une cellule triplée. Virgule = sep. milliers FR."""
    s = _dedup_cell(raw).strip().replace(" ", "").replace("\u00a0", "").replace(",", "")
    try:
        return float(s)
    except ValueError:
        return None



def _scan_j10_per_ticker(source_doc) -> list:
    """Extrait les prédictions J+10 de chaque société cotée."""
    body = source_doc.element.body
    elements = list(body.iterchildren())
    rx_ticker = re.compile(r"^\s*\d+\.\s*([A-Z]{2,6})\s*[-\u2014]")
    rx_var = re.compile(r"([+\-]?\d+(?:[.,]\d+)?)\s*%")
    current_ticker = None
    results = []
    seen_tickers = set()

    for i, child in enumerate(elements):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            style = _para_style(child)
            raw = _para_text(child).strip()
            txt = _dedup_cell(raw)
            if style in ("Titre2", "Heading2"):
                m = rx_ticker.match(txt)
                if m:
                    current_ticker = m.group(1)
        elif tag == "tbl" and current_ticker:
            data = _read_source_tbl_data(child)
            if not data or not data[0]:
                continue
            header = " ".join(_dedup_cell(c) for c in data[0])
            if "Prix" not in header:
                continue
            first_col = " ".join(_dedup_cell(r[0]) for r in data[1:] if r)
            if "J+" not in first_col:
                continue
            j10_row = None
            for row in data[1:]:
                if not row:
                    continue
                horizon = _dedup_cell(row[0]).strip()
                if horizon.startswith("J+10") and j10_row is None:
                    j10_row = row
            if j10_row is None or len(j10_row) < 5:
                continue
            try:
                prix_j10  = _parse_brvm_price(j10_row[3])
                borne_bas = _parse_brvm_price(j10_row[2])
                borne_haut = _parse_brvm_price(j10_row[4])
                var_raw = _dedup_cell(j10_row[5] if len(j10_row) > 5 else j10_row[-1])
                m_var = rx_var.search(var_raw)
                var_j10 = float(m_var.group(1).replace(",", ".")) if m_var else None
                if prix_j10 is None or var_j10 is None:
                    continue
                cours_actuel = round(prix_j10 / (1 + var_j10 / 100)) if var_j10 != -100 else None
                if current_ticker not in seen_tickers:
                    seen_tickers.add(current_ticker)
                    results.append({
                        "ticker": current_ticker, "cours_actuel": cours_actuel,
                        "prix_j10": prix_j10, "var_j10": var_j10,
                        "borne_bas": borne_bas, "borne_haut": borne_haut,
                    })
            except (ValueError, ZeroDivisionError, IndexError):
                pass
    return results



def _fmt_prix(v) -> str:
    if v is None:
        return "—"
    return f"{v:,.0f}".replace(",", " ")



def _dedup_text(s: str) -> str:
    """Déduplique un texte triplé 'XYZXYZXYZ' → 'XYZ'."""
    n = len(s)
    for d in (3, 2):
        if n % d == 0:
            p = s[:n // d]
            if s == p * d:
                return p
    return s



    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    _setup_header_footer(doc, date_str)

    _section_entete(doc, data, date_str)                            # 1 — En-tête institutionnel
    _section_synthese_generale(doc, source_buckets, source_doc)     # 2 — Synthèse + 2 graphiques
    _section_secteurs(doc, source_buckets, source_doc)              # 3 — Analyse par secteur
    _section_liquidite(doc, source_buckets, source_doc)             # 4 — Liquidité + 2 histogrammes
    _section_macro(doc, source_buckets, source_doc)                 # 5 — Macro condensé
    _section_matrice_risque(doc, source_buckets, source_doc)        # 6 — Récapitulatif risques
    _section_actualites(doc, source_buckets, source_doc)            # 7 — Actualités résumé
    _section_classement(doc, source_buckets, source_doc)            # 8 — Classement /100
    _section_portefeuilles(doc, source_buckets, source_doc)         # 9 — 3 portefeuilles
    _section_alertes(doc, source_buckets, source_doc)               # 10 — Alertes du jour
    _section_predictions_ia(doc, source_doc, source_buckets)        # 11 — Prédictions IA (J+1 → J+10)
    _section_analyse_financiere_sectorielle(doc, source_doc, source_buckets)  # 12 — Analyse financière sectorielle

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Point d'entrée ────────────────────────────────────────────────────────────


def _section_matrice_risque(doc, source_buckets, source_doc):
    """
    Section Risques — RÉCAPITULATIF DES RISQUES — TOUTES LES SOCIÉTÉS.

    Lit le bucket 'recap_risques' (ancre "RÉCAPITULATIF DES RISQUES") du rapport
    source et reproduit fidèlement :
      1. Texte d'introduction
      2. Vue d'ensemble des niveaux de risque (tableau + commentaire)
      3. Top 5 sociétés les plus risquées (tableau + détails)
      4. Top 5 sociétés les moins risquées (tableau + détails)
      5. Tableau complet — toutes les sociétés classées par score de risque
    """
    _heading(doc, "RÉCAPITULATIF DES RISQUES — TOUTES LES SOCIÉTÉS")

    blocks = source_buckets.get("recap_risques", [])
    if not blocks:
        # Fallback : ancrage alternatif
        blocks = source_buckets.get("matrice_risque", [])
    if not blocks:
        _para(doc, "Section récapitulatif des risques non identifiée dans le rapport source.")
        return

    # ── Couleurs niveau de risque ─────────────────────────────────────────────
    def _risk_color(level: str) -> tuple:
        l = level.lower()
        if "très élevé" in l or "tres eleve" in l:
            return "FF0000", "FFC7CE"
        if "élevé" in l or "eleve" in l:
            return "D93025", "FFC7CE"
        if "moyen" in l:
            return "E37400", "FFEB9C"
        return "0F9D58", "C6EFCE"  # faible par défaut

    current_h2 = None
    tables_done = {}   # heading → liste de tables déjà copiées

    for kind, el in blocks:
        if kind == "p":
            st  = _para_style(el)
            raw = _para_text(el).strip()
            txt = _dedup_text(raw)
            if not txt:
                continue

            if st in ("Heading1", "Titre1"):
                continue   # titre principal déjà affiché

            if st in ("Heading2", "Titre2"):
                current_h2 = txt
                _heading(doc, txt, 2)
                continue

            if st in ("Heading3", "Titre3"):
                _heading(doc, txt, 3)
                continue

            if st in ("ListBullet", "Listepuces", "ListParagraph"):
                p_b = doc.add_paragraph()
                p_b.paragraph_format.space_before = Pt(1)
                p_b.paragraph_format.space_after  = Pt(1)
                p_b.paragraph_format.left_indent  = Pt(12)
                r_b = p_b.add_run(f"• {txt}")
                r_b.font.size = Pt(9)
                continue

            # Paragraphe ordinaire
            _para(doc, txt)

        elif kind == "tbl":
            _copy_source_table(
                doc, el,
                header_bg="1A237E", header_fg="FFFFFF",
                max_rows=60,        # tableau complet = 48 lignes + en-tête
            )
            doc.add_paragraph()



def _build_docx(data: dict, source_doc, source_buckets: dict, date_str: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    _setup_header_footer(doc, date_str)

    _section_entete(doc, data, date_str)                            # 1 — En-tête institutionnel
    _section_synthese_generale(doc, source_buckets, source_doc)     # 2 — Synthèse + 2 graphiques
    _section_secteurs(doc, source_buckets, source_doc)              # 3 — Analyse par secteur
    _section_liquidite(doc, source_buckets, source_doc)             # 4 — Liquidité + 2 histogrammes
    _section_macro(doc, source_buckets, source_doc)                 # 5 — Macro condensé
    _section_matrice_risque(doc, source_buckets, source_doc)        # 6 — Récapitulatif risques
    _section_actualites(doc, source_buckets, source_doc)            # 7 — Actualités résumé
    _section_classement(doc, source_buckets, source_doc)            # 8 — Classement /100
    _section_portefeuilles(doc, source_buckets, source_doc)         # 9 — 3 portefeuilles
    _section_alertes(doc, source_buckets, source_doc)               # 10 — Alertes du jour
    _section_predictions_ia(doc, source_doc, source_buckets)        # 11 — Prédictions IA (J+1 → J+10)
    _section_analyse_financiere_sectorielle(doc, source_doc, source_buckets)  # 12 — Analyse financière sectorielle

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Point d'entrée ────────────────────────────────────────────────────────────

def generate(docs_bytes, freq: str = "JOUR", period_info: dict = None) -> tuple:
    """
    Génère la Note Stratégique BRVM depuis un ou plusieurs .docx source.
    docs_bytes : bytes (un seul doc) ou list[bytes] (plusieurs docs, plus récent en premier).
    Retourne (filename: str, docx_bytes: bytes).

    Le rapport source contient déjà toutes les sections (Synthèse, Sectoriel,
    Liquidité, Macro, Actualités, Classement, Portefeuilles, Alertes) sous
    forme de Heading 1 — la note les recopie/condense plutôt que de les
    ré-extraire via Claude.
    """
    if isinstance(docs_bytes, bytes):
        docs_bytes = [docs_bytes]

    date_str = date.today().strftime("%d/%m/%Y")
    date_file = date.today().strftime("%Y%m%d")

    freq_suffix = {"JOUR": "JOUR", "HEBDO": "HEBDO", "MENSUEL": "MENSUEL",
                   "TRIM": "TRIM", "ANNUEL": "ANNUEL"}.get(freq, freq)

    print(f"  [Note/{freq}] Ouverture du rapport source ({len(docs_bytes)} doc(s))...")
    source_doc = _open_source_doc(docs_bytes)
    source_buckets = _split_source_by_h1(source_doc)
    print(f"  [Note/{freq}] Sections détectées : "
          + ", ".join(k for k, v in source_buckets.items() if v))

    # Données minimales pour _section_entete (period_info / freq)
    data = {}
    if period_info:
        data["_period_info"] = period_info
        data["_freq"] = freq

    print(f"  [Note/{freq}] Construction du document Word...")
    docx_bytes = _build_docx(data, source_doc, source_buckets, date_str)

    filename = f"Note_Strategique_BRVM_{date_file}_{freq_suffix}.docx"
    return filename, docx_bytes
