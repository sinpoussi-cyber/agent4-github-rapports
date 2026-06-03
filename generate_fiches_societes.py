import io
import math
import os
import re
import random
import subprocess
import tempfile
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter

from PIL import Image as _PILImage

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from extractor import extract_all
from enricher import enrich

_MARGIN_CM = 1.5

# ── Namespaces OOXML pour extraction d'images ─────────────────────────────────
_A_NS      = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R_NS      = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_BLIP_TAG  = "{%s}blip"  % _A_NS
_EMBED_ATTR = "{%s}embed" % _R_NS

# RE pour détecter les tickers BRVM dans le texte (2-8 majuscules + suffixe pays optionnel)
_TICKER_RE = re.compile(
    r'\b([A-Z]{2,8}(?:CI|BF|SN|ML|TG|GN|BN|NE|GW)?)\b'
)
# Mots courants à exclure pour éviter les faux positifs
# (recommandations, niveaux de risque, abréviations financières, etc.)
_TICKER_STOPWORDS = {
    # Recommandations / signaux
    "ACHAT", "VENTE", "NEUTRE", "FORT", "FAIBLE", "CONSERVER", "SURVEILLER",
    "PRUDENCE", "EVITER", "ÉVITER", "POSITIF", "NEGATIF", "NÉGATIF",
    "HAUSSIER", "BAISSIER",
    # Niveaux de risque / qualité
    "RISQUE", "ELEVE", "ÉLEVÉ", "MOYEN", "BAS", "HAUTE", "FORTE",
    # Indicateurs techniques
    "RSI", "MM", "EMA", "MACD", "STOCH", "BOLL", "ADX", "OBV", "CCI",
    "WRI", "BBG", "ATH", "ATL",
    # Statistiques / mesures
    "CV", "VOL", "MAX", "MIN", "AVG", "TOP", "FLOP", "NB",
    # Entités financières génériques
    "BRVM", "FCFA", "UEMOA", "PDG", "DG", "CA", "PNB", "RBE", "ROE", "ROA",
    "PER", "SG", "BNP", "CIB", "USD", "EUR", "XOF", "BCE", "FMI", "BM",
    "SA", "SAS", "SARL", "NV", "SE", "AG", "LTD", "PLC", "INC", "LLC",
    # Secteurs / types
    "BANQUE", "ASSURANCE", "TELECOM", "INDUSTRIE", "AGRICOLE", "FINANCE",
    # Divers rapports
    "NOTE", "FICHE", "ALERTE", "SCORE", "RANG", "ANALYSE",
}
# Distance maximale (en éléments body) entre le dernier ticker vu et une image.
# Le rapport source place les graphiques ~20-50 éléments après le titre de section.
# On utilise une valeur large car on filtre ensuite avec known_tickers.
_IMG_MAX_DIST = 80


# ── Helpers bas niveau ────────────────────────────────────────────────────────

def _rgb(hex6: str) -> RGBColor:
    return RGBColor(int(hex6[:2], 16), int(hex6[2:4], 16), int(hex6[4:], 16))


def _cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_margins(cell, twips: int = 40):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(twips))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _cw(cell, text, bold=False, size=8, color=None, bg=None, align=None):
    if bg:
        _cell_bg(cell, bg)
    _cell_margins(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = _rgb(color)
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    return run


def _cp(p, before=0, after=1):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def _narrative(doc, text: str, size: int = 9, italic: bool = False, color: str = None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = _rgb(color)
    return p


# ── Helpers visuels ───────────────────────────────────────────────────────────

def _section_heading(doc, text: str, color: str = "1A73E8"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = _rgb(color)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _sub_heading(doc, text: str, color: str = "444444"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb(color)


def _key_bloc(doc, label: str, texte: str, bg: str, fg: str = "1A1A1A"):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, bg)
    _cell_margins(cell, 80)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = _rgb(fg)
    r2 = p.add_run(texte)
    r2.font.size = Pt(9)
    r2.font.color.rgb = _rgb("333333")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(3)


def _add_separator(doc, color: str = "CCCCCC"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Helpers cours / chart ─────────────────────────────────────────────────────

def _to_float(v):
    """Conversion robuste : nombre, '1 440', '1,440.5', '1.440,5'."""
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace(" ", "").replace("\u00a0", "")
    if not s:
        return None
    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


# ── Extraction des images depuis le document Word source ──────────────────────

def _detect_img_type(blob: bytes) -> str:
    """Détecte le format d'image à partir des magic bytes."""
    if blob[:4] == b"\x89PNG":
        return "png"
    if blob[:3] == b"\xff\xd8\xff":
        return "jpeg"
    if blob[:4] == b"GIF8":
        return "gif"
    if blob[:4] == bytes([0xD7, 0xCD, 0xC6, 0x9A]):
        return "wmf"
    if blob[:4] == bytes([0x01, 0x00, 0x00, 0x00]):
        return "emf"
    return "unknown"


def _blob_to_png(blob: bytes, ticker: str = "?") -> bytes | None:
    """
    Convertit un blob image (PNG, JPEG, GIF, EMF, WMF) en bytes PNG prêts
    à être insérés dans python-docx via add_picture().

    - PNG / JPEG / GIF  → PIL (direct, sans aller sur disque)
    - EMF / WMF         → LibreOffice headless (fichier temporaire)
    - Inconnu           → retour None (fallback matplotlib)
    """
    img_type = _detect_img_type(blob)

    if img_type in ("png", "jpeg", "gif"):
        try:
            img = _PILImage.open(io.BytesIO(blob))
            if img.mode not in ("RGB", "RGBA"):
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            result = buf.getvalue()
            print(f"  [Fiches/ImgExtract] {ticker} : {img_type.upper()} → PNG "
                  f"{img.size[0]}×{img.size[1]}px ({len(result)} bytes)")
            return result
        except Exception as exc:
            print(f"  [Fiches/ImgExtract] {ticker} : PIL échec sur {img_type} — {exc}")
            # Pour PNG, renvoyer le blob brut si PIL échoue (peut fonctionner quand même)
            return blob if img_type == "png" else None

    if img_type in ("emf", "wmf"):
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                in_path = os.path.join(tmpdir, f"img.{img_type}")
                with open(in_path, "wb") as fh:
                    fh.write(blob)
                r = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "png",
                     "--outdir", tmpdir, in_path],
                    capture_output=True, timeout=20,
                )
                out_path = os.path.join(tmpdir, "img.png")
                if os.path.exists(out_path):
                    with open(out_path, "rb") as fh:
                        result = fh.read()
                    print(f"  [Fiches/ImgExtract] {ticker} : {img_type.upper()} → PNG "
                          f"via LibreOffice ({len(result)} bytes)")
                    return result
                else:
                    print(f"  [Fiches/ImgExtract] {ticker} : LibreOffice n'a pas produit "
                          f"de PNG (stderr: {r.stderr[:120].decode(errors='replace')})")
        except Exception as exc:
            print(f"  [Fiches/ImgExtract] {ticker} : LibreOffice échec — {exc}")
        return None

    print(f"  [Fiches/ImgExtract] {ticker} : format inconnu ({blob[:4].hex()}) — ignoré")
    return None


def _annotate_source_chart(png_bytes: bytes, s: dict) -> bytes:
    """
    Prend le PNG brut extrait du rapport Word source (graphique avec cours réels
    + prédictions J+1→J+10) et ajoute un bandeau de synthèse sous l'image avec :
      ▲ Plus haut 100j  |  Cours actuel (perf%)  |  ▼ Plus bas 100j

    Les valeurs sont lues directement depuis le dict de la société (s).
    Si PIL échoue ou si les données sont insuffisantes, retourne le PNG original intact.
    """
    from PIL import Image as _Img, ImageDraw, ImageFont
    import io as _io

    try:
        img = _Img.open(_io.BytesIO(png_bytes)).convert("RGBA")
        W, H = img.size

        # ── Récupérer les valeurs ─────────────────────────────────────────────
        def _fval(key):
            v = s.get(key)
            if v is None:
                return None
            try:
                return float(str(v).replace(",", ".").replace(" ", "")
                             .replace("\u00a0", "").replace("%", "").strip())
            except ValueError:
                return None

        plus_haut    = _fval("plus_haut_100j")
        plus_bas     = _fval("plus_bas_100j")
        cours_actuel = _fval("cours") or _fval("cours_fin")
        perf_100j    = s.get("perf_100j", "")

        # Si données minimales absentes, retourner le PNG tel quel
        if plus_haut is None and plus_bas is None and cours_actuel is None:
            return png_bytes

        # ── Polices ───────────────────────────────────────────────────────────
        _FONT_PATH_BOLD = "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"
        _FONT_PATH_REG  = "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
        try:
            font_val = ImageFont.truetype(_FONT_PATH_BOLD, max(16, W // 52))
            font_lbl = ImageFont.truetype(_FONT_PATH_REG,  max(12, W // 72))
        except Exception:
            font_val = ImageFont.load_default()
            font_lbl = font_val

        # ── Construire le bandeau ─────────────────────────────────────────────
        BAND_H = max(52, W // 20)
        new_img = _Img.new("RGBA", (W, H + BAND_H), (255, 255, 255, 255))
        new_img.paste(img, (0, 0), img)
        draw = ImageDraw.Draw(new_img)

        # Fond du bandeau (bleu très clair)
        draw.rectangle([(0, H), (W, H + BAND_H)], fill=(240, 244, 255, 255))
        # Ligne de séparation
        draw.line([(0, H), (W, H)], fill=(180, 195, 230), width=2)
        # Séparateurs verticaux entre colonnes
        for x in (W // 3, 2 * W // 3):
            draw.line([(x, H + 6), (x, H + BAND_H - 6)], fill=(200, 210, 235), width=1)

        col_w = W // 3
        PAD   = max(14, W // 80)
        TOP_L = H + max(5, BAND_H // 9)
        TOP_V = H + max(22, BAND_H // 3)

        def _fmt(v):
            if v is None:
                return "N/D"
            return f"{v:,.0f} FCFA".replace(",", "\u00a0")

        # Colonne 1 — Plus haut (vert)
        C_HAUT = (15, 157, 88)
        draw.text((PAD, TOP_L), "\u25b2 Plus haut 100j", font=font_lbl, fill=C_HAUT)
        draw.text((PAD, TOP_V), _fmt(plus_haut),          font=font_val, fill=C_HAUT)

        # Colonne 2 — Cours actuel (bleu marine)
        C_COURS = (26, 35, 126)
        ca_label = "Cours actuel"
        perf_str = f"  ({perf_100j})" if perf_100j else ""
        draw.text((col_w + PAD, TOP_L), ca_label,                      font=font_lbl, fill=C_COURS)
        draw.text((col_w + PAD, TOP_V), _fmt(cours_actuel) + perf_str, font=font_val, fill=C_COURS)

        # Colonne 3 — Plus bas (rouge)
        C_BAS = (217, 48, 37)
        draw.text((2 * col_w + PAD, TOP_L), "\u25bc Plus bas 100j", font=font_lbl, fill=C_BAS)
        draw.text((2 * col_w + PAD, TOP_V), _fmt(plus_bas),          font=font_val, fill=C_BAS)

        # ── Sérialiser ───────────────────────────────────────────────────────
        buf = _io.BytesIO()
        new_img.save(buf, format="PNG", optimize=False)
        result = buf.getvalue()
        print(f"  [Fiches/Chart] {s.get('ticker','?')} : bandeau annoté ajouté "
              f"({W}×{H+BAND_H}px, {len(result):,} bytes)")
        return result

    except Exception as exc:
        print(f"  [Fiches/Chart] {s.get('ticker','?')} : annotation échouée ({exc}) "
              "— PNG source utilisé sans modification")
        return png_bytes


def _extract_images_from_docx(doc_bytes: bytes,
                               known_tickers: list | None = None) -> dict:
    """
    Extrait les graphiques du document Word source et les associe à leur ticker.

    Stratégie en deux passes :

    Passe 1 — collecte séquentielle
        On parcourt le body dans l'ordre et on construit deux listes parallèles :
          - ``elements_info`` : (index, type, texte, blobs)
          - En identifiant les tickers dans le texte et les images dans les éléments.

    Passe 2 — association par recherche arrière (remontée)
        Pour chaque image trouvée (index i), on remonte les éléments précédents
        jusqu'à trouver le premier ticker BRVM valide (known_tickers en priorité).
        On s'arrête au bout de _IMG_MAX_DIST éléments remontés.
        Cette approche est robuste même si le ticker est dans le titre de section,
        30-50 éléments avant l'image.

    Filtre qualité : on ignore les blobs < 5 KB (icônes, logos).
    En cas de doublon, on garde l'image la plus grande pour chaque ticker.

    Paramètres
    ----------
    doc_bytes      : bytes du .docx source
    known_tickers  : liste des tickers LLM (filtre prioritaire)

    Retourne
    --------
    dict[ticker_str -> png_bytes]
    """
    try:
        doc = Document(io.BytesIO(doc_bytes))
    except Exception as exc:
        print(f"  [Fiches/ImgExtract] Impossible d'ouvrir le document source — {exc}")
        return {}

    kt_set = set(t.upper() for t in (known_tickers or []))

    # ── Helpers internes ──────────────────────────────────────────────────────

    def _find_tickers_in_text(text: str) -> list[str]:
        """
        Retourne tous les tickers BRVM valides trouvés dans text.
        Priorité : known_tickers d'abord, puis tokens uppercase non-stopwords.
        """
        found = []
        for m in _TICKER_RE.finditer(text):
            cand = m.group(1)
            if cand in _TICKER_STOPWORDS:
                continue
            # Si on a des known_tickers, ne garder que ceux-là
            if kt_set and cand not in kt_set:
                continue
            if cand not in found:
                found.append(cand)
        return found

    def _blobs_from_element(element) -> list[bytes]:
        """Extrait tous les blobs d'images d'un élément XML (y compris imbriqués)."""
        blobs = []
        for drawing in element.findall(".//" + qn("w:drawing")):
            for blip in drawing.findall(".//" + _BLIP_TAG):
                rId = blip.get(_EMBED_ATTR)
                if not rId:
                    continue
                rel = doc.part.rels.get(rId)
                if rel is None:
                    continue
                try:
                    b = rel.target_part.blob
                    if len(b) >= 5_000:   # ignorer icônes/logos < 5 KB
                        blobs.append(b)
                    else:
                        print(f"  [Fiches/ImgExtract] blob ignoré ({len(b)} bytes — trop petit)")
                except Exception:
                    pass
        return blobs

    # ── Passe 1 : collecte ────────────────────────────────────────────────────
    # Chaque entrée : {"idx": int, "tickers": list[str], "blobs": list[bytes]}
    elements_info = []

    for idx, child in enumerate(doc.element.body.iterchildren()):
        tag = child.tag
        entry = {"idx": idx, "tickers": [], "blobs": []}

        if tag == qn("w:p"):
            texts = [n.text for n in child.iter() if n.text and n.text.strip()]
            para_text = " ".join(texts)
            entry["tickers"] = _find_tickers_in_text(para_text)
            entry["blobs"]   = _blobs_from_element(child)

        elif tag == qn("w:tbl"):
            # Parcourir toutes les cellules pour texte + images
            all_cell_text = []
            for tc in child.findall(".//" + qn("w:tc")):
                cell_texts = [n.text for n in tc.iter() if n.text and n.text.strip()]
                all_cell_text.extend(cell_texts)
                entry["blobs"].extend(_blobs_from_element(tc))
            entry["tickers"] = _find_tickers_in_text(" ".join(all_cell_text))

        elements_info.append(entry)

    # ── Passe 2 : association par remontée ────────────────────────────────────
    images_map: dict[str, bytes] = {}

    for i, entry in enumerate(elements_info):
        if not entry["blobs"]:
            continue

        # Remonter les éléments précédents pour trouver le ticker le plus proche
        associated_ticker: str | None = None
        for back in range(0, min(_IMG_MAX_DIST, i + 1)):
            prev = elements_info[i - back]
            if prev["tickers"]:
                # Prendre le dernier ticker de l'élément précédent
                # (le plus proche sémantiquement du graphique)
                associated_ticker = prev["tickers"][-1]
                break

        if associated_ticker is None:
            print(f"  [Fiches/ImgExtract] image(s) à l'index {i} : aucun ticker "
                  f"trouvé dans les {_IMG_MAX_DIST} éléments précédents — ignorée(s)")
            continue

        for blob in entry["blobs"]:
            existing = images_map.get(associated_ticker)
            if existing is not None and len(blob) <= len(existing):
                continue  # garder le plus grand
            png = _blob_to_png(blob, associated_ticker)
            if png:
                images_map[associated_ticker] = png

    nb = len(images_map)
    if nb:
        print(f"  [Fiches/ImgExtract] {nb} graphique(s) associé(s) : "
              f"{list(images_map.keys())}")
    else:
        print("  [Fiches/ImgExtract] Aucun graphique associé à un ticker.")
    return images_map


def _build_price_chart_png(s: dict, width_in: float = 6.5, height_in: float = 2.6) -> bytes:
    """
    Génère un graphique matplotlib (100 jours) ancré sur les bornes connues :
    cours_debut, cours_fin, plus_haut_100j, plus_bas_100j.
    Trace : courbe bleue reconstruite, droite de tendance pointillée, marqueurs haut/bas.
    Retourne les bytes PNG ; None si données insuffisantes.
    """
    debut = _to_float(s.get("cours_debut"))
    fin   = _to_float(s.get("cours_fin")) or _to_float(s.get("cours"))
    haut  = _to_float(s.get("plus_haut_100j"))
    bas   = _to_float(s.get("plus_bas_100j"))

    # Reconstruire cours_debut depuis perf_100j si absent :
    # cours_debut = cours_fin / (1 + perf/100)
    if debut is None and fin is not None:
        perf_str = str(s.get("perf_100j") or "").replace("%", "").replace("+", "").strip()
        if perf_str:
            try:
                perf_val = float(perf_str.replace(",", "."))
                debut = fin / (1 + perf_val / 100)
                print(f"  [Fiches/Chart] {s.get('ticker')} : cours_debut calculé "
                      f"depuis perf_100j={perf_str}% → {debut:.0f}")
            except (ValueError, ZeroDivisionError):
                pass

    if debut is None or fin is None or haut is None or bas is None:
        missing = [k for k, v in [("cours_debut", debut), ("cours_fin", fin),
                                   ("plus_haut_100j", haut), ("plus_bas_100j", bas)]
                   if v is None]
        print(f"  [Fiches/Chart] {s.get('ticker')} : données manquantes {missing} — "
              "graphique matplotlib impossible")
        return None
    if haut < bas:
        haut, bas = bas, haut

    # ── Reconstruction d'une trajectoire 100 séances ancrée sur les 4 points connus.
    n = 100
    rng = random.Random(hash(str(s.get("ticker") or "")) & 0xFFFFFFFF)
    if abs(haut - debut) >= abs(bas - debut):
        i_high, i_low = 30, 70
    else:
        i_low, i_high = 30, 70

    anchors = sorted([(0, debut), (i_high, haut), (i_low, bas), (n - 1, fin)])
    x_a = [a[0] for a in anchors]
    y_a = [a[1] for a in anchors]

    xs = list(range(n))
    ys = []
    for x in xs:
        for k in range(len(x_a) - 1):
            if x_a[k] <= x <= x_a[k + 1]:
                t = (x - x_a[k]) / max(x_a[k + 1] - x_a[k], 1)
                t_smooth = 0.5 - 0.5 * math.cos(math.pi * t)
                base = y_a[k] + (y_a[k + 1] - y_a[k]) * t_smooth
                break
        amp = (haut - bas) * 0.04
        if x in (0, i_high, i_low, n - 1):
            ys.append(base)
        else:
            noisy = base + rng.uniform(-amp, amp)
            ys.append(min(haut, max(bas, noisy)))

    # ── Tracé
    fig, ax = plt.subplots(figsize=(width_in, height_in), dpi=140)
    ax.plot(xs, ys, color="#1A73E8", linewidth=1.6, zorder=3)
    ax.fill_between(xs, ys, min(ys) - (haut - bas) * 0.05,
                    color="#1A73E8", alpha=0.07, zorder=1)

    ax.plot([0, n - 1], [debut, fin], linestyle="--", color="#888888",
            linewidth=1.2, zorder=2, label="Tendance")

    ax.scatter([i_high], [haut], color="#0F9D58", s=42, zorder=5, edgecolor="white")
    ax.scatter([i_low], [bas], color="#D93025", s=42, zorder=5, edgecolor="white")
    ax.annotate(f"+ haut\n{haut:,.0f}".replace(",", " "),
                xy=(i_high, haut), xytext=(0, 8), textcoords="offset points",
                ha="center", fontsize=7, color="#0F9D58", weight="bold")
    ax.annotate(f"+ bas\n{bas:,.0f}".replace(",", " "),
                xy=(i_low, bas), xytext=(0, -22), textcoords="offset points",
                ha="center", fontsize=7, color="#D93025", weight="bold")

    ax.scatter([0, n - 1], [debut, fin], color="#1A237E", s=22, zorder=4)

    date_d = str(s.get("date_debut_100j") or "J-100")
    date_f = str(s.get("date_fin_100j") or "J")
    ax.set_xticks([0, n // 2, n - 1])
    ax.set_xticklabels([date_d, "—", date_f], fontsize=7, color="#555555")
    ax.yaxis.set_major_formatter(FuncFormatter(
        lambda v, _: f"{v:,.0f}".replace(",", " ")))
    ax.tick_params(axis="y", labelsize=7, colors="#555555")
    for side in ("top", "right"):
        ax.spines[side].set_visible(False)
    for side in ("left", "bottom"):
        ax.spines[side].set_color("#CCCCCC")
    ax.grid(True, axis="y", linestyle=":", color="#DDDDDD", linewidth=0.6, zorder=0)
    ax.set_ylim(bas - (haut - bas) * 0.10, haut + (haut - bas) * 0.18)

    perf = s.get("perf_100j")
    title_perf = f"   ({perf})" if perf else ""
    ax.set_title(
        f"Évolution du cours sur 100 jours{title_perf}",
        fontsize=9, color="#1A237E", weight="bold", loc="left", pad=6,
    )

    fig.tight_layout(pad=0.6)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    png_bytes = buf.getvalue()
    print(f"  [Fiches/Chart] {s.get('ticker')} : PNG matplotlib {len(png_bytes)} bytes")
    return png_bytes


def _fund_val(s: dict, key: str) -> str:
    v = s.get(key)
    if v is None:
        return "N/D"
    sv = str(v).strip()
    if sv == "" or sv.lower() in ("null", "none", "—"):
        return "N/D"
    return sv


def _build_fundamentals_table(doc, s: dict):
    """
    Tableau enrichi des indicateurs fondamentaux. Inclut CA et résultat net
    sur 3 exercices (N, N-1, N-2), ratios de rentabilité (marge, ROE, ROA, PER),
    capitalisation et structure (dividende, rendement, dette nette / EBITDA),
    et croissances (1 an / 3 ans). 'N/D' quand la donnée est absente du rapport.
    """
    _sub_heading(doc, "Indicateurs fondamentaux (extraits du rapport)")

    per_actuel = _fund_val(s, "per")
    per_secto = _fund_val(s, "per_sectoriel")
    per_val = per_actuel if per_secto in ("N/D", "") else f"{per_actuel} (secteur : {per_secto})"

    div_actuel = _fund_val(s, "dividende")
    rend = _fund_val(s, "rendement_dividende")
    div_val = div_actuel if rend in ("N/D", "") else f"{div_actuel} (rendement : {rend})"

    dette_nette = _fund_val(s, "dette_nette")
    dn_ebitda = _fund_val(s, "dette_nette_ebitda")
    dette_val = dette_nette if dn_ebitda in ("N/D", "") else f"{dette_nette} | DN/EBITDA : {dn_ebitda}"

    rows = [
        ("Chiffre d'affaires (N)",      _fund_val(s, "ca"),                       _fund_val(s, "ca_date")),
        ("CA — exercice N-1",           _fund_val(s, "ca_n_1"),                   "N-1"),
        ("CA — exercice N-2",           _fund_val(s, "ca_n_2"),                   "N-2"),
        ("Résultat net (N)",            _fund_val(s, "resultat_net"),             _fund_val(s, "rn_date")),
        ("Résultat net — N-1",          _fund_val(s, "resultat_n_1"),             "N-1"),
        ("Résultat net — N-2",          _fund_val(s, "resultat_n_2"),             "N-2"),
        ("Marge nette",                 _fund_val(s, "marge_nette"),              _fund_val(s, "mn_date")),
        ("ROE",                         _fund_val(s, "roe"),                      _fund_val(s, "roe_date")),
        ("ROA",                         _fund_val(s, "roa"),                      _fund_val(s, "roa_date")),
        ("PER (actuel / sectoriel)",    per_val,                                  _fund_val(s, "per_date")),
        ("Capitalisation boursière",    _fund_val(s, "capitalisation_boursiere"), _fund_val(s, "capi_date")),
        ("Dividende & rendement",       div_val,                                  _fund_val(s, "div_date")),
        ("Dette nette / EBITDA",        dette_val,                                _fund_val(s, "date_donnees_financieres")),
        ("Croissance CA — 1 an",        _fund_val(s, "croissance_ca_1an"),        _fund_val(s, "croissance_ca_date")),
        ("Croissance CA — 3 ans",       _fund_val(s, "croissance_ca_3ans"),       _fund_val(s, "croissance_ca_date")),
    ]

    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Indicateur", "Valeur", "Date / période"]):
        _cw(tbl.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")

    for i, (label, val, dt) in enumerate(rows, start=1):
        bg_val = "F5F5F5" if val == "N/D" else "FFFFFF"
        bg_dt  = "F5F5F5" if dt  == "N/D" else "FFFFFF"
        _cw(tbl.rows[i].cells[0], label, bold=True, size=8, bg="EBF0FA")
        _cw(tbl.rows[i].cells[1], val, size=8, bg=bg_val)
        _cw(tbl.rows[i].cells[2], dt,  size=8, bg=bg_dt)

    doc.add_paragraph()
    _build_financial_commentary(doc, s)


# ── Commentaire narratif d'analyse financière ────────────────────────────────

def _pct_to_float(v):
    """'+5,2%' → 5.2 ; '5,8%' → 5.8 ; None / non-numérique → None."""
    if v is None:
        return None
    s = str(v).replace("%", "").replace(",", ".").replace("+", "").strip()
    if not s or s.lower() in ("null", "none", "—", "n/d"):
        return None
    try:
        return float(s)
    except ValueError:
        return None


def _build_financial_commentary(doc, s: dict):
    """
    Paragraphe narratif (5-8 lignes) de synthèse financière. Combine :
    - santé financière globale (rentabilité, marge, dette)
    - comparaison sectorielle (PER vs PER sectoriel)
    - points forts et risques identifiés
    - conclusion sur l'attractivité financière
    """
    _sub_heading(doc, "Commentaire — Analyse financière")

    ticker = _s(s, "ticker", "—")
    nom = _s(s, "nom", ticker)
    secteur = _s(s, "secteur", "—")

    marge = _pct_to_float(s.get("marge_nette"))
    roe = _pct_to_float(s.get("roe"))
    roa = _pct_to_float(s.get("roa"))
    croiss_1y = _pct_to_float(s.get("croissance_ca_1an")) or _pct_to_float(s.get("croissance_ca"))
    croiss_3y = _pct_to_float(s.get("croissance_ca_3ans"))
    rendement = _pct_to_float(s.get("rendement_dividende"))

    per_f = _pct_to_float(s.get("per"))
    per_sec_f = _pct_to_float(s.get("per_sectoriel"))

    # 1) Santé financière globale
    if roe is not None and marge is not None:
        if roe >= 12 and marge >= 8:
            sante = (
                f"{nom} affiche une rentabilité solide (ROE de {roe:.1f}%, "
                f"marge nette de {marge:.1f}%), signe d'une exploitation efficace "
                "et d'un retour sur fonds propres conforme aux standards du secteur."
            )
        elif roe < 5 or marge < 2:
            sante = (
                f"La rentabilité de {nom} apparaît contrainte (ROE {roe:.1f}%, "
                f"marge nette {marge:.1f}%), reflet de pressions sur les coûts ou "
                "d'une structure de capital peu optimisée."
            )
        else:
            sante = (
                f"{nom} présente une rentabilité moyenne (ROE {roe:.1f}%, "
                f"marge nette {marge:.1f}%), niveau qui laisse une marge de "
                "progression sans signal d'alerte immédiat."
            )
    else:
        sante = (
            f"{nom}, cotée dans le secteur {secteur}, présente un profil financier "
            "dont les ratios complets ne sont pas tous disponibles dans le rapport "
            "source — l'appréciation s'appuie sur les éléments partiels recensés."
        )

    # 2) Croissance & dynamique commerciale
    if croiss_3y is not None:
        if croiss_3y >= 10:
            croissance = (
                f"La trajectoire commerciale est porteuse (croissance CA "
                f"{croiss_3y:+.1f}% sur 3 ans), traduisant une dynamique "
                "structurelle de parts de marché ou de pricing power."
            )
        elif croiss_3y <= 0:
            croissance = (
                f"Le chiffre d'affaires est en contraction sur la période "
                f"({croiss_3y:+.1f}% sur 3 ans), signal d'une demande affaiblie ou "
                "d'un repositionnement concurrentiel à surveiller."
            )
        else:
            croissance = (
                f"La croissance commerciale demeure modérée ({croiss_3y:+.1f}% "
                "sur 3 ans), conforme à un secteur en phase de maturité."
            )
    elif croiss_1y is not None:
        croissance = (
            f"Sur le dernier exercice, le CA évolue de {croiss_1y:+.1f}% — "
            "à confirmer sur un horizon pluriannuel."
        )
    else:
        croissance = (
            "La dynamique de croissance n'est pas chiffrée dans le rapport source, "
            "ce qui limite la projection des flux futurs."
        )

    # 3) Comparaison sectorielle (PER)
    if per_f is not None and per_sec_f is not None:
        ecart = per_f - per_sec_f
        if ecart <= -1.5:
            comparaison = (
                f"Sur le plan de la valorisation, le PER de {per_f:.1f} ressort "
                f"en-dessous de la moyenne sectorielle ({per_sec_f:.1f}) — "
                "configuration de décote relative qui peut intéresser les "
                "investisseurs value, sous réserve d'un catalyseur de revalorisation."
            )
        elif ecart >= 1.5:
            comparaison = (
                f"La valorisation apparaît plus exigeante que la moyenne du secteur "
                f"(PER {per_f:.1f} vs {per_sec_f:.1f}) — la prime ne se justifie "
                "qu'en présence d'une croissance ou d'une rentabilité supérieures."
            )
        else:
            comparaison = (
                f"La valorisation est alignée sur le secteur (PER {per_f:.1f} vs "
                f"{per_sec_f:.1f}), ce qui n'introduit ni décote ni prime."
            )
    elif per_f is not None:
        comparaison = (
            f"Le PER actuel ressort à {per_f:.1f} ; à défaut de référence sectorielle "
            "explicite, la comparaison de valorisation reste indicative."
        )
    else:
        comparaison = (
            "La valorisation (PER) n'étant pas renseignée, la comparaison sectorielle "
            "se fonde sur les ratios de rentabilité disponibles."
        )

    # 4) Forces & risques
    forces = _sl(s, "forces_financieres")
    faiblesses = _sl(s, "faiblesses_financieres")
    forts_risques_parts = []
    if forces:
        forts_risques_parts.append("Points forts : " + " ; ".join(f.strip().rstrip(".") for f in forces[:3]) + ".")
    if faiblesses:
        forts_risques_parts.append("Risques : " + " ; ".join(f.strip().rstrip(".") for f in faiblesses[:3]) + ".")
    if not forts_risques_parts:
        if rendement is not None and rendement >= 4:
            forts_risques_parts.append(
                f"Point fort : rendement du dividende attractif ({rendement:.1f}%)."
            )
        if roa is not None and roa < 1:
            forts_risques_parts.append(
                f"Risque : ROA limité ({roa:.1f}%) signalant une efficience modeste des actifs."
            )
    forts_risques = " ".join(forts_risques_parts) if forts_risques_parts else (
        "L'analyse qualitative ne fait pas ressortir de points forts ou risques majeurs "
        "au-delà des ratios déjà commentés."
    )

    # 5) Conclusion sur l'attractivité financière
    score = _score_f(s)
    reco = _s(s, "reco", "").upper()
    if score >= 70 or "ACHAT" in reco:
        conclusion = (
            "Au global, l'attractivité financière est jugée favorable : la combinaison "
            "des fondamentaux soutient une position constructive sur le titre."
        )
    elif score <= 39 or "VENTE" in reco:
        conclusion = (
            "L'attractivité financière reste fragilisée : les ratios actuels ne "
            "soutiennent pas un repositionnement offensif sans amélioration tangible."
        )
    else:
        conclusion = (
            "L'attractivité financière est intermédiaire : un positionnement sélectif "
            "à proportion mesurée est cohérent avec le profil fondamental observé."
        )

    full = " ".join([sante, croissance, comparaison, forts_risques, conclusion])
    _narrative(doc, full)


# ── Helpers métier ────────────────────────────────────────────────────────────

def _signal_emoji(signal: str) -> str:
    s = str(signal or "").lower()
    if any(w in s for w in ("haussier", "positif", "achat", "fort", "bon", "élevé", "eleve")):
        return "🟢"
    if any(w in s for w in ("baissier", "négatif", "negatif", "vente", "faible")):
        return "🔴"
    return "🟡"


def _signal_bg(signal: str) -> str:
    s = str(signal or "").lower()
    if any(w in s for w in ("haussier", "positif", "achat", "élevé")):
        return "C6EFCE"
    if any(w in s for w in ("baissier", "négatif", "vente")):
        return "FFC7CE"
    return "FFEB9C"


def _score_label_color(score) -> tuple:
    try:
        s = float(score or 0)
    except (ValueError, TypeError):
        s = 0.0
    if s >= 75:
        return "Excellent", "0F9D58"
    if s >= 60:
        return "Bon", "1A73E8"
    if s >= 40:
        return "Moyen", "E37400"
    return "Faible", "D93025"


def _reco_bg(reco: str) -> str:
    r = str(reco).upper()
    if "ACHAT" in r:
        return "C6EFCE"
    if "VENTE" in r:
        return "FFC7CE"
    return "FFEB9C"


def _reco_fg(reco: str) -> str:
    r = str(reco).upper()
    if "ACHAT" in r:
        return "0F9D58"
    if "VENTE" in r:
        return "D93025"
    return "E37400"


def _risque_bg(risque: str) -> str:
    r = str(risque or "").lower()
    if "faible" in r:
        return "C6EFCE"
    if "élevé" in r or "eleve" in r:
        return "FFC7CE"
    return "FFEB9C"


def _var_color(var: str) -> str:
    v = str(var or "").strip()
    if v.startswith("+"):
        return "EBF7EE"
    if v.startswith("-"):
        return "FDEEEE"
    return "FFFFFF"


def _s(data, key, default="") -> str:
    v = data.get(key)
    return str(v) if v is not None else default


def _sl(data, key) -> list:
    v = data.get(key)
    return v if isinstance(v, list) else []


def _score_f(data) -> float:
    try:
        return float(data.get("score") or 0)
    except (ValueError, TypeError):
        return 0.0


# ── Extraction texte source ───────────────────────────────────────────────────

def _extract_text(doc_bytes: bytes) -> str:
    """
    Parcourt le body en ordre document pour préserver l'association paragraphes/tableaux.
    Les cellules au format 'clé\\nvaleur' (comme dans les tableaux d'indicateurs
    financiers du rapport source) sont aplaties en 'clé: valeur'.
    """
    doc = Document(io.BytesIO(doc_bytes))
    para_by_id = {id(p._element): p for p in doc.paragraphs}
    table_by_id = {id(t._element): t for t in doc.tables}

    parts = []
    table_idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_by_id.get(id(child))
            if p is None:
                continue
            txt = p.text.strip()
            if txt:
                parts.append(txt)
        elif child.tag == qn("w:tbl"):
            t = table_by_id.get(id(child))
            if t is None:
                continue
            table_idx += 1
            kv_lines = []
            for row in t.rows:
                for cell in row.cells:
                    raw = cell.text.strip()
                    if not raw:
                        continue
                    if "\n" in raw:
                        k, v = raw.split("\n", 1)
                        k, v = k.strip(), v.strip()
                        if v and v.lower() != "none":
                            kv_lines.append(f"  {k}: {v}")
                    else:
                        kv_lines.append(f"  {raw}")
            if kv_lines:
                parts.append(f"[Tableau {table_idx}]")
                parts.extend(kv_lines)
    return "\n".join(parts)


# ── Contexte multi-documents ─────────────────────────────────────────────────

def _build_context(docs_bytes: list, freq: str) -> str:
    if len(docs_bytes) == 1:
        return _extract_text(docs_bytes[0])

    max_older = {"HEBDO": 6, "MENSUEL": 9, "TRIM": 12, "ANNUEL": 14}.get(freq, 5)
    chars_older = {"HEBDO": 3000, "MENSUEL": 2000, "TRIM": 1500, "ANNUEL": 1000}.get(freq, 2000)

    recent = _extract_text(docs_bytes[0])[:25000]
    older_parts = []
    for i, db in enumerate(docs_bytes[1:max_older + 1], 1):
        excerpt = _extract_text(db)[:chars_older]
        older_parts.append(f"--- Document J-{i} ---\n{excerpt}")

    return (
        "=== DOCUMENT LE PLUS RÉCENT (J) ===\n"
        + recent
        + ("\n\n=== HISTORIQUE (extraits) ===\n" + "\n\n".join(older_parts) if older_parts else "")
    )


# ═══════════════════════════════════════════════════════════════════════════════
# FONCTIONS DE CONSTRUCTION DE LA FICHE
# ═══════════════════════════════════════════════════════════════════════════════

def build_header(doc, s: dict, date_str: str):
    """
    En-tête complet de la fiche :
    Ligne 1 — Ticker | Nom | Secteur | Date
    Ligne 2 — Score/Label | Recommandation | Indicateurs (🟢🟡🔴)
    Ligne 3 — Décision | Confiance | Tendance | Résumé
    """
    ticker = _s(s, "ticker", "???")
    nom = _s(s, "nom", ticker)
    secteur = _s(s, "secteur", "—")
    score = s.get("score")
    score_str = f"{score:.0f}" if score is not None else "—"
    score_label, score_color = _score_label_color(score)
    reco = _s(s, "reco", "NEUTRE")
    decision = _s(s, "decision", "—")
    confiance = _s(s, "confiance", "—")
    tendance = _s(s, "tendance_100j", "—")
    resume = _s(s, "resume_rapport", "—")

    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"

    # ── Ligne 1 : identité
    r0 = tbl.rows[0]
    _cw(r0.cells[0], ticker, bold=True, size=15, color="FFFFFF", bg="1A237E")
    _cw(r0.cells[1], nom[:48], size=9, color="E8EAF6", bg="1A237E")
    _cw(r0.cells[2], secteur, size=8, color="C5CAE9", bg="283593")
    _cw(r0.cells[3], f"Rapport du {date_str}", size=8, color="C5CAE9", bg="283593",
        align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Ligne 2 : score + reco + indicateurs
    r1 = tbl.rows[1]
    _cw(r1.cells[0], f"Score : {score_str}/100  —  {score_label}",
        bold=True, size=9, color=score_color, bg="E8EAF6")
    _cw(r1.cells[1], reco, bold=True, size=11,
        color=_reco_fg(reco), bg=_reco_bg(reco))
    ind = "   ".join([
        f"{_signal_emoji(_s(s, 'mm'))} MM",
        f"{_signal_emoji(_s(s, 'boll'))} Boll",
        f"{_signal_emoji(_s(s, 'macd'))} MACD",
        f"{_signal_emoji(_s(s, 'rsi'))} RSI",
        f"{_signal_emoji(_s(s, 'stoch'))} Stoch",
    ])
    r1.cells[2].merge(r1.cells[3])
    _cw(r1.cells[2], ind, size=9, bg="E8EAF6")

    # ── Ligne 3 : décision + confiance + tendance + résumé
    r2 = tbl.rows[2]
    _cw(r2.cells[0], f"Décision : {decision}", bold=True, size=8, bg="F0F4FF")
    _cw(r2.cells[1], f"Confiance : {confiance}", size=8, bg="F5F5F5")
    _cw(r2.cells[2], f"Tendance 100j : {tendance}", size=8, bg="F5F5F5")
    _cw(r2.cells[3], f"Résumé : {resume}", size=8, bg="F5F5F5")


def _fmt_int(v) -> str:
    """Format un nombre en milliers séparés par espace ; renvoie '—' si None."""
    if v is None:
        return "—"
    try:
        f = float(v)
        if f != f:  # NaN
            return "—"
        return f"{f:,.0f}".replace(",", " ")
    except (ValueError, TypeError):
        return str(v).strip() or "—"


def _validate_var_1j(val):
    """Filtre les variations journalières aberrantes (>±10%, signe de
    confusion avec une perf multi-jours type perf_100j)."""
    if val is None:
        return "—"
    try:
        pct = float(
            str(val).replace("%", "").replace("+", "").replace(",", ".").strip()
        )
        if abs(pct) > 10:
            return "—"
        return val
    except (ValueError, TypeError):
        return "—"


def _fmt_amount(val):
    """Formate un montant avec séparateur de milliers (espace).
    Si val est déjà une string portant une unité (Mds, M, FCFA, %), la
    conserve telle quelle ; si val est un nombre brut ou une chaîne purement
    numérique, formate avec espaces comme séparateur de milliers."""
    if val is None or val == "N/D" or val == "—":
        return "—"
    if isinstance(val, (int, float)):
        try:
            f = float(val)
            if f != f:
                return "—"
            return f"{f:,.0f}".replace(",", " ")
        except (ValueError, TypeError):
            return str(val)
    s = str(val).strip()
    if not s or s.lower() in ("null", "none"):
        return "—"
    if any(u in s for u in ("Md", "milliard", "Milliard", "million", "Million",
                            "FCFA", "%", " M ", " M.")):
        return s
    try:
        f = float(s.replace(" ", "").replace(",", "."))
        return f"{f:,.0f}".replace(",", " ")
    except ValueError:
        return s


def build_market_table(doc, s: dict):
    """
    Tableau des métriques de marché en 6×2 (12 indicateurs).
    Inclut capitalisation, volume moyen 30j, nombre d'actions et PER
    en plus des métriques de risque/stabilité.
    """
    _section_heading(doc, "MÉTRIQUES DE MARCHÉ")

    cours = _fmt_amount(s.get("cours"))
    var_1j = _validate_var_1j(s.get("var_1j"))
    volatilite = _s(s, "volatilite") or "—"
    beta = _s(s, "beta") or "—"
    liquidite = _s(s, "liquidite") or "—"
    risque = _s(s, "risque") or "—"
    divergence = _s(s, "divergence") or "aucune"
    stabilite = _s(s, "stabilite") or "—"
    capi = _fmt_amount(s.get("capitalisation_boursiere"))
    vol_30j = _fmt_int(s.get("volume_moyen_30j"))
    nb_act = _fmt_int(s.get("nb_actions"))
    per = _s(s, "per") or "—"

    risque_bg = _risque_bg(risque)
    var_bg = _var_color(var_1j)
    stab_bg = "C6EFCE" if "bonne" in str(stabilite).lower() else (
        "FFC7CE" if "fragile" in str(stabilite).lower() else "FFEB9C"
    )
    div_bg = "FFEB9C" if divergence.lower() not in ("aucune", "—", "") else "FFFFFF"

    pairs = [
        ("Cours actuel (FCFA)",       cours,      "F0F4FF"),
        ("Variation 1 journée",       var_1j,     var_bg),
        ("Capitalisation boursière",  capi,       "FFFFFF"),
        ("PER (Cours / BPA)",         per,        "FFFFFF"),
        ("Volume moyen 30j",          vol_30j,    "FFFFFF"),
        ("Nb actions en circulation", nb_act,     "FFFFFF"),
        ("Volatilité",                volatilite, "FFFFFF"),
        ("Bêta",                      beta,       "FFFFFF"),
        ("Liquidité",                 liquidite,  "FFFFFF"),
        ("Niveau de risque",          risque,     risque_bg),
        ("Divergence tech/fond",      divergence, div_bg),
        ("Stabilité",                 stabilite,  stab_bg),
    ]

    tbl = doc.add_table(rows=6, cols=4)
    tbl.style = "Table Grid"
    for i in range(6):
        ll, lv, lbg = pairs[i]
        rl, rv, rbg = pairs[i + 6]
        _cw(tbl.rows[i].cells[0], ll, bold=True, size=8, bg="EBF0FA")
        _cw(tbl.rows[i].cells[1], lv, size=8, bg=lbg)
        _cw(tbl.rows[i].cells[2], rl, bold=True, size=8, bg="EBF0FA")
        _cw(tbl.rows[i].cells[3], rv, size=8, bg=rbg)


def build_chart_comment(doc, s: dict, source_png: bytes | None = None):
    """
    Insère le graphique extrait du rapport Word source + commentaire analytique.
    Seul le graphique original du rapport est utilisé (cours réels + prédictions).
    Si l'image source est absente, un message d'absence est affiché.

    Paramètres
    ----------
    source_png : bytes PNG extrait par _extract_images_from_docx(), ou None.
    """
    ticker = _s(s, "ticker", "?")
    png_bytes: bytes | None = None

    # Graphique source Word — seule source acceptée
    if source_png and len(source_png) > 1000:
        # Ajouter le bandeau Plus haut / Cours actuel / Plus bas sous le graphique
        png_bytes = _annotate_source_chart(source_png, s)
    else:
        print(f"  [Fiches/Chart] {ticker} : graphique source absent "
              f"(source_png={'None' if source_png is None else str(len(source_png))+' bytes'})")

    # ── Insertion de l'image dans le document ─────────────────────────────────
    if png_bytes and len(png_bytes) > 1000:
        p_img = doc.add_paragraph()
        p_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(2)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(png_bytes)
            tmp_path = tmp.name
        try:
            p_img.add_run().add_picture(tmp_path, width=Cm(16))
        except Exception as exc:
            print(f"  [Fiches/Chart] {ticker} : add_picture() échec — {exc}")
            # Paragraphe d'erreur si l'insertion échoue
            p_img.clear()
            r_err = p_img.add_run(
                f"[Graphique disponible mais non inséré — erreur technique : {exc}]"
            )
            r_err.italic = True
            r_err.font.size = Pt(9)
            r_err.font.color.rgb = _rgb("C0392B")
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

        # Légende discrète sous le graphique
        p_leg = doc.add_paragraph()
        p_leg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_leg.paragraph_format.space_before = Pt(0)
        p_leg.paragraph_format.space_after = Pt(4)
        r_leg = p_leg.add_run(
            f"Cours réels + prédictions J+1→J+10  ·  {ticker}"
        )
        r_leg.italic = True
        r_leg.font.size = Pt(7)
        r_leg.font.color.rgb = _rgb("AAAAAA")

    else:
        if png_bytes is not None:
            print(f"  [Fiches/Chart] {ticker} : PNG trop petit "
                  f"({len(png_bytes)} bytes) — ignoré")
        p_ph = doc.add_paragraph()
        p_ph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ph.paragraph_format.space_before = Pt(8)
        p_ph.paragraph_format.space_after = Pt(4)
        r_ph = p_ph.add_run(
            "[Données de cours 100j non disponibles pour ce titre — graphique non tracé]"
        )
        r_ph.italic = True
        r_ph.font.size = Pt(9)
        r_ph.font.color.rgb = _rgb("888888")

    # ── Section commentaire ───────────────────────────────────────────────────
    _section_heading(doc, "COMMENTAIRE DE LA COURBE — ÉVOLUTION 100 JOURS")

    # ── Données clés ──────────────────────────────────────────────────────────
    ticker      = _s(s, "ticker", "?")
    cours       = _to_float(s.get("cours") or s.get("cours_fin"))
    var_1j_raw  = _validate_var_1j(s.get("var_1j"))
    perf_100j   = _s(s, "perf_100j", "—")
    plus_haut   = _to_float(s.get("plus_haut_100j"))
    plus_bas    = _to_float(s.get("plus_bas_100j"))
    tendance    = _s(s, "tendance_100j", "neutre").lower()
    volatilite  = _s(s, "volatilite", "modérée").lower()
    reco        = _s(s, "reco", "NEUTRE").upper()
    score       = _score_f(s)

    # ── Calcul distance cours actuel / plus haut et plus bas ─────────────────
    dist_haut_str = ""
    dist_bas_str  = ""
    if cours and plus_haut and plus_haut > 0:
        d = (plus_haut - cours) / plus_haut * 100
        dist_haut_str = f"{d:.1f}% sous le plus haut"
    if cours and plus_bas and plus_bas > 0:
        d = (cours - plus_bas) / plus_bas * 100
        dist_bas_str  = f"{d:.1f}% au-dessus du plus bas"

    # ── Signal dominant ───────────────────────────────────────────────────────
    signals = [s.get(k) or "" for k in ("mm", "boll", "macd", "rsi", "stoch")]
    nb_pos = sum(1 for sg in signals if any(
        w in str(sg).lower() for w in ("haussier", "positif", "achat")))
    nb_neg = sum(1 for sg in signals if any(
        w in str(sg).lower() for w in ("baissier", "négatif", "vente")))

    if nb_pos >= 3:
        signal_label = "convergence haussière"
        signal_fg    = "0F9D58"
        signal_bg    = "EBF7EE"
    elif nb_neg >= 3:
        signal_label = "convergence baissière"
        signal_fg    = "D93025"
        signal_bg    = "FDEEEE"
    else:
        signal_label = "signaux mixtes"
        signal_fg    = "E37400"
        signal_bg    = "FFF8E6"

    # ── Phrase de tendance (1 ligne max) ──────────────────────────────────────
    if "haussier" in tendance or "hausse" in tendance:
        tendance_phrase = "Structure haussière — creux ascendants confirmés."
    elif "baissier" in tendance or "baisse" in tendance:
        tendance_phrase = "Structure baissière — rebonds sans suivi vendus."
    else:
        tendance_phrase = "Range horizontal — rupture directionnelle à surveiller."

    # ── Tableau synthèse compact ──────────────────────────────────────────────
    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    headers = ["Perf. 100j", "Var. séance", "Distance plus haut", "Distance plus bas", "Volatilité"]
    values  = [
        perf_100j,
        var_1j_raw if var_1j_raw != "—" else "—",
        dist_haut_str or "—",
        dist_bas_str  or "—",
        volatilite.capitalize(),
    ]
    for i, (h, v) in enumerate(zip(headers, values)):
        _cw(tbl.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")
        # Colorer la perf selon signe
        bg = "FFFFFF"
        if i == 0:
            if perf_100j.startswith("+"):  bg = "EBF7EE"
            elif perf_100j.startswith("-"): bg = "FDEEEE"
        _cw(tbl.rows[1].cells[i], v, size=9, bg=bg)
    doc.add_paragraph()

    # ── Ligne de synthèse (1-2 phrases max) ──────────────────────────────────
    cours_str = f"{cours:,.0f}".replace(",", " ") if cours else "—"
    ph_str    = f"{plus_haut:,.0f}".replace(",", " ") if plus_haut else "—"
    pb_str    = f"{plus_bas:,.0f}".replace(",", " ") if plus_bas else "—"

    synthese = (
        f"{ticker} cote à {cours_str} FCFA (perf. 100j : {perf_100j}). "
        f"Plus haut : {ph_str} — Plus bas : {pb_str}. "
        f"{tendance_phrase} "
        f"Indicateurs techniques : {signal_label} ({nb_pos}/5 positifs, {nb_neg}/5 négatifs)."
    )
    p_synth = doc.add_paragraph()
    p_synth.paragraph_format.space_before = Pt(0)
    p_synth.paragraph_format.space_after = Pt(4)
    p_synth.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_synth = p_synth.add_run(synthese)
    r_synth.font.size = Pt(9)

    # ── PARTIE 1 : Analyse du cours (statistiques 100 jours) ────────────────
    # puis PARTIE 0 : Indicateurs de valorisation boursière
    _src_doc = s.get("_source_doc_ref")
    if _src_doc is not None:
        parties_chart = _extract_parties(_src_doc, ticker)

        # PARTIE 1 en premier : analyse statistique du cours sur 100 jours (sans titre)
        if parties_chart.get('p1'):
            p_p1 = doc.add_paragraph()
            p_p1.paragraph_format.space_before = Pt(1)
            p_p1.paragraph_format.space_after  = Pt(4)
            p_p1.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_p1 = p_p1.add_run(parties_chart['p1'])
            r_p1.font.size = Pt(9)

        # PARTIE 0 après : valorisation boursière (sans titre)
        if parties_chart.get('p0'):
            p_p0 = doc.add_paragraph()
            p_p0.paragraph_format.space_before = Pt(1)
            p_p0.paragraph_format.space_after  = Pt(4)
            p_p0.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_p0 = p_p0.add_run(parties_chart['p0'])
            r_p0.font.size = Pt(9)

    # ── Bloc signal coloré (1 ligne) ──────────────────────────────────────────
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.style = "Table Grid"
    cell2 = tbl2.rows[0].cells[0]
    _cell_bg(cell2, signal_bg)
    tcPr = cell2._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top","bottom","left","right"):
        em = OxmlElement(f"w:{side}"); em.set(qn("w:w"),"80"); em.set(qn("w:type"),"dxa"); tcMar.append(em)
    tcPr.append(tcMar)
    p2 = cell2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(f"Signal dominant : {signal_label.upper()}  |  Recommandation : {reco}  |  Score : {score:.0f}/100")
    r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = _rgb(signal_fg)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def _fmt_num(v, decimals: int = 1) -> str:
    """Format un nombre avec N décimales, espaces de milliers. '—' si vide/None."""
    if v is None:
        return "—"
    try:
        f = float(v)
        if f != f:
            return "—"
        s = f",.{decimals}f"
        return format(f, s).replace(",", " ")
    except (ValueError, TypeError):
        return str(v).strip() or "—"


def _tech_values_str(s: dict, sig_key: str) -> str:
    """Construit la chaîne 'MM20: 2850 | MM50: 2920 | Signal: BAISSIER' pour
    chaque indicateur, en injectant les valeurs numériques extraites."""
    signal = _s(s, sig_key) or "—"
    sig_upper = signal.upper() if signal != "—" else "—"
    if sig_key == "mm":
        return (
            f"MM20 : {_fmt_num(s.get('mm20_valeur'), 0)} | "
            f"MM50 : {_fmt_num(s.get('mm50_valeur'), 0)} | "
            f"Signal : {sig_upper}"
        )
    if sig_key == "boll":
        return (
            f"Bande basse : {_fmt_num(s.get('boll_inf'), 0)} | "
            f"Bande haute : {_fmt_num(s.get('boll_sup'), 0)} | "
            f"Signal : {sig_upper}"
        )
    if sig_key == "macd":
        return (
            f"MACD : {_fmt_num(s.get('macd_valeur'), 2)} | "
            f"Signal line : {_fmt_num(s.get('macd_signal_line'), 2)} | "
            f"Signal : {sig_upper}"
        )
    if sig_key == "rsi":
        return f"RSI : {_fmt_num(s.get('rsi_valeur'), 1)} | Signal : {sig_upper}"
    if sig_key == "stoch":
        return (
            f"%K : {_fmt_num(s.get('stoch_k'), 1)} | "
            f"%D : {_fmt_num(s.get('stoch_d'), 1)} | "
            f"Signal : {sig_upper}"
        )
    return f"Signal : {sig_upper}"


def build_technical_analysis(doc, s: dict):
    """
    Analyse technique structurée :
    - Tableau des 5 indicateurs (signal + appréciation + valeurs numériques + détail)
    - Synthèse globale
    - Évaluation convergence / divergence des signaux
    """
    _section_heading(doc, "ANALYSE TECHNIQUE")

    indicateurs = [
        ("Moyennes Mobiles (MM)", "mm",   "mm_signal",   "mm_detail"),
        ("Bandes de Bollinger",   "boll", "boll_signal", "boll_detail"),
        ("MACD",                  "macd", "macd_signal", "macd_detail"),
        ("RSI",                   "rsi",  "rsi_signal",  "rsi_detail"),
        ("Stochastique",          "stoch","stoch_signal","stoch_detail"),
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Indicateur", "Appréciation", "Valeurs & Signal", "Analyse détaillée"]):
        _cw(tbl.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")

    for label, sig_key, sig_label_key, detail_key in indicateurs:
        signal = _s(s, sig_key)
        sig_label = _s(s, sig_label_key) or (signal.capitalize() if signal else "—")
        detail = _s(s, detail_key) or "—"
        emoji = _signal_emoji(signal)
        values_str = _tech_values_str(s, sig_key)

        row = tbl.add_row()
        _cw(row.cells[0], f"{emoji}  {label}", bold=True, size=8, bg="EBF0FA")
        _cw(row.cells[1], sig_label, size=8, bg=_signal_bg(signal))
        _cw(row.cells[2], values_str, size=7, bg=_signal_bg(signal))
        _cw(row.cells[3], detail[:120] if detail != "—" else "—", size=8)

    doc.add_paragraph()

    synthese = _s(s, "synthese_tech") or _s(s, "analyse_tech")
    if synthese and synthese not in ("", "—"):
        _narrative(doc, synthese, italic=True)

    signals_raw = [_s(s, k) for k in ("mm", "boll", "macd", "rsi", "stoch")]
    pos = sum(1 for sg in signals_raw if any(
        w in str(sg).lower() for w in ("haussier", "positif", "achat", "élevé")))
    neg = sum(1 for sg in signals_raw if any(
        w in str(sg).lower() for w in ("baissier", "négatif", "negatif", "vente")))
    neu = len(signals_raw) - pos - neg

    if pos >= 4:
        conv_label = "CONVERGENCE HAUSSIÈRE"
        conv_text = (
            f"{pos}/5 indicateurs en signal positif. "
            "Les signaux techniques sont fortement alignés — biais acheteur confirmé. "
            "Risque principal : surextension si la progression est trop rapide."
        )
        conv_bg, conv_fg = "C6EFCE", "0F9D58"
    elif neg >= 4:
        conv_label = "CONVERGENCE BAISSIÈRE"
        conv_text = (
            f"{neg}/5 indicateurs en signal négatif. "
            "Les signaux techniques convergent vers un biais vendeur. "
            "Éviter toute entrée avant confirmation d'un signal de retournement."
        )
        conv_bg, conv_fg = "FFC7CE", "D93025"
    elif pos > neg:
        conv_label = "DOMINANTE HAUSSIÈRE"
        conv_text = (
            f"{pos} signaux positifs / {neg} négatifs / {neu} neutres. "
            "Signaux partiellement alignés. "
            "Surveiller la confirmation par le volume avant renforcement."
        )
        conv_bg, conv_fg = "E8F8F0", "155724"
    elif neg > pos:
        conv_label = "DOMINANTE BAISSIÈRE"
        conv_text = (
            f"{neg} signaux négatifs / {pos} positifs / {neu} neutres. "
            "Prudence recommandée. "
            "Ne pas renforcer tant que le bilan des signaux ne s'améliore pas."
        )
        conv_bg, conv_fg = "FFF0E6", "C0392B"
    else:
        conv_label = "SIGNAUX MIXTES"
        conv_text = (
            f"{pos} positifs / {neg} négatifs / {neu} neutres. "
            "Absence de convergence claire — pas de biais directionnel dominant. "
            "Attendre un signal confirmé avant toute décision."
        )
        conv_bg, conv_fg = "FFEB9C", "7D6608"

    _key_bloc(doc, f"SYNTHÈSE TECHNIQUE — {conv_label} :", conv_text, conv_bg, conv_fg)

    # ── PARTIE 2 : Analyse technique détaillée (depuis rapport source) ─────────
    _src_doc2 = s.get("_source_doc_ref")
    if _src_doc2 is not None:
        _ticker2 = _s(s, "ticker", "?")
        parties2 = _extract_parties(_src_doc2, _ticker2)
        if parties2.get('p2'):
            p_p2 = doc.add_paragraph()
            p_p2.paragraph_format.space_before = Pt(2)
            p_p2.paragraph_format.space_after  = Pt(4)
            p_p2.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_p2 = p_p2.add_run(parties2['p2'])
            r_p2.font.size = Pt(9)


def build_fundamental_analysis(doc, s: dict):
    """
    Analyse fondamentale :
    - Profil de la société (secteur, positionnement, score)
    - Analyse qualitative (texte enricher ou fallback)
    - Profil financier (liquidité, stabilité, confiance)
    - Risques identifiés
    - Perspectives
    """
    _section_heading(doc, "ANALYSE FONDAMENTALE")

    ticker = _s(s, "ticker", "—")
    nom = _s(s, "nom", ticker)
    secteur = _s(s, "secteur", "—")
    score_f = _score_f(s)
    score_str = f"{score_f:.0f}" if s.get("score") is not None else "—"
    score_label, _ = _score_label_color(score_f)
    liquidite = _s(s, "liquidite", "—")
    stabilite = _s(s, "stabilite", "—")
    confiance = _s(s, "confiance", "—")
    reco_src = _s(s, "reco_src", "—")

    analyse = _s(s, "analyse_fond_recente")
    if not analyse or analyse in ("", "—", "null", "None"):
        analyse = _s(s, "analyse_fond")

    if analyse and analyse not in ("", "—", "null", "None"):
        _narrative(doc, analyse)
    else:
        _narrative(doc,
                   f"{nom} ({ticker}), cotée dans le secteur {secteur} de la BRVM, "
                   f"présente un profil fondamental {score_label} avec un score composite de {score_str}/100. "
                   f"Liquidité : {liquidite}. Stabilité financière : {stabilite}. "
                   f"Confiance analytique : {confiance}. "
                   "Les données fondamentales détaillées seront intégrées lors du prochain rapport complet.")

    _build_fundamentals_table(doc, s)

    _sub_heading(doc, "Profil financier")
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Liquidité", "Stabilité", "Confiance analytique", "Source reco"]):
        _cw(tbl.rows[0].cells[i], h, bold=True, size=8, bg="EBF0FA")
    vals = [liquidite, stabilite, confiance, reco_src]
    bgs = [
        "C6EFCE" if "haute" in str(liquidite).lower() else ("FFC7CE" if "faible" in str(liquidite).lower() else "FFEB9C"),
        _risque_bg(stabilite.replace("bonne", "faible").replace("fragile", "élevé")),
        "C6EFCE" if "élevée" in str(confiance).lower() else ("FFC7CE" if "faible" in str(confiance).lower() else "FFEB9C"),
        "FFFFFF",
    ]
    for i, (v, b) in enumerate(zip(vals, bgs)):
        _cw(tbl.rows[1].cells[i], v, size=8, bg=b)
    doc.add_paragraph()

    risques = _sl(s, "risques")
    if risques:
        _key_bloc(doc,
                  "⚠  RISQUES IDENTIFIÉS :",
                  "  |  ".join(str(r) for r in risques[:3]),
                  "FFF0E6", "C0392B")

    # ── PARTIE 3 : texte d'analyse fondamentale depuis le rapport source ─────
    # (rapports trimestriels, tendances récentes, recommandation source)
    _src_doc_f = s.get("_source_doc_ref")
    if _src_doc_f is not None:
        _ticker_f = _s(s, "ticker", "?")
        parties_f = _extract_parties(_src_doc_f, _ticker_f)
        if parties_f.get("p3"):
            _sub_heading(doc, "Analyse fondamentale — données récentes (source rapport)")
            p_f3 = doc.add_paragraph()
            p_f3.paragraph_format.space_before = Pt(2)
            p_f3.paragraph_format.space_after  = Pt(4)
            p_f3.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_f3 = p_f3.add_run(parties_f["p3"])
            r_f3.font.size = Pt(9)

    persp = _s(s, "perspectives")
    if persp and persp not in ("", "—"):
        _key_bloc(doc, "PERSPECTIVES :", persp, "E8F8F0", "155724")


def build_financial_analysis(doc, s: dict):
    """
    Analyse financière détaillée (pages 42-66 du rapport source) :
    - Tableau ratios : PER, ROE, ROA, marge nette, croissance CA, dette/CP
    - Évolution sur 2-3 ans si disponible (CA, résultat net, ROE)
    - Forces et faiblesses financières (3 points max chacun)
    - Conclusion financière en 2-3 lignes
    - Date des données toujours affichée
    """
    _section_heading(doc, "ANALYSE FINANCIÈRE DÉTAILLÉE")

    date_donnees = _fund_val(s, "date_donnees_financieres")
    _narrative(doc,
               f"Données financières de référence — {date_donnees}.",
               size=8, italic=True, color="666666")

    _sub_heading(doc, "1.  Ratios financiers clés")
    ratios_rows = [
        ("PER (Price/Earnings)",  _fund_val(s, "per"),           _fund_val(s, "per_date")),
        ("ROE (Return on Equity)", _fund_val(s, "roe"),          _fund_val(s, "roe_date")),
        ("ROA (Return on Assets)", _fund_val(s, "roa"),          _fund_val(s, "roa_date")),
        ("Marge nette",            _fund_val(s, "marge_nette"),  _fund_val(s, "mn_date")),
        ("Croissance CA",          _fund_val(s, "croissance_ca"), _fund_val(s, "croissance_ca_date")),
        ("Dette / Capitaux propres", _fund_val(s, "dette_cp"),   _fund_val(s, "dette_cp_date")),
    ]
    tbl_r = doc.add_table(rows=1 + len(ratios_rows), cols=3)
    tbl_r.style = "Table Grid"
    for i, h in enumerate(["Ratio", "Valeur", "Date"]):
        _cw(tbl_r.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")
    for i, (label, val, dt) in enumerate(ratios_rows, start=1):
        bg_val = "F5F5F5" if val == "N/D" else "FFFFFF"
        bg_dt  = "F5F5F5" if dt  == "N/D" else "FFFFFF"
        _cw(tbl_r.rows[i].cells[0], label, bold=True, size=8, bg="EBF0FA")
        _cw(tbl_r.rows[i].cells[1], val, size=8, bg=bg_val)
        _cw(tbl_r.rows[i].cells[2], dt,  size=8, bg=bg_dt)
    doc.add_paragraph()

    ca_n   = _fund_val(s, "ca")
    ca_n_1 = _fund_val(s, "ca_n_1")
    ca_n_2 = _fund_val(s, "ca_n_2")
    rn_n   = _fund_val(s, "resultat_net")
    rn_n_1 = _fund_val(s, "resultat_n_1")
    rn_n_2 = _fund_val(s, "resultat_n_2")
    roe_n   = _fund_val(s, "roe")
    roe_n_1 = _fund_val(s, "roe_n_1")
    roe_n_2 = _fund_val(s, "roe_n_2")

    has_history = any(v != "N/D" for v in (ca_n_1, ca_n_2, rn_n_1, rn_n_2, roe_n_1, roe_n_2))
    if has_history:
        _sub_heading(doc, "2.  Évolution sur 2-3 ans")
        tbl_h = doc.add_table(rows=4, cols=4)
        tbl_h.style = "Table Grid"
        for i, h in enumerate(["Indicateur", "N-2", "N-1", "N (réf)"]):
            _cw(tbl_h.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")
        evo_rows = [
            ("Chiffre d'affaires", ca_n_2, ca_n_1, ca_n),
            ("Résultat net",       rn_n_2, rn_n_1, rn_n),
            ("ROE",                roe_n_2, roe_n_1, roe_n),
        ]
        for i, (lbl, v2, v1, v0) in enumerate(evo_rows, start=1):
            _cw(tbl_h.rows[i].cells[0], lbl, bold=True, size=8, bg="EBF0FA")
            _cw(tbl_h.rows[i].cells[1], v2, size=8, bg="F5F5F5" if v2 == "N/D" else "FFFFFF")
            _cw(tbl_h.rows[i].cells[2], v1, size=8, bg="F5F5F5" if v1 == "N/D" else "FFFFFF")
            _cw(tbl_h.rows[i].cells[3], v0, size=8, bg="F5F5F5" if v0 == "N/D" else "FFFFFF")
        doc.add_paragraph()
    else:
        _narrative(doc,
                   "Historique sur 2-3 ans non disponible dans le rapport source.",
                   size=8, italic=True, color="888888")

    forces = _sl(s, "forces_financieres")
    faibles = _sl(s, "faiblesses_financieres")

    # ── Enrichir forces/faiblesses depuis les ratios extraits ────────────────
    # Générer des forces/faiblesses additionnelles depuis les champs numériques
    def _enrich_from_ratios(s):
        extra_forces, extra_faibles = [], []
        def _pf(v):
            try: return float(str(v).replace('%','').replace(',','.').replace(' ',''))
            except: return None

        roe  = _pf(s.get("roe"))
        roa  = _pf(s.get("roa"))
        marge = _pf(s.get("marge_nette"))
        per  = _pf(s.get("per"))
        croiss = _pf(s.get("croissance_ca_1an") or s.get("croissance_ca"))
        cap_ex = _pf(s.get("coefficient_exploitation"))
        cout_r = _pf(s.get("cout_du_risque"))
        autono = _pf(s.get("autonomie_financiere"))

        if roe and roe > 15: extra_forces.append(f"ROE élevé ({roe:.1f}%) — forte rentabilité des fonds propres")
        elif roe and roe < 8: extra_faibles.append(f"ROE faible ({roe:.1f}%) — rentabilité des fonds propres insuffisante")

        if roa and roa > 2: extra_forces.append(f"ROA solide ({roa:.2f}%) — bonne rentabilité des actifs")
        elif roa and roa < 0.8: extra_faibles.append(f"ROA faible ({roa:.2f}%) — actifs peu rentabilisés")

        if marge and marge > 30: extra_forces.append(f"Marge nette élevée ({marge:.1f}%) — profitabilité excellente")
        elif marge and marge < 10: extra_faibles.append(f"Marge nette faible ({marge:.1f}%) — pression sur la rentabilité")

        if croiss and croiss > 10: extra_forces.append(f"Forte croissance CA ({croiss:+.1f}%) — dynamisme commercial")
        elif croiss and croiss < 0: extra_faibles.append(f"CA en recul ({croiss:+.1f}%) — contraction de l'activité")

        if cap_ex and cap_ex < 50: extra_forces.append(f"Coefficient d'exploitation maîtrisé ({cap_ex:.1f}%) — efficacité opérationnelle")
        elif cap_ex and cap_ex > 65: extra_faibles.append(f"Coefficient d'exploitation élevé ({cap_ex:.1f}%) — coûts structurels lourds")

        if cout_r and cout_r > 15: extra_faibles.append(f"Coût du risque élevé ({cout_r:.1f}%) — provisions importantes")
        if autono and autono < 10: extra_faibles.append(f"Autonomie financière limitée ({autono:.1f}%) — forte dépendance externe")
        elif autono and autono > 20: extra_forces.append(f"Bonne autonomie financière ({autono:.1f}%) — structure bilancielle solide")

        return extra_forces, extra_faibles

    ef, ew = _enrich_from_ratios(s)
    # Fusionner sans doublons (max 6 par catégorie)
    all_forces  = list(dict.fromkeys(forces  + ef))[:6]
    all_faibles = list(dict.fromkeys(faibles + ew))[:6]

    _sub_heading(doc, "3.  Forces et faiblesses financières")
    if all_forces or all_faibles:
        max_rows = max(len(all_forces), len(all_faibles), 1)
        tbl_ff = doc.add_table(rows=1 + max_rows, cols=2)
        tbl_ff.style = "Table Grid"
        _cw(tbl_ff.rows[0].cells[0], "✚  FORCES", bold=True, size=9, bg="0F9D58", color="FFFFFF")
        _cw(tbl_ff.rows[0].cells[1], "−  FAIBLESSES", bold=True, size=9, bg="D93025", color="FFFFFF")
        for i in range(max_rows):
            f_txt = all_forces[i]  if i < len(all_forces)  else "—"
            w_txt = all_faibles[i] if i < len(all_faibles) else "—"
            _cw(tbl_ff.rows[i + 1].cells[0], f"• {f_txt}", size=8, bg="EBF7EE" if f_txt != "—" else "F5F5F5")
            _cw(tbl_ff.rows[i + 1].cells[1], f"• {w_txt}", size=8, bg="FDEEEE" if w_txt != "—" else "F5F5F5")
        doc.add_paragraph()
    else:
        _narrative(doc,
                   "Forces et faiblesses financières non identifiées dans le rapport source.",
                   size=8, italic=True, color="888888")

    _sub_heading(doc, "4.  Conclusion financière")
    synth = _s(s, "synthese_financiere")
    if synth and synth not in ("", "—", "null", "None"):
        _key_bloc(doc, "BILAN FINANCIER :", synth, "E8F0FB", "1558A7")
    else:
        parts = []
        if _fund_val(s, "per") != "N/D":
            parts.append(f"PER {_fund_val(s, 'per')}")
        if _fund_val(s, "roe") != "N/D":
            parts.append(f"ROE {_fund_val(s, 'roe')}")
        if _fund_val(s, "marge_nette") != "N/D":
            parts.append(f"marge nette {_fund_val(s, 'marge_nette')}")
        if _fund_val(s, "croissance_ca") != "N/D":
            parts.append(f"croissance CA {_fund_val(s, 'croissance_ca')}")
        if parts:
            txt = (
                "Profil financier synthétique — " + " · ".join(parts)
                + f". Données arrêtées au {date_donnees}. "
                "Une analyse complémentaire sera intégrée dès la disponibilité "
                "de la synthèse qualitative du rapport source."
            )
        else:
            txt = (
                "Aucune synthèse financière n'a été extraite pour cette société dans "
                f"le rapport de référence ({date_donnees}). "
                "Les données seront enrichies lors de la prochaine publication."
            )
        _key_bloc(doc, "BILAN FINANCIER :", txt, "E8F0FB", "1558A7")


# ── Données financières structurées (5 sous-sections) ───────────────────────

def _emoji_bg(emoji: str) -> str:
    """Couleur de fond de la cellule signal selon l'émoji."""
    if "🟢" in emoji:
        return "C6EFCE"
    if "🟡" in emoji:
        return "FFEB9C"
    if "🔴" in emoji:
        return "FFC7CE"
    return "E3F2FD"  # 🔵 informatif


def _ratio_emoji(val, good, vigilance, lower_better=False) -> str:
    """Émoji selon la position de val par rapport aux seuils."""
    f = _pct_to_float(val)
    if f is None:
        return "🔵"
    if lower_better:
        if f <= good:
            return "🟢"
        if f <= vigilance:
            return "🟡"
        return "🔴"
    if f >= good:
        return "🟢"
    if f >= vigilance:
        return "🟡"
    return "🔴"


def _build_financial_subtable(doc, s: dict, rows):
    """
    Tableau 4 colonnes (Indicateur · Valeur · Interprétation · Signal) pour
    une sous-section financière.
    """
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Indicateur", "Valeur", "Interprétation", "Signal"]):
        _cw(tbl.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")

    for i, (label, key, interp, emoji) in enumerate(rows, start=1):
        val = _fmt_amount(s.get(key))
        bg_val = "F5F5F5" if val == "—" else "FFFFFF"
        em = emoji if val != "—" else "🔵"
        _cw(tbl.rows[i].cells[0], label, bold=True, size=8, bg="EBF0FA")
        _cw(tbl.rows[i].cells[1], val, size=8, bg=bg_val)
        _cw(tbl.rows[i].cells[2], interp, size=7, bg="FFFFFF")
        _cw(tbl.rows[i].cells[3], em, size=10, bg=_emoji_bg(em),
            align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def build_financial_data_complete(doc, s: dict, source_doc=None):
    """
    📊 DONNÉES FINANCIÈRES STRUCTURÉES & INTERPRÉTÉES
    Copie fidèle depuis le rapport source (tous les tableaux + titres de section).
    Fallback sur les données LLM si source_doc absent ou section non trouvée.
    """
    ticker = _s(s, "ticker", "?")

    _section_heading(doc, "📊 DONNÉES FINANCIÈRES STRUCTURÉES & INTERPRÉTÉES")

    # ── Priorité 1 : copie depuis le rapport source ───────────────────────────
    if source_doc is not None:
        fin_blocks = _extract_donnees_financieres_tables(source_doc, ticker)
        if fin_blocks:
            # Afficher source du rapport
            p_src = doc.add_paragraph()
            p_src.paragraph_format.space_before = Pt(0)
            p_src.paragraph_format.space_after  = Pt(4)
            r_src = p_src.add_run("Source : rapport source BRVM — données officielles")
            r_src.italic = True; r_src.font.size = Pt(8)
            r_src.font.color.rgb = _rgb("888888")

            for kind, el in fin_blocks:
                if kind == 'p':
                    # Paragraphe titre de section (📌 1. BILAN...)
                    from docx.oxml.ns import qn as _qn2
                    txt_raw = ''.join(n.text or '' for n in el.iter() if n.text).strip()
                    def _dd2(s):
                        n=len(s)
                        for d in (3,2):
                            if n%d==0 and s==s[:n//d]*d: return s[:n//d]
                        return s
                    txt = _dd2(txt_raw)
                    if not txt: continue
                    if 'DONNÉES FINANCIÈRES STRUCTURÉES' in txt.upper():
                        continue  # déjà affiché en titre
                    p_t = doc.add_paragraph()
                    p_t.paragraph_format.space_before = Pt(6)
                    p_t.paragraph_format.space_after  = Pt(2)
                    r_t = p_t.add_run(txt)
                    r_t.bold = True; r_t.font.size = Pt(9)
                    r_t.font.color.rgb = _rgb("1A237E")

                elif kind == 'tbl':
                    # Copier le tableau depuis la source
                    data = _read_source_tbl_for_copy(el)
                    if not data: continue
                    ncols = max(len(r) for r in data)
                    if ncols == 0: continue
                    tbl = doc.add_table(rows=len(data), cols=ncols)
                    tbl.style = "Table Grid"
                    for ri, row in enumerate(data):
                        bg_h = "1A237E" if ri == 0 else ("F5F5F5" if ri % 2 == 0 else "FFFFFF")
                        fg_h = "FFFFFF" if ri == 0 else "1A1A1A"
                        for ci, cell_text in enumerate(row):
                            if ci >= ncols: continue
                            # Colorer selon emoji signal
                            bg = bg_h
                            if ri > 0:
                                if '🟢' in cell_text: bg = "EBF7EE"
                                elif '🔴' in cell_text: bg = "FDEEEE"
                                elif '⚠️' in cell_text or '🟡' in cell_text: bg = "FFF8E6"
                                elif '🔵' in cell_text: bg = "E8F0FB"
                            _cw(tbl.rows[ri].cells[ci], cell_text[:150],
                                bold=(ri == 0), size=8 if ri > 0 else 8,
                                bg=bg, color=fg_h if ri == 0 else "1A1A1A")
                    doc.add_paragraph()
            return  # données source utilisées — pas de fallback

    # ── Fallback : données LLM ────────────────────────────────────────────────
    _section_heading(doc, "Données financières (LLM)", color="888888")

    date_donnees = _fund_val(s, "date_donnees_financieres")
    secteur = _s(s, "secteur", "—")
    _narrative(
        doc,
        f"Source : rapport d'analyse BRVM   ·   Secteur : {secteur}   ·   "
        f"Exercice de référence : {date_donnees}.",
        size=8, italic=True, color="666666",
    )

    _sub_heading(doc, "1.  BILAN ACTIF")
    _build_financial_subtable(doc, s, [
        ("Caisse & Banque Centrale", "caisse_banque_centrale",
         "Liquidités disponibles immédiatement.", "🟢"),
        ("Effets publics", "effets_publics",
         "Titres d'État détenus — actifs sûrs et liquides.", "🔵"),
        ("Créances interbancaires", "creances_interbancaires",
         "Prêts accordés aux autres établissements.", "🔵"),
        ("Créances clientèle", "creances_clientele",
         "Encours de crédits — moteur principal du PNB.", "🟢"),
        ("Immob. incorporelles", "immob_incorporelles",
         "Actifs immatériels (licences, logiciels).", "🔵"),
        ("Immob. corporelles", "immob_corporelles",
         "Bâtiments, agences, équipements.", "🔵"),
        ("Trésorerie active", "tresorerie_actif",
         "Disponibilités à court terme.", "🟢"),
        ("Total actif", "total_actif",
         "Taille de bilan globale.", "🔵"),
    ])

    _sub_heading(doc, "2.  BILAN PASSIF")
    _build_financial_subtable(doc, s, [
        ("Capital souscrit", "capital_souscrit",
         "Apports en capital des actionnaires.", "🔵"),
        ("Réserves", "reserves",
         "Bénéfices accumulés non distribués.", "🟢"),
        ("Capitaux propres", "capitaux_propres",
         "Fonds propres totaux — solidité financière.", "🟢"),
        ("Capitaux permanents", "capitaux_permanents",
         "Capitaux propres + dettes longues — stabilité.", "🟢"),
        ("Dettes interbancaires", "dettes_interbancaires",
         "Refinancement auprès d'autres banques.", "🟡"),
        ("Dettes clientèle", "dettes_clientele",
         "Dépôts collectés — ressource principale.", "🔵"),
        ("Dettes financières totales", "dettes_financieres_totales",
         "Endettement financier global.", "🟡"),
        ("Dettes totales", "dettes_totales",
         "Passif exigible total.", "🟡"),
    ])

    _sub_heading(doc, "3.  COMPTE DE RÉSULTAT")
    _build_financial_subtable(doc, s, [
        ("PNB (Produit Net Bancaire)", "pnb",
         "Chiffre d'affaires bancaire — agrégat d'activité.", "🟢"),
        ("Intérêts produits", "interets_produits",
         "Revenus des prêts et placements.", "🟢"),
        ("Intérêts charges", "interets_charges",
         "Coût de la collecte et du refinancement.", "🔴"),
        ("Commissions produits", "commissions_produits",
         "Revenus de services bancaires et frais.", "🟢"),
        ("Charges générales", "charges_generales",
         "Frais administratifs et d'exploitation.", "🟡"),
        ("Charges de personnel", "charges_personnel",
         "Masse salariale — premier poste de charge.", "🟡"),
        ("RBE (Résultat Brut Expl.)", "rbe",
         "Performance opérationnelle brute.", "🟢"),
        ("Résultat d'exploitation", "resultat_exploitation",
         "Rentabilité de l'activité courante.", "🟢"),
        ("Provisions", "provisions",
         "Coût du risque comptabilisé sur l'exercice.", "🟡"),
        ("Résultat net", "resultat_net",
         "Bénéfice final attribuable aux actionnaires.", "🟢"),
    ])

    _sub_heading(doc, "4.  RATIOS DE RENTABILITÉ")
    _build_financial_subtable(doc, s, [
        ("Marge opérationnelle", "marge_operationnelle",
         "Profitabilité de l'exploitation (cible ≥ 30%).",
         _ratio_emoji(s.get("marge_operationnelle"), good=30, vigilance=15)),
        ("Coefficient d'exploitation", "coefficient_exploitation",
         "Charges / PNB — efficience (cible < 60%).",
         _ratio_emoji(s.get("coefficient_exploitation"),
                      good=60, vigilance=70, lower_better=True)),
        ("Coût du risque", "cout_du_risque",
         "Provisions / encours — qualité du portefeuille (< 1%).",
         _ratio_emoji(s.get("cout_du_risque"),
                      good=1, vigilance=2, lower_better=True)),
        ("Marge nette", "marge_nette",
         "Résultat net / PNB — rentabilité finale.",
         _ratio_emoji(s.get("marge_nette"), good=15, vigilance=5)),
        ("ROE", "roe",
         "Rentabilité des fonds propres (cible ≥ 12%).",
         _ratio_emoji(s.get("roe"), good=12, vigilance=5)),
        ("ROA", "roa",
         "Rentabilité des actifs (cible ≥ 1,5%).",
         _ratio_emoji(s.get("roa"), good=1.5, vigilance=0.5)),
    ])

    _sub_heading(doc, "5.  STRUCTURE & LIQUIDITÉ")
    _build_financial_subtable(doc, s, [
        ("Autonomie financière", "autonomie_financiere",
         "Capitaux propres / total bilan — indépendance.",
         _ratio_emoji(s.get("autonomie_financiere"), good=15, vigilance=8)),
        ("Dépendance financière", "dependance_financiere",
         "Dettes / total bilan — exposition au passif externe.",
         _ratio_emoji(s.get("dependance_financiere"),
                      good=70, vigilance=85, lower_better=True)),
        ("Gearing", "gearing",
         "Dettes / capitaux propres — levier d'endettement.",
         _ratio_emoji(s.get("gearing"),
                      good=1, vigilance=2, lower_better=True)),
        ("Solvabilité générale", "solvabilite_generale",
         "Actif / dettes — capacité de remboursement.",
         _ratio_emoji(s.get("solvabilite_generale"),
                      good=1.2, vigilance=1.05)),
        ("Liquidité immédiate", "liquidite_immediate",
         "Capacité à honorer les engagements courts (cible ≥ 1).",
         _ratio_emoji(s.get("liquidite_immediate"),
                      good=1, vigilance=0.5)),
        ("Couverture des intérêts", "couverture_interets",
         "Résultat d'exploitation / charges d'intérêts.",
         _ratio_emoji(s.get("couverture_interets"),
                      good=3, vigilance=1.5)),
    ])


def build_analyse_fondamentale_partie3(doc, s: dict, source_doc=None):
    """PARTIE 3 : Analyse fondamentale depuis le rapport source."""
    _section_heading(doc, "PARTIE 3 — Analyse fondamentale (section critique)")
    ticker = _s(s, "ticker", "?")
    if source_doc is not None:
        parties3 = _extract_parties(source_doc, ticker)
        if parties3.get('p3'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(parties3['p3'])
            r.font.size = Pt(9)
            return
    # Fallback LLM
    analyse = _s(s, "synthese_financiere")
    if analyse and analyse not in ("", "—"):
        _narrative(doc, analyse, size=9)


def build_conclusion(doc, s: dict):
    """
    Conclusion d'investissement :
    0. Synthèse PARTIE 4 depuis le rapport source (si disponible)
    1. Matrice Risque × Horizon de placement
    2. Divergences majeures
    3. Recommandation finale avec action claire
    """
    _section_heading(doc, "CONCLUSION D'INVESTISSEMENT")

    # ── PARTIE 4 : Synthèse depuis le rapport source ──────────────────────────
    # Utiliser _extract_parties pour avoir le texte complet et non tronqué
    _src_doc_c = s.get("_source_doc_ref")
    ticker_c   = _s(s, "ticker", "?")
    _src_indicators = s.get("_src_indicators") or {}

    if _src_doc_c is not None:
        parties_c = _extract_parties(_src_doc_c, ticker_c)
        p4_text = parties_c.get('p4', '')
        if p4_text:
            p_s4 = doc.add_paragraph()
            p_s4.paragraph_format.space_before = Pt(2)
            p_s4.paragraph_format.space_after  = Pt(6)
            p_s4.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_s4 = p_s4.add_run(p4_text)
            r_s4.font.size = Pt(9)
            doc.add_paragraph()
    elif _src_indicators.get('partie4'):
        # Fallback sur la version de _extract_source_indicators
        p_s4 = doc.add_paragraph()
        p_s4.paragraph_format.space_before = Pt(2)
        p_s4.paragraph_format.space_after  = Pt(6)
        p_s4.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_s4 = p_s4.add_run(_src_indicators['partie4'])
        r_s4.font.size = Pt(9)
        doc.add_paragraph()

    ticker = _s(s, "ticker", "—")
    score_f = _score_f(s)
    score_str = f"{score_f:.0f}" if s.get("score") is not None else "—"
    score_label, score_color = _score_label_color(score_f)
    reco = _s(s, "reco", "NEUTRE")
    decision = _s(s, "decision", "SURVEILLER")
    risque = _s(s, "risque", "modéré")
    divergence = _s(s, "divergence", "aucune")
    confiance = _s(s, "confiance", "Modérée")
    resume = _s(s, "resume_rapport", "—")

    if score_f >= 70:
        horizon = "Long terme (> 12 mois)"
        profil_inv = "Investissement core / portefeuille défensif"
    elif score_f >= 55:
        horizon = "Moyen terme (6-12 mois)"
        profil_inv = "Opportunité de portage / portefeuille équilibré"
    elif score_f >= 40:
        horizon = "Court terme (3-6 mois)"
        profil_inv = "Position tactique / portefeuille offensif"
    else:
        horizon = "Très court terme ou abstention"
        profil_inv = "Profil spéculatif — risque élevé"

    _sub_heading(doc, "1.  Matrice Risque × Horizon de placement")

    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Valeur", "Risque", "Horizon recommandé", "Profil investisseur", "Score"]):
        _cw(tbl.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")

    risque_bg = _risque_bg(risque)
    _cw(tbl.rows[1].cells[0], ticker, bold=True, size=9, bg="E8EAF6")
    _cw(tbl.rows[1].cells[1], risque.capitalize(), bold=True, size=9, bg=risque_bg)
    _cw(tbl.rows[1].cells[2], horizon, size=9, bg="F5F5F5")
    _cw(tbl.rows[1].cells[3], profil_inv, size=8, bg="F5F5F5")
    _cw(tbl.rows[1].cells[4], f"{score_str}/100 — {score_label}", bold=True, size=9, bg=risque_bg)

    doc.add_paragraph()

    _sub_heading(doc, "2.  Divergences identifiées")

    div_lower = str(divergence).lower()
    if div_lower not in ("aucune", "—", "", "none"):
        _key_bloc(doc,
                  "⚡  DIVERGENCE :",
                  f"{divergence}. "
                  "Ce signal dissonant doit être pris en compte avant toute décision — "
                  "les signaux techniques et fondamentaux ne convergent pas.",
                  "FFEB9C", "7D6608")
    else:
        p_nd = doc.add_paragraph()
        p_nd.paragraph_format.space_after = Pt(4)
        r_nd = p_nd.add_run("✔  Aucune divergence majeure identifiée. Cohérence technique/fondamental.")
        r_nd.font.size = Pt(9)
        r_nd.italic = True
        r_nd.font.color.rgb = _rgb("0F9D58")

    _sub_heading(doc, "3.  Recommandation d'investissement")

    action_map = {
        "ACHAT FORT":  ("RENFORCER SIGNIFICATIVEMENT", "C6EFCE", "0F9D58",
                        "Tous les signaux sont alignés. C'est le moment d'augmenter l'exposition. "
                        "Priorité absolue dans l'allocation."),
        "ACHAT":       ("RENFORCER PROGRESSIVEMENT", "C6EFCE", "0F9D58",
                        "Signaux favorables. Entrée progressive recommandée sur 2-3 séances. "
                        "Ne pas entrer en totalité en une seule transaction."),
        "SURVEILLER":  ("CONSERVER ET SURVEILLER", "FFEB9C", "7D6608",
                        "Maintenir la position existante. Ne pas renforcer pour l'instant. "
                        "Attendre la prochaine confirmation directionnelle."),
        "PRUDENCE":    ("ALLÉGER PARTIELLEMENT", "FFF0E6", "C0392B",
                        "Réduire l'exposition de 25-50%. Sécuriser une partie des plus-values. "
                        "Conserver le solde en attente de signal de retournement."),
        "ÉVITER":      ("ALLÉGER OU SORTIR", "FFC7CE", "D93025",
                        "Signaux défavorables confirmés. Réduire ou clôturer la position. "
                        "Ne pas initier de nouvelle entrée sur ce titre."),
    }
    action_label, action_bg, action_fg, action_text = action_map.get(
        str(decision).upper(),
        ("SURVEILLER", "FFEB9C", "7D6608", "Maintenir la position. Aucune action urgente requise.")
    )

    tbl2 = doc.add_table(rows=1, cols=1)
    cell2 = tbl2.rows[0].cells[0]
    _cell_bg(cell2, action_bg)
    _cell_margins(cell2, 140)
    p2 = cell2.paragraphs[0]
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_action = p2.add_run(f"▶  {action_label}")
    r_action.bold = True
    r_action.font.size = Pt(13)
    r_action.font.color.rgb = _rgb(action_fg)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

    _narrative(doc, action_text)

    _narrative(doc,
               f"{ticker} — Score {score_str}/100 ({score_label}). "
               f"Reco : {reco}. Confiance : {confiance}. "
               f"Risque : {risque} — Horizon : {horizon.lower()}.")

    if resume and resume not in ("", "—"):
        _key_bloc(doc, "RÉSUMÉ :", resume, "E8F0FB", "1558A7")


# ── Pied de page ──────────────────────────────────────────────────────────────

def _pied(doc, date_str: str, freq: str = "JOUR", period_info: dict = None):
    p = doc.add_paragraph()
    _cp(p, 4, 0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "CCCCCC")
    pBdr.append(top)
    pPr.append(pBdr)
    period_suffix = ""
    if freq != "JOUR" and period_info:
        period_suffix = (
            f"   |   {period_info.get('freq_label', freq)} : "
            f"{period_info.get('date_debut', '—')} → {period_info.get('date_fin', '—')}"
            f" ({period_info.get('nb_seances', '—')} séances)"
        )
    r = p.add_run(f"Document confidentiel — Analyse BRVM — {date_str}{period_suffix}")
    r.font.size = Pt(7)
    r.italic = True
    r.font.color.rgb = _rgb("999999")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Assemblage de la fiche ────────────────────────────────────────────────────

def build_risques(doc, s: dict, source_doc=None):
    """
    Section PROFIL DE RISQUE de la fiche société.

    Affiche 5 blocs :
      1. SCORE DE RISQUE calculé (extrait du rapport source : score/100, barre,
         tableau des 5 critères — Volatilité, Bêta, Liquidité, Divergence, Stabilité)
      2. Indicateurs de risque globaux (niveau, horizon, volatilité, bêta)
      3. Faiblesses financières + risques (LLM)
      4. Divergence technique/fondamentale
      5. Texte narratif PARTIE 3 & 4 (rapport source)
    """
    _section_heading(doc, "PROFIL DE RISQUE")

    ticker    = _s(s, "ticker", "?")
    risque    = _s(s, "risque", "—")
    horizon   = _s(s, "horizon", "—")
    volat     = _s(s, "volatilite", "—")
    # Formater le bêta avec 4 décimales (valeur réelle depuis le tableau score risque)
    beta_raw = _s(s, "beta", "—")
    try:
        beta_f = float(str(beta_raw).split("(")[0].replace(",", ".").strip())
        beta = f"{beta_f:.4f}"
    except (ValueError, AttributeError):
        beta = beta_raw
    confiance = _s(s, "confiance", "—")
    divergence = _s(s, "divergence", "—")

    # ── Couleur selon niveau de risque ────────────────────────────────────────
    r_low = risque.lower()
    if "faible" in r_low:
        risk_bg, risk_fg = "C6EFCE", "0F9D58"
        risk_emoji = "🟢"
    elif "élevé" in r_low or "eleve" in r_low:
        risk_bg, risk_fg = "FFC7CE", "D93025"
        risk_emoji = "🔴"
    else:
        risk_bg, risk_fg = "FFEB9C", "E37400"
        risk_emoji = "🟡"

    # ── 0. Score de risque calculé (depuis le rapport source) ────────────────
    if source_doc is not None:
        # Ajouter la Stabilité des rendements dans les indicateurs globaux
        stab_rendements = _extract_stabilite_rendements(source_doc, ticker)
        if stab_rendements:
            # Nettoyer la valeur
            stab_clean = re.sub(r'^[^a-zA-Z0-9σ]*', '', stab_rendements).strip()
        else:
            stab_clean = None

        risk_data = _extract_risk_score_data(source_doc, ticker)
        if risk_data:
            score_v = risk_data["score_str"]
            niveau_v = risk_data["niveau"]
            barre_v  = risk_data["barre"]

            # En-tête score
            p_sc = doc.add_paragraph()
            p_sc.paragraph_format.space_before = Pt(4)
            p_sc.paragraph_format.space_after  = Pt(1)
            r_sc = p_sc.add_run(f"⚠️  SCORE DE RISQUE — {score_v}/100  →  {niveau_v}")
            r_sc.bold = True; r_sc.font.size = Pt(11); r_sc.font.color.rgb = _rgb(risk_fg)

            # Barre ASCII
            if barre_v:
                p_bar = doc.add_paragraph()
                p_bar.paragraph_format.space_before = Pt(0)
                p_bar.paragraph_format.space_after  = Pt(4)
                r_bar = p_bar.add_run(barre_v)
                r_bar.font.name = "Courier New"; r_bar.font.size = Pt(10)
                r_bar.font.color.rgb = _rgb(risk_fg)

            # Tableau des 5 critères
            if risk_data["criteres"]:
                _sub_heading(doc, "Détail des critères de risque")
                tbl_r = doc.add_table(rows=1, cols=3)
                tbl_r.style = "Table Grid"
                for ci, h in enumerate(["Critère", "Valeur mesurée", "Interprétation"]):
                    _cw(tbl_r.rows[0].cells[ci], h, bold=True, size=8, bg="1A237E", color="FFFFFF")
                for crit in risk_data["criteres"]:
                    tr = tbl_r.add_row()
                    # Nettoyer le nom du critère (enlever emoji + poids)
                    nom_clean = re.sub(r'^[^\w]+', '', crit["nom"])
                    nom_clean = re.sub(r'\s*\(\d+%\)', '', nom_clean).strip()
                    _cw(tr.cells[0], nom_clean[:40], bold=True, size=8, bg="EBF0FA")
                    _cw(tr.cells[1], crit["valeur"][:70], size=8)
                    interp_clean = re.sub(r'^[⚠️🟢🔴🔵ℹ️]+\s*', '', crit["interp"]).strip()
                    _cw(tr.cells[2], interp_clean[:60], size=8,
                        bg="EBF7EE" if "🟢" in crit["interp"] else
                           ("FDEEEE" if "🔴" in crit["interp"] else "FFF8E6"))
                doc.add_paragraph()

    # ── 1. Bloc indicateurs globaux ───────────────────────────────────────────
    # Construire les colonnes selon disponibilité de stab_rendements
    _headers = ["Niveau de risque", "Horizon", "Volatilité", "Bêta", "Confiance"]
    _vals    = [f"{risk_emoji} {risque}", horizon, volat, beta, confiance]
    _bgs     = [risk_bg, "F5F5F5", "F5F5F5", "F5F5F5", "F5F5F5"]
    _fgs     = [risk_fg,  "333333", "333333", "333333", "333333"]
    if stab_clean:
        _headers.append("📉 Stabilité rendements")
        _vals.append(stab_clean[:50])
        _bgs.append("F5F5F5"); _fgs.append("333333")

    ncols = len(_headers)
    tbl = doc.add_table(rows=2, cols=ncols)
    tbl.style = "Table Grid"
    for i, h in enumerate(_headers):
        _cw(tbl.rows[0].cells[i], h, bold=True, size=8, bg="1A237E", color="FFFFFF")
    for i, (v, bg, fg) in enumerate(zip(_vals, _bgs, _fgs)):
        _cw(tbl.rows[1].cells[i], v, bold=(i == 0), size=9, bg=bg, color=fg)
    doc.add_paragraph()

    # ── 2. Faiblesses financières (LLM) ──────────────────────────────────────
    faiblesses = _sl(s, "faiblesses_financieres")
    risques_list = _sl(s, "risques")
    all_risks = list(dict.fromkeys(faiblesses + risques_list))  # dédoublonner

    if all_risks:
        _sub_heading(doc, "Facteurs de risque identifiés")
        tbl2 = doc.add_table(rows=len(all_risks), cols=2)
        tbl2.style = "Table Grid"
        for i, risk_txt in enumerate(all_risks[:8]):
            # Numéro
            _cw(tbl2.rows[i].cells[0], str(i + 1), bold=True, size=9,
                bg=risk_bg, color=risk_fg)
            # Texte du risque
            _cw(tbl2.rows[i].cells[1], str(risk_txt).strip(), size=9,
                bg="FFF8F8" if "élevé" in risque.lower() else "FFFFF0")
        doc.add_paragraph()

    # ── 3. Divergence technique/fondamentale ──────────────────────────────────
    if divergence and divergence.lower() not in ("aucune", "—", "", "none"):
        tbl3 = doc.add_table(rows=1, cols=1)
        tbl3.style = "Table Grid"
        cell3 = tbl3.rows[0].cells[0]
        _cell_bg(cell3, "FFF3CD")
        tcPr = cell3._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side in ("top","bottom","left","right"):
            em = OxmlElement(f"w:{side}")
            em.set(qn("w:w"), "100")
            em.set(qn("w:type"), "dxa")
            tcMar.append(em)
        tcPr.append(tcMar)
        p3 = cell3.paragraphs[0]
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(0)
        r3a = p3.add_run("⚡ Divergence technique / fondamentale : ")
        r3a.bold = True; r3a.font.size = Pt(9); r3a.font.color.rgb = _rgb("7D5200")
        r3b = p3.add_run(divergence)
        r3b.font.size = Pt(9); r3b.font.color.rgb = _rgb("444444")
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # ── 4. Texte narratif du rapport source ───────────────────────────────────
    # Extraire les paragraphes "PARTIE 3 & 4" de la section société dans le doc source
    if source_doc is not None:
        _sub_heading(doc, "Analyse du risque (source rapport)")
        risk_paras = _extract_risk_narrative(source_doc, ticker)
        if risk_paras:
            for txt in risk_paras[:4]:   # max 4 phrases
                p_r = doc.add_paragraph()
                p_r.paragraph_format.space_before = Pt(1)
                p_r.paragraph_format.space_after  = Pt(3)
                p_r.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
                r_r = p_r.add_run(txt)
                r_r.font.size = Pt(9)
        else:
            # Fallback : synthèse financière
            synth = _s(s, "synthese_financiere")
            if synth and synth not in ("", "—"):
                _narrative(doc, synth, size=9, italic=True, color="555555")


def _extract_company_section(source_doc, ticker: str) -> tuple:
    """
    Localise la section individuelle d'une société dans le rapport source.
    Retourne (start_idx, end_idx, elements_list) ou (None, None, []).
    """
    import re as _re
    elements = list(source_doc.element.body.iterchildren())

    def _ps(el):
        pPr = el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        return pPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '') if pPr is not None else ''
    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s

    ticker_start = None
    ticker_end   = None
    rx = _re.compile(r'^\s*\d+\.\s*' + _re.escape(ticker) + r'\b')

    for i, child in enumerate(elements):
        if child.tag.split('}')[-1] != 'p': continue
        s = _ps(child)
        t = _dd(_pt(child).strip())
        if s in ('Titre2', 'Heading2') and rx.match(t):
            ticker_start = i
        elif ticker_start is not None and s in ('Titre2', 'Heading2') and not rx.match(t):
            ticker_end = i
            break

    if ticker_start is None:
        return None, None, elements
    return ticker_start, (ticker_end or min(ticker_start + 150, len(elements))), elements

def _extract_risk_score_data(source_doc, ticker: str) -> dict | None:
    """
    Extrait les données du SCORE DE RISQUE d'une société depuis le rapport source.
    Structure source :
      Para  : "⚠️ SCORE DE RISQUE — 21.8/100 → Moyen"   (triplé)
      Para  : "████░░░  21.8/100"                        (barre ASCII)
      Table : 5 critères × 4 colonnes [Critère, Valeur, Formule, Interprétation]

    Retourne dict {score_str, niveau, barre, criteres: [{nom, valeur, formule, interp}]}
    ou None si non trouvé.
    """
    ticker_start, ticker_end, elements = _extract_company_section(source_doc, ticker)
    if ticker_start is None:
        return None

    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s
    def _read_tbl(tbl_el):
        rows = []
        for tr in tbl_el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
            cells = [_dd(''.join(n.text or '' for n in tc.iter() if n.text).strip())
                     for tc in tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')]
            if cells: rows.append(cells)
        return rows

    score_str = None
    niveau    = None
    barre     = None
    criteres  = []

    for i in range(ticker_start, ticker_end):
        child = elements[i]
        tag   = child.tag.split('}')[-1]
        txt   = _dd(_pt(child).strip()) if tag == 'p' else ''

        if tag == 'p' and 'SCORE DE RISQUE' in txt.upper():
            # Extraire score et niveau : "⚠️ SCORE DE RISQUE — 21.8/100 → Moyen"
            m = re.search(r'(\d+[\.,]\d+)/100\s*[→\->\s]+\s*(\w+)', txt)
            if m:
                score_str = m.group(1).replace(',', '.')
                niveau    = m.group(2).strip()

        elif tag == 'p' and barre is None and score_str and '░' in txt or (tag == 'p' and '█' in txt and score_str):
            barre = txt[:60]

        elif tag == 'tbl' and score_str and not criteres:
            data = _read_tbl(child)
            if not data or not data[0]: continue
            h = ' '.join(data[0]).lower()
            if 'critère' in h or 'critere' in h or 'poids' in h:
                for row in data[1:]:
                    if len(row) >= 2:
                        criteres.append({
                            'nom':    row[0][:60] if row[0] else '—',
                            'valeur': row[1][:80] if len(row) > 1 else '—',
                            'formule': row[2][:60] if len(row) > 2 else '—',
                            'interp':  row[3][:80] if len(row) > 3 else '—',
                        })
                break   # on a trouvé le tableau critères

    if score_str is None:
        return None
    return {'score_str': score_str, 'niveau': niveau or '—', 'barre': barre or '', 'criteres': criteres}


def _extract_cours_commentary(source_doc, ticker: str) -> str | None:
    """
    Extrait le commentaire d'analyse du cours (PARTIE 1) depuis le rapport source.
    Le texte est dans un paragraphe contenant "PARTIE 1" ou "ANALYSE DU COURS"
    suivi du texte narratif.
    Retourne le texte nettoyé (max 400 chars) ou None.
    """
    ticker_start, ticker_end, elements = _extract_company_section(source_doc, ticker)
    if ticker_start is None:
        return None

    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s

    for i in range(ticker_start, ticker_end):
        child = elements[i]
        if child.tag.split('}')[-1] != 'p': continue
        txt = _dd(_pt(child).strip())
        # Détecter le paragraphe PARTIE 1 (contient le texte d'analyse du cours)
        if ('PARTIE 1' in txt.upper() and 'COURS' in txt.upper() and len(txt) > 80):
            # Nettoyer le texte (enlever "### **PARTIE 1 ...**" markdown)
            clean = re.sub(r'\*{1,3}[^*]+\*{1,3}', '', txt)
            clean = re.sub(r'#+\s*PARTIE\s+\d+[^.\n]*', '', clean)
            clean = re.sub(r'\s{2,}', ' ', clean).strip()
            if len(clean) > 60:
                return clean[:400] + ('…' if len(clean) > 400 else '')
        # Aussi détecter le paragraphe qui parle directement du cours sur 100 jours
        if 'sur les 100 derniers jours' in txt.lower() or '100 derniers jours' in txt.lower():
            clean = re.sub(r'\*{1,3}', '', txt).strip()
            return clean[:400] + ('…' if len(clean) > 400 else '')
    return None


def _extract_risk_narrative(source_doc, ticker: str) -> list:
    """
    Parcourt le document source, trouve la section du ticker,
    et extrait les phrases de risque des paragraphes narratifs
    (PARTIE 3 Fondamentale + PARTIE 4 Conclusion).
    Retourne une liste de phrases courtes (max 200 chars chacune).
    """
    ticker_start, ticker_end, elements = _extract_company_section(source_doc, ticker)
    if ticker_start is None:
        return []

    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _ps(el):
        pPr = el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        return pPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val','') if pPr is not None else ''
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s

    risk_keywords = ('risque', 'volatil', 'fragil', 'surveill', 'vigilance',
                     'attention', 'incertitude', 'faiblesse', 'endett',
                     'pression', 'dégradation', 'correction')
    risk_sentences = []

    for i in range(ticker_start, ticker_end):
        child = elements[i]
        if child.tag.split('}')[-1] != 'p': continue
        s = _ps(child)
        if s in ('Titre2', 'Heading2', 'Titre3', 'Heading3'): continue
        txt = _dd(_pt(child).strip())
        if not txt or len(txt) < 20: continue
        if not any(k in txt.lower() for k in risk_keywords): continue

        sentences = re.split(r'(?<=[.!?])\s+', txt)
        for sent in sentences:
            sent = re.sub(r'\*+', '', sent).strip()
            sent = re.sub(r'#+\s*PARTIE\s+\d+[^.]*\.?', '', sent).strip()
            if len(sent) < 30: continue
            if not any(k in sent.lower() for k in risk_keywords): continue
            if len(sent) > 200: sent = sent[:197] + '…'
            if sent not in risk_sentences:
                risk_sentences.append(sent)
            if len(risk_sentences) >= 4: break
        if len(risk_sentences) >= 4: break

    return risk_sentences


def _extract_stabilite_rendements(source_doc, ticker: str) -> str | None:
    """
    Extrait la ligne '📉 Stabilité des rendements (15%)' du tableau score risque.
    Retourne la cellule complète (ex: 'σ_j=1.56%/j (Stable) — Vol. annualisée≈24.8%')
    ou None si non trouvé.
    """
    ticker_start, ticker_end, elements = _extract_company_section(source_doc, ticker)
    if ticker_start is None:
        return None

    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s
    def _read_tbl(tbl_el):
        rows = []
        for tr in tbl_el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
            cells = [_dd(''.join(n.text or '' for n in tc.iter() if n.text).strip())
                     for tc in tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')]
            if cells: rows.append(cells)
        return rows

    for i in range(ticker_start, ticker_end):
        child = elements[i]
        if child.tag.split('}')[-1] != 'tbl': continue
        data = _read_tbl(child)
        for row in data:
            if row and 'tabilit' in row[0].lower() and 'rendement' in row[0].lower():
                return row[1] if len(row) > 1 else row[0]
    return None


def _extract_donnees_financieres_tables(source_doc, ticker: str) -> list:
    """
    Extrait la section '📊 DONNÉES FINANCIÈRES STRUCTURÉES & INTERPRÉTÉES'
    du rapport source pour un ticker donné.
    Retourne list[ (type, element) ] : type = 'p' ou 'tbl'
    Inclut tous les paragraphes (titres de section) et tous les tableaux.
    S'arrête au premier Heading3 qui n'est pas dans la section données financières.
    """
    ticker_start, ticker_end, elements = _extract_company_section(source_doc, ticker)
    if ticker_start is None:
        return []

    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s
    def _ps(el):
        pPr = el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        return pPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '') if pPr is not None else ''

    blocks = []
    in_section = False

    for i in range(ticker_start, ticker_end):
        child = elements[i]
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            s = _ps(child)
            t = _dd(_pt(child).strip())
            if not t: continue

            # Détecter le début de la section données financières
            if 'DONNÉES FINANCIÈRES STRUCTURÉES' in t.upper() or 'DONNÉES FINANCIERES STRUCTURÉES' in t.upper():
                in_section = True
                blocks.append(('p', child))
                continue

            if in_section:
                # Arrêter aux sections suivantes
                # — Heading3 hors sections de données financières
                # — Textes qui marquent le début des PARTIES narratives
                if s in ('Heading3', 'Titre3') and not any(
                    w in t for w in ('BILAN', 'COMPTE', 'CASH', 'RATIO', 'STRUCTURE', 'DÉLAI', 'DONNÉE')
                ):
                    break
                # Arrêter si on tombe sur les PARTIES narratives (0-4) ou doublon rapport
                if any(m in t.upper() for m in ('PARTIE 0', 'PARTIE 1', 'PARTIE 2',
                                                  'PARTIE 3', 'PARTIE 4', 'ABSOLUMENT',
                                                  'RAPPORT D\'ANALYSE', '### **RAPPORT')):
                    break
                # Garder paragraphes de labels (📌 1. BILAN...) et textes
                blocks.append(('p', child))

        elif tag == 'tbl' and in_section:
            blocks.append(('tbl', child))

    return blocks


def _extract_parties(source_doc, ticker: str) -> dict:
    """
    Extrait les textes PARTIE 0, 1, 2, 3, 4 du rapport source pour un ticker.
    Structure source :
      **PARTIE 0 : INDICATEURS DE VALORISATION BOURSIÈRE**
      texte (1 paragraphe)
      **PARTIE 1 : ANALYSE DU COURS — STATISTIQUES ET ÉVOLUTION (100 derniers jours)**
      texte (1 paragraphe)
      **PARTIE 2 : ANALYSE TECHNIQUE DÉTAILLÉE**
      texte... (PLUSIEURS paragraphes : MM, Bollinger, MACD, RSI, Stoch, Conclusion)
      **PARTIE 3 : ANALYSE FONDAMENTALE (SECTION CRITIQUE)**
      texte (1-2 paragraphes)
      **PARTIE 4 : CONCLUSION D'INVESTISSEMENT**
      texte (plusieurs paragraphes)

    Retourne dict : {'p0': str, 'p1': str, 'p2': str, 'p3': str, 'p4': str}
    SANS limite de longueur — texte complet tel que dans le rapport source.
    """
    ticker_start, ticker_end, elements = _extract_company_section(source_doc, ticker)
    if ticker_start is None:
        return {}

    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s

    result = {'p0': '', 'p1': '', 'p2': '', 'p3': '', 'p4': ''}
    current_part = None

    for i in range(ticker_start, ticker_end):
        child = elements[i]
        if child.tag.split('}')[-1] != 'p': continue
        raw = _pt(child).strip()
        txt = _dd(raw)
        if not txt: continue

        import re as _re2
        t_up = txt.upper()

        # Détecter les marqueurs de parties : le mot "PARTIE X" doit être en DÉBUT du texte
        # (ex: "### **PARTIE 0 : ..." ou "**PARTIE 0 :**")
        # On refuse les faux positifs type "La valorisation (Partie 0) est..."
        is_part_marker = bool(_re2.match(
            r'^\s*[#*\s]*PARTIE\s*([0-4])\s*[:\–\-]', txt, _re2.IGNORECASE
        ))

        if is_part_marker and 'PARTIE 0' in t_up and 'VALORISATION' in t_up:
            current_part = 'p0'; continue
        if is_part_marker and 'PARTIE 1' in t_up and 'COURS' in t_up:
            current_part = 'p1'; continue
        if is_part_marker and 'PARTIE 2' in t_up and 'TECHNIQUE' in t_up:
            current_part = 'p2'; continue
        if is_part_marker and 'PARTIE 3' in t_up and 'FONDAMENTALE' in t_up:
            current_part = 'p3'; continue
        if is_part_marker and 'PARTIE 4' in t_up and 'CONCLUSION' in t_up:
            current_part = 'p4'; continue

        # Arrêter aux sections non narratives
        stop_markers = ('📎', '📋', '· · ·', '────', '════', '---')
        if any(txt.startswith(m) for m in stop_markers):
            if current_part == 'p4':
                current_part = None
            continue
        if txt.startswith('[') and ('Télécharger' in txt or 'Année :' in txt):
            if current_part == 'p4':
                current_part = None
            continue
        # Arrêter PARTIE 4 sur les lignes de documents/rapports annexes
        if current_part == 'p4' and any(m in txt for m in (
            'Recommandation source :', '⚠️ Risques :', '🔮 Perspectives :',
            'Indicateurs financiers :', 'Points clés :', 'Année : 20'
        )):
            current_part = None
            continue

        if current_part:
            # Nettoyer markdown **gras** → texte
            clean = re.sub(r'[*]{1,3}([^*]+)[*]{1,3}', r'\1', txt)
            clean = re.sub(r'^[*#-]\s+', '', clean)
            clean = re.sub(r'\s{2,}', ' ', clean).strip()
            if clean and not clean.startswith('---') and not clean.startswith('Absolument'):
                sep = '\n' if current_part == 'p2' else ' '
                result[current_part] += (sep if result[current_part] else '') + clean

    return result


def _read_source_tbl_for_copy(tbl_el) -> list:
    """
    Lit un tableau source et retourne ses données pour copie.
    Retourne list[list[str]] dédupliqué.
    """
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s
    rows = []
    for tr in tbl_el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
        cells = [_dd(''.join(n.text or '' for n in tc.iter() if n.text).strip())
                 for tc in tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')]
        if cells:
            rows.append(cells)
    return rows



def _extract_source_indicators(source_doc, ticker: str) -> dict:
    """
    Extrait depuis le rapport source :
    - Le bêta RÉEL (depuis le tableau score risque : "β=0.0563 (Défensif)...")
    - Les valeurs numériques des indicateurs techniques :
        mm20_val, mm50_val, boll_inf, boll_sup, macd_val, macd_sig, rsi_val, stoch_k, stoch_d
    - Le texte PARTIE 4 (conclusion d'investissement / synthèse)

    Retourne dict avec toutes ces valeurs, ou dict vide si non trouvé.
    """
    import re as _re
    ticker_start, ticker_end, elements = _extract_company_section(source_doc, ticker)
    if ticker_start is None:
        return {}

    def _pt(el): return ''.join(n.text or '' for n in el.iter() if n.text)
    def _dd(s):
        n = len(s)
        for d in (3, 2):
            if n % d == 0 and s == s[:n//d]*d: return s[:n//d]
        return s
    def _read_tbl(tbl_el):
        rows = []
        for tr in tbl_el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
            cells = [_dd(''.join(n.text or '' for n in tc.iter() if n.text).strip())
                     for tc in tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')]
            if cells: rows.append(cells)
        return rows

    result = {
        'beta_reel': None,
        'mm20_val': None, 'mm50_val': None,
        'boll_inf': None, 'boll_sup': None,
        'macd_val': None, 'macd_sig': None,
        'rsi_val': None,
        'stoch_k': None, 'stoch_d': None,
        'partie4': '',
    }

    def _num(s):
        """Extrait un float depuis une string comme '9 091,75' ou '271.251'"""
        if not s: return None
        s2 = s.replace(' ', '').replace('\u00a0', '')
        # Format français : virgule = décimale, espace = milliers
        if ',' in s2 and '.' not in s2:
            s2 = s2.replace(',', '.')
        elif ',' in s2 and '.' in s2:
            # ex: "9,091.75" → enlever virgule de milliers
            s2 = s2.replace(',', '')
        try:
            return float(s2)
        except (ValueError, TypeError):
            return None

    in_partie4 = False
    score_found = False

    for i in range(ticker_start, ticker_end):
        child = elements[i]
        tag = child.tag.split('}')[-1]
        raw = _pt(child).strip() if tag == 'p' else ''
        txt = _dd(raw) if raw else ''

        # ── BÊTA depuis tableau score risque ──────────────────────────────────
        if tag == 'tbl' and result['beta_reel'] is None:
            data = _read_tbl(child)
            full = str(data)
            if 'êta' in full or 'Bêt' in full:
                for row in data:
                    cell = ' '.join(row)
                    m = _re.search(r'β\s*=\s*([-]?[\d.,]+)', cell)
                    if m:
                        result['beta_reel'] = _num(m.group(1))
                        break

        # ── Valeurs numériques indicateurs techniques ─────────────────────────
        if tag == 'p' and txt:
            t_low = txt.lower()

            # MM20 / MM50
            if 'mm20' in t_low or 'mm 20' in t_low or 'moyenne mobile' in t_low:
                m20 = _re.search(r'(?:MM20|M20)[^0-9]*?([\d\s]+[,.][\d]+)', txt, _re.IGNORECASE)
                m50 = _re.search(r'(?:MM50|M50)[^0-9]*?([\d\s]+[,.][\d]+)', txt, _re.IGNORECASE)
                if m20 and result['mm20_val'] is None:
                    result['mm20_val'] = _num(m20.group(1))
                if m50 and result['mm50_val'] is None:
                    result['mm50_val'] = _num(m50.group(1))

            # Bollinger
            if 'bollinger' in t_low or 'bande' in t_low or 'borne' in t_low:
                b_inf = _re.search(r'(?:inf[eé]rieure?|basse?|basse des bandes)[^0-9]*?([\d\s]+[,.][\d]+)', txt, _re.IGNORECASE)
                b_sup = _re.search(r'(?:sup[eé]rieure?|haute?|haute des bandes)[^0-9]*?([\d\s]+[,.][\d]+)', txt, _re.IGNORECASE)
                # Alternative : deux nombres consécutifs pour borne inf/sup
                if not b_inf and not b_sup:
                    nums = _re.findall(r'([\d]{4,6}[,.][\d]+)', txt)
                    if len(nums) >= 2:
                        b_inf_v = _num(nums[0])
                        b_sup_v = _num(nums[-1])
                        if b_inf_v and b_sup_v and b_inf_v < b_sup_v:
                            if result['boll_inf'] is None: result['boll_inf'] = b_inf_v
                            if result['boll_sup'] is None: result['boll_sup'] = b_sup_v
                if b_inf and result['boll_inf'] is None:
                    result['boll_inf'] = _num(b_inf.group(1))
                if b_sup and result['boll_sup'] is None:
                    result['boll_sup'] = _num(b_sup.group(1))

            # MACD
            if 'macd' in t_low:
                m_macd = _re.search(r'MACD\s*[\(\[]?([\d\s]+[,.][\d]+)', txt, _re.IGNORECASE)
                m_sig  = _re.search(r'(?:signal|ligne de signal)\s*[\(\[]?([\d\s]+[,.][\d]+)', txt, _re.IGNORECASE)
                if m_macd and result['macd_val'] is None:
                    result['macd_val'] = _num(m_macd.group(1))
                if m_sig and result['macd_sig'] is None:
                    result['macd_sig'] = _num(m_sig.group(1))

            # RSI
            if 'rsi' in t_low:
                m_rsi = _re.search(r'RSI\s*(?:à|de|est|=|:)?\s*(?:une valeur de\s*)?([\d]+[,.][\d]+)', txt, _re.IGNORECASE)
                if m_rsi and result['rsi_val'] is None:
                    result['rsi_val'] = _num(m_rsi.group(1))

            # Stochastique
            if 'stoch' in t_low or '%k' in t_low:
                m_k = _re.search(r'%K\s*[\(\[]?([\d]+[,.][\d]+)', txt, _re.IGNORECASE)
                m_d = _re.search(r'%D\s*[\(\[]?([\d]+[,.][\d]+)', txt, _re.IGNORECASE)
                if m_k and result['stoch_k'] is None:
                    result['stoch_k'] = _num(m_k.group(1))
                if m_d and result['stoch_d'] is None:
                    result['stoch_d'] = _num(m_d.group(1))

            # PARTIE 4 — maintenant gérée par _extract_parties (plus complète)
            if 'PARTIE 4' in txt.upper() and 'CONCLUSION' in txt.upper():
                in_partie4 = True
                continue
            if in_partie4:
                def _ps2(el):
                    pPr = el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                    return pPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '') if pPr is not None else ''
                s = _ps2(child)
                if s in ('Heading3','Titre3','Heading2','Titre2') or txt.startswith('📎') or txt.startswith('📋') or txt.startswith('['):
                    in_partie4 = False
                elif txt and not txt.startswith('---') and not txt.startswith('Absolument'):
                    clean4 = _re.sub(r'[*#]+', '', txt).strip()
                    # Dédupliquer : ne pas ajouter si déjà présent (texte triplé dans source)
                    if clean4 and clean4 not in result['partie4']:
                        result['partie4'] += ('\n' if result['partie4'] else '') + clean4

    return result


def _build_fiche_docx(s: dict, date_str: str, freq: str = "JOUR",
                      period_info: dict = None,
                      images_map: dict | None = None,
                      source_doc=None) -> bytes:
    """
    Assemble la fiche Word d'une société.

    Paramètre supplémentaire
    ------------------------
    images_map : dict[ticker -> png_bytes] produit par _extract_images_from_docx().
                 Si None ou si le ticker de cette société n'y figure pas, on bascule
                 sur le graphique matplotlib de secours.
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    m = Cm(_MARGIN_CM)
    section.top_margin = m
    section.bottom_margin = m
    section.left_margin = m
    section.right_margin = m

    if doc.paragraphs:
        p0 = doc.paragraphs[0]
        p0.clear()
        _cp(p0, 0, 0)

    # Récupérer le PNG source pour ce ticker (None si absent)
    ticker = str(s.get("ticker") or "").strip().upper()
    source_png = (images_map or {}).get(ticker)

    # ── Extraire et injecter les indicateurs RÉELS depuis le rapport source ───
    # Bêta réel, MM20/MM50, Bollinger, MACD, RSI, Stoch, PARTIE 4
    if source_doc is not None and s.get("_src_indicators") is None:
        _si = _extract_source_indicators(source_doc, ticker)
        s["_src_indicators"] = _si
        # Bêta réel (ex: 0.0563 au lieu de 1.0 du LLM)
        if _si.get("beta_reel") is not None:
            s["beta"] = f"{_si['beta_reel']:.4f}"
        # Valeurs numériques indicateurs techniques
        for _s_key, _i_key in [
            ("mm20_valeur",    "mm20_val"),
            ("mm50_valeur",    "mm50_val"),
            ("boll_inf",       "boll_inf"),
            ("boll_sup",       "boll_sup"),
            ("macd_valeur",    "macd_val"),
            ("macd_signal_line", "macd_sig"),
            ("rsi_valeur",     "rsi_val"),
            ("stoch_k",        "stoch_k"),
            ("stoch_d",        "stoch_d"),
        ]:
            if s.get(_s_key) is None and _si.get(_i_key) is not None:
                s[_s_key] = _si[_i_key]

    build_header(doc, s, date_str)           # En-tête : ticker, score, reco, indicateurs
    _add_separator(doc)
    build_market_table(doc, s)               # Métriques de marché
    _add_separator(doc)
    build_chart_comment(doc, s, source_png)  # Graphique source Word + bandeau annoté
    _add_separator(doc)
    build_technical_analysis(doc, s)         # Analyse technique (tableau + convergence)
    _add_separator(doc)
    build_fundamental_analysis(doc, s)       # Analyse fondamentale + risques + perspectives
    _add_separator(doc)
    build_risques(doc, s, source_doc)        # Profil de risque (source rapport + LLM)
    _add_separator(doc)
    build_financial_analysis(doc, s)         # Analyse financière détaillée
    _add_separator(doc)
    build_financial_data_complete(doc, s, source_doc)  # 📊 Données financières (copie source)
    _add_separator(doc)
    build_analyse_fondamentale_partie3(doc, s, source_doc)  # PARTIE 3 fondamentale
    _add_separator(doc)
    build_conclusion(doc, s)                 # Conclusion : matrice + divergences + action
    _pied(doc, date_str, freq, period_info)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Point d'entrée ────────────────────────────────────────────────────────────

def generate(docs_bytes, freq: str = "JOUR", period_info: dict = None) -> list:
    """
    Génère une fiche Word par société depuis un ou plusieurs .docx source.

    Pipeline :
      0. Extraction des images du document source (nouveau)
      1. Extraction du texte
      2. Extraction LLM (JSON minimal)
      3. Enrichissement Python + génération Word

    docs_bytes : bytes (un seul doc) ou list[bytes] (plusieurs docs, plus récent en premier).
    Retourne list de (filename: str, docx_bytes: bytes).
    """
    if isinstance(docs_bytes, bytes):
        docs_bytes = [docs_bytes]

    date_str  = date.today().strftime("%d/%m/%Y")
    date_file = date.today().strftime("%Y%m%d")
    freq_suffix = {"JOUR": "JOUR", "HEBDO": "HEBDO", "MENSUEL": "MENSUEL",
                   "TRIM": "TRIM", "ANNUEL": "ANNUEL"}.get(freq, freq)

    # ── Étape 0 : extraction des images + ouverture source_doc ───────────────
    # On ouvre le document source une seule fois pour :
    #   - extraire les graphiques (images embarquées)
    #   - extraire les textes de risque par société (section narrative)
    print(f"  [Fiches/{freq}] Étape 0/3 : Extraction des images du document source...")
    images_map: dict = {}
    source_doc = None
    try:
        from docx import Document as _DocxDocument
        source_doc = _DocxDocument(io.BytesIO(docs_bytes[0]))
        images_map = _extract_images_from_docx(docs_bytes[0])
        print(f"  [Fiches/{freq}] {len(images_map)} graphique(s) extrait(s) : "
              f"{list(images_map.keys()) or '(aucun)'}")
    except Exception as exc:
        print(f"  [Fiches/{freq}] AVERTISSEMENT extraction images : {exc}")

    # ── Étape 1 : extraction du texte ─────────────────────────────────────────
    print(f"  [Fiches/{freq}] Étape 1/3 : Extraction du texte ({len(docs_bytes)} doc(s))...")
    full_text = _build_context(docs_bytes, freq)
    print(f"  [Fiches/{freq}] Texte source : {len(full_text)} chars")

    # ── Étape 2 : extraction LLM ──────────────────────────────────────────────
    print(f"  [Fiches/{freq}] Étape 2/3 : Extraction LLM (JSON minimal)...")
    raw_companies = extract_all(full_text, freq, period_info)
    print(f"  [Fiches/{freq}] LLM → {len(raw_companies)} société(s) extraite(s).")

    if not raw_companies:
        print(f"  [Fiches/{freq}] ERREUR : extraction invalide ou aucune société — "
              "abandon (voir logs Extractor).")
        return []

    # Étape 0b : ré-extraction TOUJOURS avec les tickers LLM confirmés.
    # La première passe (sans filtre) peut avoir associé de faux tickers
    # (ACHAT, RISQUE, CV...). On réextrait maintenant avec la liste stricte
    # des 47 sociétés pour garantir la bonne association image ↔ ticker.
    known = [str(c.get("ticker") or "").strip().upper() for c in raw_companies if c.get("ticker")]
    if known:
        print(f"  [Fiches/{freq}] Étape 0b : ré-extraction images avec {len(known)} tickers LLM...")
        try:
            images_map = _extract_images_from_docx(docs_bytes[0], known_tickers=known)
            print(f"  [Fiches/{freq}] {len(images_map)}/{len(known)} graphique(s) : "
                  f"{list(images_map.keys()) or '(aucun)'}")
        except Exception as exc:
            print(f"  [Fiches/{freq}] AVERTISSEMENT ré-extraction : {exc}")

    # ── Étape 3 : enrichissement + génération Word ────────────────────────────
    print(f"  [Fiches/{freq}] Étape 3/3 : Enrichissement Python + génération Word...")
    all_companies = enrich(raw_companies)
    print(f"  [Fiches/{freq}] Enrichissement → {len(all_companies)} société(s).")

    results = []
    for company in all_companies:
        ticker = str(company.get("ticker") or "").strip()
        if not ticker:
            print(f"  [Fiches/{freq}] SKIP : société sans ticker")
            continue
        has_source_img = ticker.upper() in images_map
        try:
            # Injecter le source_doc dans le dict société pour accès dans build_chart_comment
            company["_source_doc_ref"] = source_doc
            docx_bytes = _build_fiche_docx(
                company, date_str, freq, period_info,
                images_map=images_map,
                source_doc=source_doc,
            )
            filename = f"Fiche_{ticker}_{date_file}_{freq_suffix}.docx"
            results.append((filename, docx_bytes))
            img_status = "✓ graphique source Word" if has_source_img else "~ graphique matplotlib"
            print(f"  [Fiches/{freq}] ✓ {filename}  [{img_status}]")
        except Exception as e:
            print(f"  [Fiches/{freq}] AVERTISSEMENT : fiche {ticker} ignorée — {e}")

    print(f"  [Fiches/{freq}] TOTAL : {len(results)} fiche(s) générée(s).")
    return results
