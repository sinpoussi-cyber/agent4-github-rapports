"""
period_synthesis.py — Synthèse d'évolution multi-périodes pour la Note Stratégique
==================================================================================
La Note Stratégique (generate_note_strategique.py) ne lit que le document le plus
récent : mensuel et annuel ne sont donc, aujourd'hui, que le rapport du jour avec
un autre titre. Ce module ajoute une VRAIE synthèse d'évolution sur toute la
période, produite par le LLM à partir du texte de CHAQUE document, puis :

  - insérée en tête d'une note existante (HEBDO / MENSUEL / TRIM), via insert_synthesis()
  - ou rendue comme note autonome (ANNUEL), via build_annual_note()

Tout est défensif : en cas d'échec, on renvoie la note d'origine inchangée
(jamais d'exception propagée à l'appelant).
"""

import io

from llm_client import call as llm_call

_FREQ_PHRASE = {
    "HEBDO": "de la semaine écoulée",
    "MENSUEL": "du mois écoulé",
    "TRIM": "du trimestre écoulé",
    "ANNUEL": "de l'année écoulée",
}

_SECTION_HEADS = {
    "TENDANCE GÉNÉRALE",
    "FAITS MARQUANTS DE LA PÉRIODE",
    "ROTATION SECTORIELLE",
    "RISQUES & POINTS DE VIGILANCE",
    "RISQUES ET POINTS DE VIGILANCE",
    "PERSPECTIVES",
}


# ── Extraction texte ──────────────────────────────────────────────────────────

def _extract_text(doc_bytes):
    from docx import Document
    doc = Document(io.BytesIO(doc_bytes))
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _sample(n, max_docs):
    """Indices d'un échantillon uniforme couvrant TOUTE la série [0..n-1]."""
    if n <= max_docs:
        return list(range(n))
    step = (n - 1) / (max_docs - 1)
    return sorted({round(i * step) for i in range(max_docs)})


# ── Génération de la synthèse (texte) ─────────────────────────────────────────

def build_synthesis_text(docs_bytes, freq, period_info=None, is_notes=False):
    """
    Produit le texte de synthèse (str), ou "" si indisponible.
    docs_bytes : liste de bytes .docx en ordre chronologique CROISSANT (ancien→récent).
    is_notes   : True si les documents sont des notes mensuelles (cas annuel).
    """
    if not docs_bytes:
        return ""

    max_docs = {"HEBDO": 7, "MENSUEL": 12, "TRIM": 14, "ANNUEL": 12}.get(freq, 8)
    idx = _sample(len(docs_bytes), max_docs)
    per_doc = min(60000 // max(len(idx), 1), 6000)

    blocks = []
    for pos in idx:
        try:
            txt = _extract_text(docs_bytes[pos])[:per_doc]
        except Exception:
            continue
        blocks.append(f"===== ÉLÉMENT {pos + 1} =====\n{txt}")
    corpus = "\n\n".join(blocks)[:90000]
    if not corpus.strip():
        return ""

    pi = period_info or {}
    periode = f"{pi.get('date_debut', '?')} → {pi.get('date_fin', '?')}"
    nb = pi.get("nb_seances", len(docs_bytes))
    quoi = "notes stratégiques mensuelles" if is_notes else "rapports journaliers BRVM"
    phrase = _FREQ_PHRASE.get(freq, "de la période")

    prompt = f"""Tu es un analyste financier senior de la BRVM (marché régional UEMOA).
On te fournit une SÉRIE de {quoi} couvrant la période {periode} ({nb} éléments,
en ordre chronologique croissant). Rédige une SYNTHÈSE {phrase} centrée sur
l'ÉVOLUTION sur TOUTE la période — surtout PAS un instantané du dernier jour.
Compare le début et la fin de période, identifie les tendances et les ruptures.

Réponds en français, sans markdown ni astérisques, avec EXACTEMENT ces intitulés
de section (en majuscules, seuls sur leur ligne) :

TENDANCE GÉNÉRALE
Trois à quatre phrases : direction du BRVM Composite et de la capitalisation sur
la période, ampleur du mouvement début→fin (chiffré si possible), points hauts/bas.

FAITS MARQUANTS DE LA PÉRIODE
Quatre à six lignes commençant chacune par «- » : plus fortes hausses/baisses,
mouvements sectoriels, ruptures de tendance, événements notables datés.

ROTATION SECTORIELLE
Deux à trois phrases : secteurs en surperformance vs sous-performance sur la période.

RISQUES & POINTS DE VIGILANCE
Trois à quatre lignes commençant chacune par «- ».

PERSPECTIVES
Deux à trois phrases prospectives et prudentes, sans recommandation d'achat ou de
vente nominative.

DOCUMENTS SOURCES (ordre chronologique croissant) :
{corpus}
"""
    try:
        return llm_call(
            prompt,
            max_tokens=1800,
            system="Analyste BRVM/UEMOA. Synthèse d'évolution multi-périodes, factuelle et prudente.",
        ).strip()
    except Exception as e:
        print(f"  [Synthèse/{freq}] AVERTISSEMENT : synthèse LLM indisponible ({e}).")
        return ""


# ── Rendu Word ────────────────────────────────────────────────────────────────

def _write_synthesis(doc, synthesis_text, freq, period_info, collect_paras=None):
    """Écrit les paragraphes de synthèse dans `doc`. Si collect_paras est une liste,
    les objets Paragraph créés y sont ajoutés (pour repositionnement ultérieur)."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1A, 0x23, 0x7E)
    GREEN = RGBColor(0x0F, 0x9D, 0x58)
    GREY = RGBColor(0x66, 0x66, 0x66)
    LGREY = RGBColor(0x99, 0x99, 0x99)

    label = {
        "HEBDO": "SYNTHÈSE HEBDOMADAIRE — ÉVOLUTION DE LA PÉRIODE",
        "MENSUEL": "SYNTHÈSE MENSUELLE — ÉVOLUTION DE LA PÉRIODE",
        "TRIM": "SYNTHÈSE TRIMESTRIELLE — ÉVOLUTION DE LA PÉRIODE",
        "ANNUEL": "SYNTHÈSE ANNUELLE — ÉVOLUTION DE LA PÉRIODE",
    }.get(freq, "SYNTHÈSE DE LA PÉRIODE")
    pi = period_info or {}
    sub = (f"Période : {pi.get('date_debut', '—')} → {pi.get('date_fin', '—')}"
           f"   •   {pi.get('nb_seances', '—')} élément(s) analysé(s)")

    def mk(text, size=10, bold=False, color=None, heading=False,
           align=None, sb=1, sa=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold or heading
        run.font.size = Pt(13 if heading else size)
        if heading:
            run.font.color.rgb = NAVY
        elif color:
            run.font.color.rgb = color
        pf = p.paragraph_format
        pf.space_before = Pt(sb)
        pf.space_after = Pt(sa)
        if align:
            p.alignment = align
        if collect_paras is not None:
            collect_paras.append(p)
        return p

    mk(label, heading=True, sb=10, sa=2)
    mk(sub, size=8, color=GREY, sa=8)

    for raw in synthesis_text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.strip().upper() in _SECTION_HEADS:
            mk(line.strip(), size=11, bold=True, color=GREEN, sb=8, sa=2)
        else:
            mk(line, size=10, sb=1, sa=1)

    mk("— — —", size=9, color=LGREY, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=8)


def insert_synthesis(note_bytes, synthesis_text, freq, period_info=None):
    """Réouvre la note .docx et insère la synthèse juste après le titre principal.
    Retourne de nouveaux bytes ; renvoie note_bytes inchangé si échec/synthèse vide."""
    if not synthesis_text:
        return note_bytes
    try:
        from docx import Document

        doc = Document(io.BytesIO(note_bytes))

        # Ancre : premier paragraphe non vide (généralement le titre)
        anchor = None
        for p in doc.paragraphs:
            if p.text.strip():
                anchor = p
                break

        created = []
        _write_synthesis(doc, synthesis_text, freq, period_info, collect_paras=created)

        # Repositionner les paragraphes (ajoutés en fin de corps) après l'ancre
        if anchor is not None and created:
            ref = anchor._p
            for p in created:
                ref.addnext(p._p)
                ref = p._p

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        print(f"  [Synthèse/{freq}] AVERTISSEMENT : insertion impossible ({e}).")
        return note_bytes


def build_annual_note(monthly_notes, year, period_info=None):
    """Construit une note ANNUELLE autonome à partir des notes mensuelles archivées.
    monthly_notes : liste de {nom, contenu_bytes} (ordre chronologique croissant).
    Retourne (filename, docx_bytes). Lève si aucune note fournie."""
    if not monthly_notes:
        raise ValueError("Aucune note mensuelle fournie pour le rapport annuel.")

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    docs = [m["contenu_bytes"] for m in monthly_notes]
    synth = build_synthesis_text(docs, "ANNUEL", period_info, is_notes=True)

    doc = Document()

    title = doc.add_paragraph()
    r = title.add_run(f"NOTE STRATÉGIQUE ANNUELLE BRVM — {year}")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    st = doc.add_paragraph()
    rs = st.add_run(f"Bilan consolidé sur {len(monthly_notes)} note(s) mensuelle(s)")
    rs.font.size = Pt(10)
    rs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(12)

    if synth:
        _write_synthesis(doc, synth, "ANNUEL", period_info)
    else:
        p = doc.add_paragraph()
        p.add_run("Synthèse annuelle indisponible (aucune donnée exploitable "
                  "extraite des notes mensuelles).").italic = True

    recap = doc.add_paragraph()
    rr = recap.add_run("NOTES MENSUELLES CONSOLIDÉES")
    rr.bold = True
    rr.font.size = Pt(12)
    rr.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    recap.paragraph_format.space_before = Pt(10)
    for m in monthly_notes:
        b = doc.add_paragraph(style=None)
        b.add_run(f"• {m['nom']}").font.size = Pt(9)
        b.paragraph_format.space_after = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    filename = f"Note_Strategique_BRVM_ANNUEL_{year}.docx"
    return filename, buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  RAPPORT STRUCTURÉ PAR SECTION — reprend TOUTES les parties du rapport
#  journalier, chacune résumée avec son évolution sur la période.
#  Les titres sont rendus en style "Heading1" (comme le rapport source), donc
#  un rapport annuel peut re-parser les rapports mensuels de la même façon
#  (récursion jour → semaine/mois → année).
# ══════════════════════════════════════════════════════════════════════════════

from docx.oxml.ns import qn  # noqa: E402

# Doit rester IDENTIQUE aux ancres de generate_note_strategique.py (même ordre,
# pour reproduire la précédence de détection des sections).
_SOURCE_ANCHORS = {
    "synthese":           "SYNTHÈSE GÉNÉRALE",
    "secteurs":           "ANALYSE PAR SECTEUR",
    "matrice_signaux":    "MATRICE DE CONVERGENCE DES SIGNAUX",
    "liquidite":          "ANALYSE DE LIQUIDITÉ",
    "top_divergences":    "TOP 10 DES DIVERGENCES MAJEURES",
    "matrice_risque":     "MATRICE RISQUE vs HORIZON",
    "macro":              "ANALYSE MACRO",
    "macro_actu":         "1. ACTUALITÉS MACRO",
    "macro_pol":          "2. ACTUALITÉS POLITIQUES",
    "macro_fin":          "3. ACTUALITÉS FINANCIÈRES",
    "macro_synth":        "SYNTHÈSE & RECOMMANDATION FINALE",
    "actualites":         "ACTUALITÉS DU MARCHÉ BRVM",
    "classement":         "CLASSEMENT DES SOCIÉTÉS",
    "portefeuilles":      "PORTEFEUILLES MODÈLES",
    "alertes":            "ALERTES DU JOUR",
    "recap_risques":      "RÉCAPITULATIF DES RISQUES",
    "toc_detail":         "TABLE DES MATIÈRES - ANALYSES DÉTAILLÉES",
    "predictions":        "PRÉDICTIONS",
    "analyse_financiere": "ANALYSE FINANCIÈRE",
}

# Sections rendues dans le rapport périodique (ordre du rapport journalier).
# On saute la table des matières (navigation).
_ORDERED_SECTIONS = [
    ("synthese",           "SYNTHÈSE GÉNÉRALE"),
    ("secteurs",           "ANALYSE PAR SECTEUR"),
    ("matrice_signaux",    "MATRICE DE CONVERGENCE DES SIGNAUX"),
    ("liquidite",          "ANALYSE DE LIQUIDITÉ"),
    ("top_divergences",    "TOP 10 DES DIVERGENCES MAJEURES"),
    ("matrice_risque",     "MATRICE RISQUE vs HORIZON"),
    ("macro",              "ANALYSE MACRO"),
    ("macro_actu",         "1. ACTUALITÉS MACRO"),
    ("macro_pol",          "2. ACTUALITÉS POLITIQUES"),
    ("macro_fin",          "3. ACTUALITÉS FINANCIÈRES"),
    ("macro_synth",        "SYNTHÈSE & RECOMMANDATION FINALE"),
    ("actualites",         "ACTUALITÉS DU MARCHÉ BRVM"),
    ("classement",         "CLASSEMENT DES SOCIÉTÉS"),
    ("portefeuilles",      "PORTEFEUILLES MODÈLES"),
    ("alertes",            "ALERTES DU JOUR"),
    ("recap_risques",      "RÉCAPITULATIF DES RISQUES"),
    ("predictions",        "PRÉDICTIONS"),
    ("analyse_financiere", "ANALYSE FINANCIÈRE"),
]

_FREQ_TITLE = {"HEBDO": "HEBDOMADAIRE", "MENSUEL": "MENSUEL",
               "TRIM": "TRIMESTRIEL", "ANNUEL": "ANNUEL"}


# ── Découpage par section (réplique exacte de _split_source_by_h1) ────────────

def _para_style_val(p_elem):
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return ""
    pStyle = pPr.find(qn("w:pStyle"))
    return pStyle.get(qn("w:val")) if pStyle is not None else ""


def _para_text_xml(p_elem):
    return "".join(t.text for t in p_elem.iter(qn("w:t")) if t.text)


def _tbl_text_xml(tbl_elem):
    lines = []
    for tr in tbl_elem.iter(qn("w:tr")):
        cells = []
        for tc in tr.iter(qn("w:tc")):
            txt = "".join(t.text for t in tc.iter(qn("w:t")) if t.text).strip()
            if txt:
                cells.append(txt)
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _split_by_h1_text(doc):
    """{anchor_key: texte_aplati} — même logique de détection que le module source
    (paragraphe de style 'Heading1' contenant l'ancre)."""
    body = doc.element.body
    buckets = {k: [] for k in _SOURCE_ANCHORS}
    current = None
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag not in ("p", "tbl"):
            continue
        if tag == "p" and _para_style_val(child) == "Heading1":
            txt = _para_text_xml(child)
            current = None
            for k, anchor in _SOURCE_ANCHORS.items():
                if anchor in txt:
                    current = k
                    break
            continue
        if current is not None:
            t = _para_text_xml(child).strip() if tag == "p" else _tbl_text_xml(child)
            if t:
                buckets[current].append(t)
    return {k: "\n".join(v) for k, v in buckets.items() if v}


# ── Résumé d'une section sur la période ───────────────────────────────────────

def _summarize_section(title, series, freq, periode, total):
    """series : liste (position, texte_section). Retourne un résumé d'évolution."""
    per_day = 1400
    blocks = [f"[Point {i + 1}]\n{txt[:per_day]}" for i, (_pos, txt) in enumerate(series)]
    corpus = "\n\n".join(blocks)[:16000]
    prompt = f"""Tu es analyste financier de la BRVM (marché UEMOA). Voici la section
« {title} » du rapport, capturée à {len(series)} dates de la période {periode}
(ordre chronologique, sur {total} document(s)). Rédige un RÉSUMÉ de cette section
pour la PÉRIODE, centré sur l'ÉVOLUTION : variations début→fin, tendances,
indicateurs clés et leur progression, éléments apparus ou disparus. Reste factuel.
4 à 8 lignes maximum, en français, sans markdown ni astérisques ; utilise «- » pour
d'éventuelles puces.

CONTENU CHRONOLOGIQUE :
{corpus}
"""
    try:
        return llm_call(
            prompt, max_tokens=700,
            system="Analyste BRVM/UEMOA. Résumé de section, évolution sur la période, factuel et concis.",
        ).strip()
    except Exception as e:
        print(f"  [Section/{freq}] AVERTISSEMENT « {title} » : {e}")
        return ""


# ── Constructeur du rapport périodique structuré ──────────────────────────────

def build_period_report(docs_bytes_asc, freq, period_info=None, is_notes=False):
    """Construit un rapport {freq} reprenant TOUTES les sections du rapport
    journalier, chacune résumée avec son évolution sur la période.
    docs_bytes_asc : bytes .docx en ordre chronologique CROISSANT.
    Retourne (filename, docx_bytes). Lève si aucun document."""
    if not docs_bytes_asc:
        raise ValueError("Aucun document fourni pour le rapport périodique.")

    import io as _io
    from datetime import date as _date
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pi = period_info or {}
    periode = f"{pi.get('date_debut', '?')} → {pi.get('date_fin', '?')}"
    total = len(docs_bytes_asc)
    max_days = {"HEBDO": 7, "MENSUEL": 10, "TRIM": 12, "ANNUEL": 12}.get(freq, 8)
    idx = _sample(total, max_days)

    parsed = []
    for i in idx:
        try:
            parsed.append(_split_by_h1_text(Document(_io.BytesIO(docs_bytes_asc[i]))))
        except Exception:
            parsed.append({})

    NAVY = RGBColor(0x1A, 0x23, 0x7E)
    GREY = RGBColor(0x66, 0x66, 0x66)
    doc = Document()

    ftitle = _FREQ_TITLE.get(freq, freq)
    t = doc.add_paragraph()
    r = t.add_run(f"RAPPORT {ftitle} BRVM")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    quoi = "note(s) mensuelle(s)" if is_notes else "rapport(s) journalier(s)"
    st = doc.add_paragraph()
    rs = st.add_run(f"Synthèse par section — évolution sur la période   •   "
                    f"{periode}   •   {total} {quoi}")
    rs.font.size = Pt(10)
    rs.font.color.rgb = GREY
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(12)

    # 1) Synthèse exécutive : tendances d'évolution sur la période
    try:
        exec_txt = build_synthesis_text(docs_bytes_asc, freq, pi, is_notes=is_notes)
        if exec_txt:
            _write_synthesis(doc, exec_txt, freq, pi)
    except Exception as e:
        print(f"  [Rapport/{freq}] AVERTISSEMENT synthèse exécutive : {e}")

    # 2) Chaque section du rapport journalier, résumée avec évolution
    n_sections = 0
    for key, title in _ORDERED_SECTIONS:
        series = [(pos, parsed[j].get(key, "").strip())
                  for j, pos in enumerate(idx) if parsed[j].get(key, "").strip()]
        if not series:
            continue
        summary = _summarize_section(title, series, freq, periode, total)

        h = doc.add_paragraph(title)
        h.style = doc.styles["Heading 1"]  # -> pStyle "Heading1" (re-parsable)

        if summary:
            for line in summary.splitlines():
                line = line.rstrip()
                if not line:
                    continue
                p = doc.add_paragraph()
                p.add_run(line).font.size = Pt(10)
                p.paragraph_format.space_after = Pt(1)
        else:
            p = doc.add_paragraph()
            p.add_run("Résumé indisponible pour cette section.").italic = True
        n_sections += 1

    if n_sections == 0:
        p = doc.add_paragraph()
        p.add_run("Aucune section reconnue dans les documents source "
                  "(vérifier la structure des rapports).").italic = True

    stamp = str(pi.get("annee")) if (freq == "ANNUEL" and pi.get("annee")) \
        else _date.today().strftime("%Y%m%d")
    filename = f"Note_Strategique_BRVM_{freq}_{stamp}.docx"

    buf = _io.BytesIO()
    doc.save(buf)
    return filename, buf.getvalue()
